"""
Erzeugt eine kompakte DIN-A4-Seite als Word-Datei:
Welche Dateien wurden verwendet, welche nicht, welche vielleicht?
"""
import os
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Projekt-Root (eine Ebene über /scripts) ───────────────────────
PROJEKT_ROOT = Path(__file__).resolve().parent.parent
os.chdir(PROJEKT_ROOT)

# ── Farben ────────────────────────────────────────────────────────
C_WEISS = RGBColor(0xFF, 0xFF, 0xFF)
F_GRUEN = "E2EFDA"   # verwendet
F_GELB  = "FFF2CC"   # moeglicherweise
F_ROT   = "FCE4D6"   # nicht verwendet
F_GRAU  = "F2F2F2"   # Lookup/Key

def set_cell_bg(cell, fill_hex):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)

def set_hdr_cell(cell, fill_hex):
    set_cell_bg(cell, fill_hex)
    for run in cell.paragraphs[0].runs:
        run.font.color.rgb = C_WEISS
        run.bold = True
        run.font.size = Pt(9)

def add_tbl(doc, headers, rows, cw, hdr_color, row_colors):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    # Header
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = h
        set_hdr_cell(cell, hdr_color)
    # Daten
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = tbl.rows[ri + 1].cells[ci]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(8.5)
            set_cell_bg(cell, row_colors[ri % len(row_colors)])
    # Breiten
    for row in tbl.rows:
        for i, w in enumerate(cw):
            row.cells[i].width = Cm(w)
    return tbl

# ── Dokument ──────────────────────────────────────────────────────
doc = Document()
doc.core_properties.author = "Datenanalyse-Team"
doc.core_properties.created = datetime.now()
doc.core_properties.modified = datetime.now()
for s in doc.sections:
    s.top_margin    = Cm(1.8)
    s.bottom_margin = Cm(1.5)
    s.left_margin   = Cm(2.0)
    s.right_margin  = Cm(1.5)
    s.page_width    = Cm(21.0)   # DIN A4
    s.page_height   = Cm(29.7)

# ── Begleitkommentar ─────────────────────────────────────────────
from datetime import date
kommentar_box = doc.add_paragraph()
rk = kommentar_box.add_run(
    f"An: Alle  |  Von: Datanalyse-Team  |  Datum: {date.today().strftime('%d.%m.%Y')}\n"
    "Betreff: Erste Analyse — Datei-Klassifikation des Qualitäts-Muster-Finder-Datensatzes\n\n"
    "Liebe Kollegen,\n\n"
    "anbei unsere erste Einschätzung, welche der 86 CSV-Dateien des Datensatzes für unser \n"
    "Projekt relevant sind — und warum.\n\n"
    "Hintergrund: Wir analysieren, ob Strukturmerkmale eines Krankenhauses \n"
    "(Bettenzahl, Träger, Personal, Region) mit der Häufigkeit seiner auffälligen \n"
    "Qualitätsindikatoren zusammenhängen. Datenquelle: offizielle Qualitätsberichte \n"
    "2023 aller deutschen Krankenhäuser (IQTIG / G-BA), 86 CSV-Dateien, ca. 1,2 GB.\n\n"
    "Was diese Übersicht zeigt:\n"
    "  ✅  Seite 1, Block 1: Dateien, die wir tatsächlich in die Analyse eingebunden haben\n"
    "  ⚠️  Seite 1, Block 2: Dateien, die potenziell interessant sind, aber noch nicht geprüft wurden\n"
    "  ❌  Seite 1, Block 3: Dateien, die wir bewusst ausgeschlossen haben\n"
    "  📄  Seite 2: Ausführliche Begründung für jede Entscheidung\n\n"
    "Wichtig: Dies ist eine erste Schätzung auf Basis einer systematischen Datei-Sichtung \n"
    "(Dateiname → Dateigröße → Spaltenheader → Beispielzeilen → Join-Schlüssel-Check). \n"
    "Die endgültige Entscheidung, welche weiteren Dateien eingebunden werden, kann sich \n"
    "im Projektverlauf noch ändern — insbesondere für die Dateien in Block ⚠️.\n\n"
    "Fragen, Ergänzungen oder Korrekturen sind sehr willkommen!"
)
rk.font.size = Pt(10)
pPr_k = kommentar_box._p.get_or_add_pPr()
shd_k = OxmlElement("w:shd"); shd_k.set(qn("w:fill"), "DEEAF1"); shd_k.set(qn("w:val"), "clear")
pPr_k.append(shd_k)
doc.add_paragraph().paragraph_format.space_after = Pt(6)

# ── Trennlinie ────────────────────────────────────────────────────
p_hr = doc.add_paragraph()
p_hr.add_run("─" * 110).font.size = Pt(8)
doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ── Titel ─────────────────────────────────────────────────────────
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
rn = t.add_run("Datei-Übersicht: Qualitäts-Muster-Finder")
rn.bold = True; rn.font.size = Pt(14); rn.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.add_run(
    "Datensatz: 86 CSV-Dateien · Berichtsjahr 2023 · Quelle: IQTIG / G-BA  |  "
    "Schlüssel: SO.QBID  |  ✅ verwendet   ⚠️ moeglicherweise   ❌ nicht relevant"
).font.size = Pt(8.5)
doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ══════════════════════════════════════════════════════════════════
# BLOCK 1: Verwendet (VERW)
# ══════════════════════════════════════════════════════════════════
p1 = doc.add_paragraph()
r1 = p1.add_run("✅  Verwendet — direkt in die Analyse eingeflossen")
r1.bold = True; r1.font.size = Pt(10); r1.font.color.rgb = RGBColor(0x37, 0x5E, 0x23)

add_tbl(doc,
    ["Datei", "Inhalt", "Verwendung im Projekt"],
    [
        ("SO.csv",                     "Stammdaten aller 2.310 Krankenhäuser",           "MERKMALE: Betten, Träger, Bundesland, Uni, Koordinaten, Standortnummer · Primärschlüssel SO.QBID"),
        ("QS.Qualitätsindikator.csv",  "~150 QI-Bewertungen pro Haus  (911 MB)",        "ZIEL-VARIABLE: auffaellig_quote > Median (76,92 %) → hat_viele_Probleme"),
        ("QS.Fortbildung.csv",         "Fortbildungsnachweise der Ärzte",                "MERKMAL: fortbildungsquote = Erbrachte / Pflichtige"),
        ("FA.csv",                     "Fachabteilungen (Brückentabelle)",               "Verbindet ABTID → FA.QBID = SO.QBID — nötig für Ärzte-Join"),
        ("FA.Personalliste.csv",       "Personal pro Fachabteilung  (14,6 MB)",          "MERKMAL: aerzte_pro_bett — stärkstes Merkmal (FI 53,6 %) im Decision Tree"),
        ("SO.Personalliste.csv",       "Personal auf Standortebene  (2026-07-29)",       "MERKMAL: pflege_pro_bett — 2. wichtigstes Merkmal (FI 23,8 %); bewusst statt AQ.Pflege.csv gewählt"),
        ("Konzern.csv",                "Konzernzugehörigkeit  (2026-07-29)",             "MERKMAL: ist_konzern — 358/1.824 Konzernhäuser, aber kein signif. Zusammenhang (Chi² p=0,90)"),
    ],
    cw=[4.0, 4.8, 9.2], hdr_color="375E23", row_colors=[F_GRUEN]
)
doc.add_paragraph().paragraph_format.space_after = Pt(3)

# ══════════════════════════════════════════════════════════════════
# BLOCK 2: Moeglicherweise (MOEG) — gruppiert
# ══════════════════════════════════════════════════════════════════
p2 = doc.add_paragraph()
r2 = p2.add_run("⚠️  Möglicherweise relevant — identifiziert, aber nicht eingebunden")
r2.bold = True; r2.font.size = Pt(10); r2.font.color.rgb = RGBColor(0xBF, 0x5A, 0x00)

add_tbl(doc,
    ["Gruppe", "Dateien", "Warum nicht eingebunden?"],
    [
        ("QS-Berichtsbasis",          "QS.csv",                                                          "Ursprünglich als notwendige Brückentabelle vermutet — wird tatsächlich nie geladen, da QS.Qualitätsindikator.csv SO.QBID bereits selbst trägt. QS.Typ (bund/land) könnte für eine Filterung genutzt werden, bisher nicht geprüft."),
        ("Leistungsbereiche",         "QS.Leistungsbereich.csv",                                         "QSLB.Dokumentationsrate = potenzielle Qualitätskennzahl. Weiterhin nicht eingebunden — aussichtsreichster Kandidat für eine künftige Ergänzung."),
        ("Externe QS-Ergebnisse",     "QS.Extern.Sonstige.csv",                                           "Keine einheitliche R*/N*-Bewertung — Überschneidung mit QS.Qualitätsindikator unklar"),
        ("Wettbewerbsumfeld",         "Mitbewerber_Betten.csv",                                          "Bettenzahl benachbarter Wettbewerber — Marktdichte-/Konkurrenzmerkmal, im gesamten Projekt bisher nicht in Betracht gezogen (entdeckt 2026-07-29)."),
        ("Mindestmengen  (MM.*)",     "MM.csv · MM.Ausnahme.csv · MM.Leistungsberechtigung.Prognose.csv", "Pot. Strukturmerkmal (Mindestmengen erfüllt?). Nicht verwendet."),
        ("Hygienedaten  (H*)",        "HB.csv · HD.csv (40 MB) · HM.csv · WeitereHygiene.csv",           "Mögliches Qualitätsmerkmal. HD.csv zu groß für direkte Sichtung. Nicht gesichtet."),
        ("Risikomanagement  (RM.*)",  "RM.csv · RM.Fallbesprechung.csv",                                  "Strukturmerkmal. Noch nicht gesichtet."),
        ("Ausstattung  (AM/AA/BM/BF/CQ)", "AM.csv · AM.Leistung · AM.VAVU · AA.csv · BM.csv · BF.csv · CQ.csv", "Strukturmerkmale. Bedeutung für Analyse noch nicht geprüft."),
        ("Personal / Qualifikation",  "AQ.Ärzte.csv · FA.Personen.csv · MP.csv", "Alternativquellen zu FA.Personalliste. AQ.Ärzte: Qualifikation, keine Anzahlen."),
        ("Psychiatrie-spezifisch",    "QS.Pso.csv · QS.Psy.csv · QS.Struktur.Station.csv",               "Nur psychiatrische Einrichtungen (QS.Einrichtung.ID statt SO.QBID). Zu spezifisch."),
        ("Sonstige Strukturmerkmale", "AMTS.* · GIQI · KISS · Lenkungsgremium · VAVU · ZV\nNotfallversorgung · DMP · EF · IF · QS.Behandlungsumfang · QS.Landesrecht\nSchutzkonzept · Praevention · BewertungStrukDialog · Personen.csv", "Teils zu spezifisch, teils noch nicht gesichtet. Können bei vertiefter Analyse ergänzt werden."),
    ],
    cw=[3.5, 5.3, 9.2], hdr_color="BF5A00", row_colors=[F_GELB]
)
doc.add_paragraph().paragraph_format.space_after = Pt(3)

# ══════════════════════════════════════════════════════════════════
# BLOCK 3: Nicht relevant (NEIN) + KEY — gruppiert
# ══════════════════════════════════════════════════════════════════
p3 = doc.add_paragraph()
r3 = p3.add_run("❌  Nicht relevant — bewusst ausgeschlossen")
r3.bold = True; r3.font.size = Pt(10); r3.font.color.rgb = RGBColor(0x7B, 0x0D, 0x0D)

add_tbl(doc,
    ["Gruppe", "Dateien", "Ausschlussgrund"],
    [
        ("Lookup-Tabellen",           "ICD.Code.csv · OPS.csv (106 MB) · OPS.Code.csv\nQS.Einrichtungstypen · QS.Berufsgruppen",             "Reine Code-Dekodierung. Keine Analysedaten."),
        ("Alle *.Key.csv  (16 Stk.)", "AA/AM/AMTS/AQZF/BF/CQ/EF/HM/IF/LK/MP/NM/PQZP/RM/VAVU.Key.csv", "Schlüssel-/Lookup-Tabellen ohne eigene Messwerte."),
        ("Ersatz durch bessere Quelle", "AQ.Pflege.csv",                                                "Enthält nur Pflege-Qualifikationsnachweise, keine Anzahlen. SO.Personalliste.csv liefert pflege_pro_bett direkter und wurde stattdessen verwendet."),
        ("Nicht-medizin. Angebote",   "NM.csv",                                                          "Parkplatz, WLAN, Catering — kein Analysebezug."),
        ("URL-Links",                 "Link.csv · LinkVersorgunggebieteSO.csv · Weiterführender_Link.csv","Nur Weblinks. Keine Zahlenwerte."),
        ("Technische / Admin-Daten",  "QS.Nachweis.csv · Error.csv · Akademische_Lehre.csv\nErfPersVorgaben · Pflegepersonalregelung · Sicherstellungszuschläge\nAbt.Zugang · Abt301 · Neuartige_Therapien", "Verwaltungsdaten, Metadaten, zu spezifisch — kein Bezug zur Projektfragestellung."),
    ],
    cw=[3.5, 5.3, 9.2], hdr_color="7B0D0D", row_colors=[F_ROT]
)
doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ── Fußzeile ──────────────────────────────────────────────────────
foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
foot.add_run(
    "Kriterium: VERW = Strukturmerkmal lt. Aufgabenstellung ODER Qualitätsindikator-Bewertungen  |  "
    "MOEG = pot. relevant, aber unklar/nicht gesichtet  |  NEIN = kein Analysewert"
).font.size = Pt(7.5)

# ══════════════════════════════════════════════════════════════════
# SEITE 2: Begründungen
# ══════════════════════════════════════════════════════════════════
doc.add_page_break()

def add_section(doc, title, color_rgb, text):
    p = doc.add_paragraph()
    r = p.add_run(title)
    r.bold = True; r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(*color_rgb)
    p = doc.add_paragraph()
    rn = p.add_run(text)
    rn.font.size = Pt(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

# Seitentitel
t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
rt2 = t2.add_run("Entscheidungsbegründung: Wie wurde die Dateiklassifikation ermittelt?")
rt2.bold = True; rt2.font.size = Pt(13); rt2.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ── Methodik ─────────────────────────────────────────────────────
add_section(doc, "1  Methodik der Datei-Sichtung", (0x1F, 0x49, 0x7D),
    "Alle 86 CSV-Dateien wurden systematisch nach einem 6-Schritte-Verfahren analysiert "
    "(durchgeführt in 01_Exploration.ipynb, Zellen 2–5):\n\n"
    "  Schritt 1 — Dateiname lesen: Präfix-Logik gibt ersten Hinweis auf Inhalt "
    "(SO.* = Standort, QS.* = Qualitätssicherung, FA.* = Fachabteilung, *.Key.csv = Lookup).\n\n"
    "  Schritt 2 — Dateigröße prüfen: Dateien über 50 MB können in VS Code nicht direkt "
    "geöffnet werden. Nur 2 Dateien überschreiten diese Grenze: QS.Qualitätsindikator.csv "
    "(911 MB) und OPS.csv (106 MB). Für diese wurde Python (pd.read_csv mit nrows=5) verwendet.\n\n"
    "  Schritt 3 — Header lesen: Spaltennamen (Zeile 1) zeigen Struktur und Inhalt der Datei. "
    "CSV ist Klartext — kein Python nötig.\n\n"
    "  Schritt 4 — Beispielzeilen lesen (Zeilen 2–5): Konkrete Datenwerte klären die Bedeutung. "
    "Erst der Wert 'Aerzte' in FA.Personal.Bereich zeigte, wie nach Ärzte-Zeilen gefiltert wird.\n\n"
    "  Schritt 5 — Key-Check: Enthält die Datei SO.QBID? "
    "SO.QBID erscheint in ~60 von 86 Dateien. Dateien ohne SO.QBID sind entweder "
    "Lookup-Tabellen oder psychiatrie-spezifische Dateien mit eigenem Schlüssel (QS.Einrichtung.ID).\n\n"
    "  Schritt 6 — Relevanz-Entscheidung: Hilft die Datei, den Zusammenhang zwischen "
    "Struktur (A-Teil) und Qualitätsauffälligkeiten (C-Teil) zu untersuchen?\n\n"
    "Grundsatz: Im Zweifel wurde eine Datei als MOEG (möglicherweise relevant) eingestuft "
    "statt sofort als NEIN ausgeschlossen."
)

# ── Begründung VERW ───────────────────────────────────────────────
add_section(doc, "2  Warum wurden diese 7 Dateien verwendet? (VERW)", (0x37, 0x5E, 0x23),
    "SO.csv: Einzige Datei mit allen Strukturmerkmalen der Aufgabenstellung in einer Tabelle — "
    "Betten, Träger, Bundesland, Uni-Status, Geo-Koordinaten, Standortnummer. Enthält zudem SO.QBID als "
    "universellen Join-Schlüssel. Ohne SO.csv gibt es weder Merkmale noch Verknüpfungen.\n\n"
    "QS.Qualitätsindikator.csv: Einzige Datei im gesamten Datensatz mit standardisierter, "
    "für alle 1.824 bewerteten Häuser einheitlich erhobener Bewertung (R* = auffällig, N* = unauffällig). "
    "Grundlage der Ziel-Variable. Exploration (Schritt 2 in 01_Exploration.ipynb) identifizierte "
    "QSErgBewStrukDialog als Schlüsselspalte. Wichtig: N99-Einträge wurden ausgeschlossen, da "
    "'nicht bewertet' ≠ 'nicht auffällig'. Median der auffaellig_quote = 76,92 %. Trägt SO.QBID "
    "bereits selbst — ein Join über QS.csv ist entgegen ursprünglicher Annahme nicht nötig.\n\n"
    "QS.Fortbildung.csv: Laut Aufgabenstellung (Fragestellung.docx) ist Fortbildungsquote "
    "explizit als zu untersuchendes Merkmal genannt. Diese Datei ist die einzige im Datensatz "
    "mit den Zählern Fortbildungspflichtige / Erbrachte — Formel: Erbrachte / Pflichtige.\n\n"
    "FA.csv: Brückentabelle zwischen FA.Personalliste.csv (hat nur ABTID) und SO.csv (braucht SO.QBID). "
    "Ohne FA.csv können Personaldaten nicht dem richtigen Krankenhaus zugeordnet werden.\n\n"
    "FA.Personalliste.csv: Einzige Quelle für Vollzeit-Ärzteanzahl auf Abteilungsebene. "
    "Ergebnis: aerzte_pro_bett ist der stärkste Prädiktor im Decision Tree (Feature Importance 53,6 %). "
    "Fallstrick: FA.Personal.Anzahl ist Komma-Dezimal ('13,47') — muss vor Aggregation zu float konvertiert werden. "
    "Tageskliniken (SO.Betten = 0) erhalten NaN — korrekt, da kein stationäres Bettenprofil.\n\n"
    "SO.Personalliste.csv (ergänzt 2026-07-29): Liefert direkt SO.QBID + SO.Personal.Anzahl mit "
    "Bereich 'Pflege' — kein Umweg über FA.csv nötig. Quelle für pflege_pro_bett, das 2. wichtigste "
    "Merkmal im Decision Tree (Feature Importance 23,8 %). Bewusst statt AQ.Pflege.csv gewählt, "
    "da diese nur Qualifikationsnachweise ohne Anzahlen enthält.\n\n"
    "Konzern.csv (ergänzt 2026-07-29): Liefert ist_konzern über den Schlüssel SO.Standortnummer "
    "(nicht SO.QBID!). Beim ersten Join-Versuch wurde versehentlich gegen SO.QBID verglichen (0 Treffer) "
    "— nach Korrektur: 358 von 1.824 Häusern (19,6 %) sind Konzernhäuser. Chi²-Test zeigt aber keinen "
    "signifikanten Zusammenhang mit Qualitätsproblemen (p=0,90); der Decision Tree bestätigt das mit "
    "0 % Feature Importance. Bewusst trotzdem im Modell belassen — kein Zusammenhang ist ein valider Befund."
)

# ── Begründung MOEG ───────────────────────────────────────────────
add_section(doc, "3  Warum sind diese Dateien nur möglicherweise relevant? (MOEG)", (0xBF, 0x5A, 0x00),
    "Gemeinsames Kriterium für MOEG: Die Datei enthält potenziell wertvolle Informationen, "
    "wurde aber aus mindestens einem der folgenden Gründe nicht eingebunden:\n\n"
    "a) Noch nicht gesichtet (Zeitkapazität): Viele Strukturdateien wie HD.csv (40 MB Hygiene-Detail), "
    "VAVU.csv (14 MB Versorgungsauftrag), RM.csv (Risikomanagement), Notfallversorgung.csv oder "
    "AM.csv (Ausstattungsmerkmale) wurden zwar identifiziert, aber noch nicht im Detail analysiert. "
    "Sie könnten bei einer Vertiefung der Analyse relevante Strukturmerkmale liefern.\n\n"
    "b) Redundanz mit besserer Quelle: AQ.Ärzte.csv enthält Qualifikationsdaten, aber keine "
    "Anzahlen — FA.Personalliste.csv ist detaillierter und wurde bevorzugt. AQ.Pflege.csv wurde "
    "aus demselben Grund durch SO.Personalliste.csv ersetzt und ist daher unter ❌ NEIN geführt, "
    "nicht mehr unter ⚠️ MOEG.\n\n"
    "c) Zu spezifisch für Teilgruppen: QS.Pso.csv, QS.Psy.csv und QS.Struktur.Station.csv "
    "nutzen QS.Einrichtung.ID statt SO.QBID und erfassen nur psychiatrische Einrichtungen. "
    "Eine Einbindung würde die allgemeine Analyse auf Psychiatrie-Häuser einschränken.\n\n"
    "d) Bedeutung unklar: BF.csv (Behandlungsfelder), CQ.csv (Zertifizierungen) "
    "und MM.csv (Mindestmengen) wurden zwar als interessante Strukturmerkmale identifiziert, "
    "aber der Zusammenhang zur Ziel-Variable wurde noch nicht untersucht. Sie sind Kandidaten "
    "für Folgeanalysen — ebenso wie QS.Leistungsbereich.csv (Dokumentationsrate als potenzielle "
    "Qualitätskennzahl) und das neu entdeckte Mitbewerber_Betten.csv (Marktdichte-Merkmal).\n\n"
    "e) Ursprünglich falsch eingeschätzt: QS.csv wurde zunächst als notwendige Brückentabelle "
    "geführt (siehe ältere Versionen dieses Dokuments), ist aber tatsächlich unnötig — "
    "QS.Qualitätsindikator.csv trägt SO.QBID bereits selbst."
)

# ── Begründung NEIN ───────────────────────────────────────────────
add_section(doc, "4  Warum wurden diese Dateien ausgeschlossen? (NEIN)", (0x7B, 0x0D, 0x0D),
    "Lookup-Tabellen (ICD.Code.csv, OPS.csv, OPS.Code.csv, QS.Einrichtungstypen.csv, "
    "QS.Berufsgruppen.csv, alle 16 *.Key.csv-Dateien): Diese Dateien übersetzen nur "
    "Codes in lesbare Bezeichnungen — sie enthalten keine eigenen Messwerte, Merkmale "
    "oder Qualitätsaussagen. Kein direkter Analysewert.\n\n"
    "AQ.Pflege.csv (ersetzt 2026-07-29): Enthält nur Pflege-Qualifikationsnachweise, keine "
    "Personal-Anzahlen. SO.Personalliste.csv liefert SO.QBID + Anzahl direkt und wurde "
    "stattdessen für pflege_pro_bett verwendet — kein offener Punkt, sondern bewusster Ersatz.\n\n"
    "NM.csv (Nicht-medizinische Angebote): Enthält Informationen zu Parkplatz, WLAN, "
    "Cafeteria und Telefon. Keinerlei Bezug zu Qualität oder Krankenhausstruktur "
    "im medizinischen Sinne.\n\n"
    "URL-Dateien (Link.csv, LinkVersorgunggebieteSO.csv, Weiterführender_Link.csv): "
    "Ausschließlich Weblinks zu externen Krankenhausseiten. Keine Zahlenwerte.\n\n"
    "Verwaltungsdaten (ErfPersVorgaben.csv, Pflegepersonalregelung.csv, "
    "Sicherstellungszuschlaege.csv, Abt.Zugang.csv, Abt301.csv): "
    "Administrative Daten zur gesetzlichen Personalplanung und Verwaltungsstruktur. "
    "Kein inhaltlicher Bezug zur medizinischen Qualitäts- oder Strukturanalyse.\n\n"
    "Zu spezifische Einzelthemen (Neuartige_Therapien.csv, "
    "Praevention_Missbrauch_und_Gewalt.csv, Schutzkonzept.csv): Sehr enge Themenbereiche "
    "ohne Relevanz für die Projektfragestellung (Zusammenhang Struktur ↔ Qualitätsauffälligkeit).\n\n"
    "Technische Metadaten (QS.Nachweis.csv, Error.csv): "
    "QS.Nachweis.csv enthält nur Zeiträume für QS-Erhebungen — reine technische Information. "
    "Error.csv enthält Fehlerprotokolle des Datensatzes.\n\n"
    "OPS.csv (106 MB): Größte nicht-relevante Datei. Enthält den Operationen- und "
    "Prozedurenschlüssel als Lookup-Tabelle. Keine Analysedaten — und durch die Größe von "
    "106 MB nur per Python ladbar, was den Ausschluss noch klarer macht."
)

# ── Hinweis offen ────────────────────────────────────────────────
p_off = doc.add_paragraph()
roff = p_off.add_run(
    "📌  Offener Punkt (Stand 2026-07-30): QS.Leistungsbereich.csv (QSLB.Dokumentationsrate) "
    "ist weiterhin identifiziert, aber nicht in die Analysetabelle eingebunden — aussichtsreichster "
    "Kandidat für eine künftige Ergänzung. Außerdem ungeklärt: ob die in QS.csv enthaltene "
    "Unterscheidung QS.Typ (bund/land) die Ziel-Variable verzerrt, da sie bisher nicht zur "
    "Filterung genutzt wird."
)
roff.font.size = Pt(10); roff.italic = True

# ── Speichern ─────────────────────────────────────────────────────
out = "Datei_Uebersicht.docx"
doc.save(out)
print(f"✅ Gespeichert: {out}")
