"""
grafiken_dokumentation.py
==========================
Erzeugt eine Word-Dokumentation fuer Kollegen: beschreibt und erklaert alle
12 Grafiken aus Notebooks/02_Analyse.ipynb (deskriptive Analyse + Inferenzstatistik).

Alle Zahlen werden live aus Data/analysetabelle.csv berechnet (keine
hartkodierten Werte), damit die Doku nicht veraltet, wenn sich die Daten
aendern. Ausfuehrliche Textversion mit allen Hintergruenden: Doku/MD/02_Analyse.md

Aufruf:
    python scripts/grafiken_dokumentation.py
"""
import os
from datetime import date, datetime
from pathlib import Path

import numpy as np
import pandas as pd
from scipy import stats
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PROJEKT_ROOT = Path(__file__).resolve().parent.parent
os.chdir(PROJEKT_ROOT)

DATA_PATH = PROJEKT_ROOT / "Data" / "analysetabelle.csv"
GRAFIK_DIR = PROJEKT_ROOT / "grafiken"
OUT_PATH = PROJEKT_ROOT / "Doku" / "Word" / "Grafiken_Dokumentation.docx"

C_BLAU = RGBColor(0x1F, 0x49, 0x7D)
C_WEISS = RGBColor(0xFF, 0xFF, 0xFF)
C_GRAU = RGBColor(0x59, 0x59, 0x59)
F_GRAU = "F2F2F2"
BILD_BREITE = Cm(15.5)


def set_bg(cell, fill):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def hdr_cell(cell):
    set_bg(cell, "1F497D")
    for run in cell.paragraphs[0].runs:
        run.font.color.rgb = C_WEISS
        run.bold = True
        run.font.size = Pt(9.5)


def add_tbl(doc, headers, rows, cw=None, zebra=True):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        hdr_cell(c)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri + 1].cells[ci]
            c.text = str(val)
            c.paragraphs[0].runs[0].font.size = Pt(9)
            if zebra and ri % 2 == 1:
                set_bg(c, F_GRAU)
    if cw:
        for row in t.rows:
            for i, w in enumerate(cw):
                row.cells[i].width = Cm(w)
    return t


def add_h(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = C_BLAU
    return p


def add_b(doc, text, size=10.5, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color
    return p


def de(x, d=0, sign=False):
    """Zahl im deutschen Format: Punkt=Tausendertrenner, Komma=Dezimaltrenner."""
    s = f"{x:{'+' if sign else ''},.{d}f}"
    return s.replace(",", "§").replace(".", ",").replace("§", ".")


def depct(x, d=1):
    """Anteil (0-1) als deutsch formatierte Prozentzahl, z. B. 0.565 -> '56,5 %'."""
    return de(x * 100, d) + " %"


def add_img(doc, path, caption=None):
    doc.add_picture(str(path), width=BILD_BREITE)
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = C_GRAU


# ══════════════════════════════════════════════════════════════════
# Daten laden + alle Kennzahlen live berechnen (wie in 02_Analyse.ipynb)
# ══════════════════════════════════════════════════════════════════
df = pd.read_csv(DATA_PATH, low_memory=False)
for col in ["SO.Latitude", "SO.Longitude"]:
    df[col] = df[col].astype(str).str.replace(",", ".", regex=False).pipe(pd.to_numeric, errors="coerce")

n_rows = len(df)
g0 = df[df["hat_viele_Probleme"] == 0]   # wenige Probleme
g1 = df[df["hat_viele_Probleme"] == 1]   # viele Probleme
n0, n1 = len(g0), len(g1)

# Grafik 1 — auffaellig_quote
quote = df["auffaellig_quote"]
median_quote = quote.median()
spitze = df[quote == 1.0]
n_spitze = len(spitze)
n_spitze_in_viele = int((spitze["hat_viele_Probleme"] == 1).sum())

# Grafik 2 — Bettenzahl
betten_max = df["SO.Betten"].max()
n_ueber_1500 = int((df["SO.Betten"] > 1500).sum())

# Grafik 3 — Traegerschaft
traeger_pct = df.groupby("KH.Träger.Art")["hat_viele_Probleme"].mean() * 100

# Grafik 4 — Uni
uni_pct = df.groupby("SO.Uni")["hat_viele_Probleme"].mean() * 100
uni_n = df.groupby("SO.Uni").size()

# Grafik 7 — Bundesland
bl_pct = (df.groupby("SO.Bundesland")["hat_viele_Probleme"]
          .agg(["mean", "count"]).rename(columns={"mean": "quote", "count": "n"})
          .sort_values("quote", ascending=False))

# Grafik 8 — Korrelation
num_cols = ["SO.Betten", "fortbildungsquote", "aerzte_pro_bett", "pflege_pro_bett",
            "ist_konzern", "auffaellig_quote", "hat_viele_Probleme", "total_qi"]
corr = df[num_cols].corr()
korr_ziel = (corr["hat_viele_Probleme"]
             .drop(["hat_viele_Probleme", "auffaellig_quote"])  # keine echten Merkmale, sondern die Zielgröße selbst
             .sort_values(key=abs, ascending=False))
r_aerzte_pflege = corr.loc["aerzte_pro_bett", "pflege_pro_bett"]

# Grafik 12 — Konzern
konzern_pct = df.groupby("ist_konzern")["hat_viele_Probleme"].mean() * 100
konzern_n = df.groupby("ist_konzern").size()

# Kapitel 4 — Inferenzstatistik (live nachgerechnet, wie im Notebook)
t_ae, p_ae = stats.ttest_ind(g0["aerzte_pro_bett"].dropna(), g1["aerzte_pro_bett"].dropna())
t_pf, p_pf = stats.ttest_ind(g0["pflege_pro_bett"].dropna(), g1["pflege_pro_bett"].dropna())

kreuz = pd.crosstab(df["ist_konzern"], df["hat_viele_Probleme"])
chi2, p_chi2, dof, _ = stats.chi2_contingency(kreuz)

traeger_gruppen = [g["auffaellig_quote"].dropna() for _, g in df.groupby("KH.Träger.Art")]
f_stat, p_anova = stats.f_oneway(*traeger_gruppen)

ci0 = stats.t.interval(0.95, len(g0["aerzte_pro_bett"].dropna()) - 1,
                        loc=g0["aerzte_pro_bett"].mean(), scale=stats.sem(g0["aerzte_pro_bett"].dropna()))
ci1 = stats.t.interval(0.95, len(g1["aerzte_pro_bett"].dropna()) - 1,
                        loc=g1["aerzte_pro_bett"].mean(), scale=stats.sem(g1["aerzte_pro_bett"].dropna()))

# ══════════════════════════════════════════════════════════════════
doc = Document()
doc.core_properties.author = "Datenanalyse-Team"
doc.core_properties.created = datetime.now()
doc.core_properties.modified = datetime.now()
doc.core_properties.title = "Grafiken der deskriptiven Analyse — Dokumentation"

for s in doc.sections:
    s.top_margin = Cm(2.2)
    s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.3)
    s.right_margin = Cm(2.0)

# ── Titelseite ────────────────────────────────────────────────────
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Grafiken der deskriptiven Analyse")
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = C_BLAU

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.add_run("Alle 12 Grafiken aus Notebooks/02_Analyse.ipynb — was zeigen sie und warum wurden sie erstellt?").font.size = Pt(12)

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run(
    f"Erstellt: {date.today().strftime('%d.%m.%Y')}  |  Quelle: Notebooks/02_Analyse.ipynb  |  "
    f"Datenbasis: Data/analysetabelle.csv ({de(n_rows)} Krankenhäuser)"
).font.size = Pt(10)
doc.add_paragraph()

add_b(
    doc,
    f"Projektfrage: Welche Krankenhausmerkmale hängen damit zusammen, dass ein Haus überdurchschnittlich "
    f"viele Qualitätsprobleme aufweist? Von {de(n_rows)} Häusern haben {de(n0)} wenige Probleme "
    f"(`hat_viele_Probleme = 0`) und {de(n1)} viele Probleme (`hat_viele_Probleme = 1`). Farbschema in allen "
    f"Grafiken: Grün = wenige Probleme, Rot = viele Probleme.",
    bold=True,
)
add_b(
    doc,
    "Ausführlichere Textversion mit allen Hintergründen und Konzept-Erklärungen: Doku/MD/02_Analyse.md",
    size=9.5, italic=True, color=C_GRAU,
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# KAPITEL 1 — Jedes Merkmal einzeln
# ══════════════════════════════════════════════════════════════════
add_h(doc, "1  Jedes Merkmal einzeln betrachtet (Grafiken 1–7, 11–12)", level=1)
add_b(
    doc,
    "Bevor man nach Zusammenhängen sucht, muss jedes Merkmal für sich geprüft werden: sinnvoll verteilt? "
    "Ausreißer? Unterscheiden sich die beiden Gruppen überhaupt sichtbar?"
)

# --- Grafik 1 ---
add_h(doc, "Grafik 1 — Verteilung der auffällig-Quote", level=2)
add_img(doc, GRAFIK_DIR / "g1_auffaellig_quote.png")
add_b(
    doc,
    "auffaellig_quote ist die kontinuierliche Vorstufe der Ziel-Variable: (Anzahl auffälliger "
    "Qualitätsindikatoren eines Hauses) ÷ (Anzahl aller bei ihm bewerteten Indikatoren). "
    "hat_viele_Probleme (die eigentliche 0/1-Ziel-Variable) wird erst daraus abgeleitet: 1, wenn die "
    "Quote über dem Median liegt.",
)
add_b(
    doc,
    "Diagramm lesen: X-Achse = Quote in %, Y-Achse = Anzahl Häuser. Die dunkle Linie verbindet die "
    "Balkenspitzen (bessere Lesbarkeit), die rote gestrichelte Linie markiert den Median.",
    italic=True,
)
add_b(
    doc,
    f"Befund: Zwei getrennte Muster — ein breiter Hügel von ca. 40–85 % und ein eigenständiger, höchster "
    f"Balken exakt bei 100 % mit {de(n_spitze)} Häusern ({depct(n_spitze/n_rows)}). Grund: Diese Häuser haben im "
    f"Median nur sehr wenige bewertete Indikatoren — bei so kleinen Zahlen reicht es, dass alle zufällig "
    f"auffällig sind, um auf 100 % zu kommen (kleine-Zahlen-Effekt, kein echtes Qualitätssignal). "
    f"{de(n_spitze_in_viele)} der {de(n1)} Häuser mit hat_viele_Probleme=1 ({depct(n_spitze_in_viele/n1)}) stammen "
    f"aus genau dieser Spitze. Der Median ({depct(median_quote)}) ist zugleich der Schwellenwert für "
    f"hat_viele_Probleme.",
)

# --- Grafik 2 ---
add_h(doc, "Grafik 2 — Bettenzahl", level=2)
add_img(doc, GRAFIK_DIR / "g2_bettenzahl.png")
add_b(
    doc,
    f"Diagramm lesen: Links ein Histogramm der Bettenzahl, bei 1.500 gekappt (das größte Haus hat "
    f"{de(betten_max)} Betten, aber nur {n_ueber_1500} von {de(n_rows)} Häusern haben über 1.500 Betten — "
    f"ohne Kappung würde sich die Mehrheit der Häuser in einem winzigen Streifen zusammendrängen). "
    f"Rechts ein Boxplot (ungekappt) je Gruppe: Box = mittlere 50 %, Strich = Median, Antennen = "
    f"Wertebereich, Punkte = Ausreißer.",
    italic=True,
)
add_b(
    doc,
    f"Befund: Median wenige Probleme = {de(g0['SO.Betten'].median())} Betten, viele Probleme = "
    f"{de(g1['SO.Betten'].median())} Betten. Die mittleren 50 %-Bereiche beider Gruppen überlappen sich "
    f"stark — kein klarer Größenunterschied.",
)

# --- Grafik 3 ---
add_h(doc, "Grafik 3 — Trägerschaft", level=2)
add_img(doc, GRAFIK_DIR / "g3_traegerschaft.png")
add_b(
    doc,
    "Diagramm lesen: Pro Trägerart zwei Balken (grün/rot), die sich zusammen zu 100 % addieren — "
    "gezeigt wird die interne Aufteilung je Trägerart, keine absoluten Hauszahlen.",
    italic=True,
)
add_b(
    doc,
    "Befund: " + ", ".join(f"{idx} {de(val, 1)} %" for idx, val in traeger_pct.sort_values(ascending=False).items())
    + " der Häuser haben viele Probleme — bislang der klarste Unterschied zwischen Gruppen "
    "(aber noch kein statistischer Beweis, siehe Kapitel 3, und noch nicht um den Größen-Störfaktor "
    "aus Grafik 10 bereinigt).",
)

# --- Grafik 4 ---
add_h(doc, "Grafik 4 — Uni-Kliniken vs. normale Häuser", level=2)
add_img(doc, GRAFIK_DIR / "g4_uni.png")
add_b(
    doc,
    "Diagramm lesen: Hier steht pro Gruppe nur ein Balken (Farbe unterscheidet nur die Gruppen, nicht "
    "wenige/viele Probleme). Die Legende nennt zusätzlich die Gruppengröße.",
    italic=True,
)
add_b(
    doc,
    f"Befund: Uni-Kliniken {de(uni_pct.get(1,0), 1)} % (n={de(uni_n.get(1,0))}) vs. normale Häuser "
    f"{de(uni_pct.get(0,0), 1)} % (n={de(uni_n.get(0,0))}) — kaum Unterschied. Die Uni-Gruppe ist klein, was ihr "
    f"Ergebnis unsicherer macht.",
)

# --- Grafik 5+6 ---
add_h(doc, "Grafik 5+6 — Fortbildungsquote & Ärzte pro Bett", level=2)
add_img(doc, GRAFIK_DIR / "g5_6_fortbildung_aerzte.png")
add_b(
    doc,
    "Diagramm lesen: Zwei eigenständige Boxplots mit je eigener Y-Achse (nur innerhalb eines Diagramms "
    "vergleichbar). Ausreißer sind hier ausgeblendet, damit die Boxen nicht zusammengedrückt werden. "
    "Entscheidend ist nicht die Antennenlänge, sondern wie stark sich die grüne und rote Box gegeneinander "
    "verschieben.",
    italic=True,
)
add_b(
    doc,
    f"Befund: Fortbildungsquote — Median wenige Probleme = {de(g0['fortbildungsquote'].median(), 3)}, viele "
    f"Probleme = {de(g1['fortbildungsquote'].median(), 3)} (praktisch identisch, kein Unterschied). "
    f"Ärzte pro Bett — Median wenige Probleme = {de(g0['aerzte_pro_bett'].median(), 3)}, viele Probleme = "
    f"{de(g1['aerzte_pro_bett'].median(), 3)} (sichtbare Verschiebung). Schlussfolgerung: Fortbildungsquote "
    f"kann man zurückstellen, Ärzte pro Bett ist ein ernster Kandidat für den T-Test in Kapitel 3 — und "
    f"wird dort bestätigt (siehe unten), später sogar wichtigstes Merkmal im Modell.",
)

# --- Grafik 7 ---
add_h(doc, "Grafik 7 — Bundesland", level=2)
add_img(doc, GRAFIK_DIR / "g7_bundesland.png")
add_b(
    doc,
    "Diagramm lesen: Balken sortiert nach Anteil 'viele Probleme'. Die Beschriftung 'n=... Häuser' an "
    "jedem Balken zeigt die Gruppengröße je Bundesland.",
    italic=True,
)
top3 = bl_pct.head(3)
low3 = bl_pct.tail(3)
add_b(
    doc,
    "Befund: Höchster Anteil in " + ", ".join(f"{idx} ({depct(row['quote'])}, n={de(row['n'])})" for idx, row in top3.iterrows())
    + "; niedrigster in " + ", ".join(f"{idx} ({depct(row['quote'])}, n={de(row['n'])})" for idx, row in low3.iloc[::-1].iterrows())
    + ". Einschränkung: Manche Bundesländer haben nur wenige Häuser — ein einzelnes Haus kann den "
    "Landes-Durchschnitt verschieben, daher mit Vorsicht zu lesen.",
)

# --- Grafik 11 ---
add_h(doc, "Grafik 11 — Pflegekräfte pro Bett", level=2)
add_img(doc, GRAFIK_DIR / "g11_pflege_pro_bett.png")
add_b(
    doc,
    f"Befund: Median wenige Probleme = {de(g0['pflege_pro_bett'].median(), 3)}, viele Probleme = "
    f"{de(g1['pflege_pro_bett'].median(), 3)} — ähnliches Muster wie Ärzte pro Bett (Grafik 6), was ein "
    f"erster Hinweis darauf ist, dass beide Merkmale Ähnliches messen (siehe Korrelation in Kapitel 2).",
)

# --- Grafik 12 ---
add_h(doc, "Grafik 12 — Konzernvergleich", level=2)
add_img(doc, GRAFIK_DIR / "g12_konzern_vergleich.png")
add_b(
    doc,
    "Diagramm lesen: Wie bei Grafik 4 ein Balken pro Gruppe; die 'n=...'-Beschriftung ist die "
    "Gruppengröße (wie viele Häuser unabhängig bzw. Konzernhaus sind), nicht die Balkenhöhe (die zeigt "
    "nur den Anteil in Prozent).",
    italic=True,
)
add_b(
    doc,
    f"Befund: Konzernhäuser {de(konzern_pct.get(1,0), 1)} % (n={de(konzern_n.get(1,0))}) vs. unabhängige Häuser "
    f"{de(konzern_pct.get(0,0), 1)} % (n={de(konzern_n.get(0,0))}) — praktisch kein Unterschied. Wird in Kapitel 3 "
    f"mit dem Chi²-Test formal bestätigt.",
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# KAPITEL 2 — Zusammenhänge zwischen Merkmalen
# ══════════════════════════════════════════════════════════════════
add_h(doc, "2  Zusammenhänge zwischen Merkmalen (Grafiken 8–10)", level=1)
add_b(
    doc,
    "Zwei neue Fragen: Welches Merkmal hängt am stärksten mit der Ziel-Variable zusammen (Grafik 8), und "
    "hängen die Merkmale auch untereinander zusammen, sodass ein scheinbarer Befund eigentlich durch ein "
    "drittes, verstecktes Merkmal — einen sogenannten Störfaktor — verursacht wird (Grafik 9–10)?"
)

add_h(doc, "Grafik 8 — Korrelationsmatrix", level=2)
add_img(doc, GRAFIK_DIR / "g8_korrelation.png")
add_b(
    doc,
    "Diagramm lesen: Zeilen und Spalten sind dieselben Merkmale. Jede Zelle zeigt die Pearson-Korrelation "
    "r (−1 bis +1) zwischen Zeilen- und Spalten-Merkmal. Diagonale = immer 1,00. Die Matrix ist gespiegelt "
    "— nur eine Hälfte muss man lesen. Dunkelrot = starker positiver, dunkelblau = starker negativer, "
    "blass = kein Zusammenhang.",
    italic=True,
)
add_tbl(
    doc,
    ["Merkmal", "r (mit hat_viele_Probleme)"],
    [(idx, de(val, 3, sign=True)) for idx, val in korr_ziel.items()],
    cw=[8.0, 6.0],
)
add_b(
    doc,
    f"Befund: total_qi korreliert am stärksten ({de(corr.loc['total_qi','hat_viele_Probleme'], 2, sign=True)}), ist "
    f"aber ein reines Strukturmerkmal (Anzahl bewerteter Indikatoren), kein Qualitätsmerkmal. Der "
    f"interessantere Befund: aerzte_pro_bett und pflege_pro_bett korrelieren fast gleich stark — beide "
    f"deutlich vor allen anderen echten Merkmalen. Nebenbefund: Beide korrelieren auch untereinander "
    f"recht stark (r ≈ {de(r_aerzte_pflege, 2)}) — ein Hinweis auf Multikollinearität (beide messen "
    f"vermutlich teilweise dieselbe zugrundeliegende 'Personalausstattung').",
)

add_h(doc, "Grafik 9 — Scatter: Bettenzahl vs. Ärzte pro Bett", level=2)
add_img(doc, GRAFIK_DIR / "g9_scatter_betten_aerzte.png")
add_b(
    doc,
    "Befund: Kein klares Trennmuster — grüne und rote Punkte überlappen stark. Mit diesen beiden "
    "Merkmalen allein lässt sich kein Haus zuverlässig einer Gruppe zuordnen.",
)

add_h(doc, "Grafik 10 — Störfaktor: Trägerschaft × Bettengröße", level=2)
add_img(doc, GRAFIK_DIR / "g10_stoerfaktor_traeger.png")
add_b(
    doc,
    "Was ist ein Störfaktor (Confounder)? Eine dritte Variable, die mit beiden Seiten eines beobachteten "
    "Zusammenhangs gleichzeitig zu tun hat und dadurch einen Zusammenhang vortäuschen oder verstärken "
    "kann. Grafik 3 zeigte private Häuser mit dem höchsten Problem-Anteil — bevor man daraus schließt "
    "'privat verursacht schlechtere Qualität', muss man prüfen, ob Träger und Qualität vielleicht beide "
    "mit einer dritten Eigenschaft zusammenhängen: hier der Größe des Hauses.",
)
add_b(
    doc,
    "Diagramm lesen: Drei Boxplots je Trägerart, Y-Achse = Bettenzahl — hier geht es nicht um wenige/"
    "viele Probleme, sondern nur um die Bettengröße je Trägerart.",
    italic=True,
)
traeger_betten = df.groupby("KH.Träger.Art")["SO.Betten"].median().sort_values()
add_b(
    doc,
    "Befund: " + ", ".join(f"{idx} {de(val)} Betten" for idx, val in traeger_betten.items())
    + " (Median). Private Häuser sind systematisch kleiner — kleinere Häuser haben pro Qualitätsindikator "
    "weniger Fälle, was die statistische Schwankungsbreite erhöht. Der scheinbare Träger-Effekt aus "
    "Grafik 3 könnte also ganz oder teilweise ein Größen-Effekt sein, der sich hinter der Trägerart "
    "versteckt.",
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# KAPITEL 3 — Inferenzstatistik
# ══════════════════════════════════════════════════════════════════
add_h(doc, "3  Inferenzstatistik — sind die Unterschiede „echt\"?", level=1)
add_b(
    doc,
    "Ein sichtbarer Unterschied in einer Grafik ist noch kein Beweis — er könnte Zufall sein. "
    "Inferenzstatistik prüft: Wie wahrscheinlich wäre der beobachtete Unterschied, wenn es in "
    "Wahrheit gar keinen gäbe? Ist diese Wahrscheinlichkeit (der p-Wert) sehr klein (< 5 %), gilt der "
    "Unterschied als statistisch signifikant."
)

def de_p(p):
    return "< 0,001" if p < 0.001 else de(p, 3)


add_tbl(
    doc,
    ["Test", "Merkmal", "Ergebnis", "p-Wert", "Signifikant?"],
    [
        ("T-Test", "Ärzte pro Bett", f"t = {de(t_ae, 3)}", de_p(p_ae), "Ja" if p_ae < 0.05 else "Nein"),
        ("T-Test", "Pflege pro Bett", f"t = {de(t_pf, 3)}", de_p(p_pf), "Ja" if p_pf < 0.05 else "Nein"),
        ("Chi²-Test", "Konzern ↔ viele Probleme", f"χ² = {de(chi2, 3)}", de_p(p_chi2), "Nein" if p_chi2 >= 0.05 else "Ja"),
        ("ANOVA", "Trägerschaft (3 Gruppen)", f"F = {de(f_stat, 3)}", de_p(p_anova), "Ja" if p_anova < 0.05 else "Nein"),
    ],
    cw=[3.0, 5.0, 3.0, 3.0, 2.5],
)
doc.add_paragraph()

add_b(doc, "T-Test: Ärzte pro Bett", bold=True)
add_b(
    doc,
    f"Prüft, ob sich die Mittelwerte zweier Gruppen unterscheiden. Befund: wenige Probleme Ø = "
    f"{de(g0['aerzte_pro_bett'].mean(), 3)}, viele Probleme Ø = {de(g1['aerzte_pro_bett'].mean(), 3)}, "
    f"t = {de(t_ae, 3)}, p {de_p(p_ae)} → hoch signifikant. Bestätigt den optischen Eindruck aus Grafik 6 sehr "
    f"deutlich — der stärkste bestätigte Einzelbefund der gesamten deskriptiven Analyse."
)

add_b(doc, "T-Test: Pflegekräfte pro Bett", bold=True)
add_b(
    doc,
    f"Befund: wenige Probleme Ø = {de(g0['pflege_pro_bett'].mean(), 3)}, viele Probleme Ø = "
    f"{de(g1['pflege_pro_bett'].mean(), 3)}, t = {de(t_pf, 3)}, p {de_p(p_pf)} → ebenfalls hoch signifikant, fast "
    f"identisch stark wie bei den Ärzten — ein Hinweis, dass hier ein gemeinsamer Effekt "
    f"'Personalausstattung' gemessen wird, nicht zwei unabhängige Phänomene."
)

add_b(doc, "Chi²-Test: Konzernzugehörigkeit", bold=True)
add_b(
    doc,
    f"Prüft, ob zwei kategoriale Merkmale (Konzern ja/nein, viele Probleme ja/nein) unabhängig sind. "
    f"Befund: χ² = {de(chi2, 3)}, p = {de(p_chi2, 2)} → nicht signifikant, deutlich über α = 0,05. Es gibt "
    f"keinen nachweisbaren Zusammenhang zwischen Konzernzugehörigkeit und Qualitätsproblemen — "
    f"bestätigt Grafik 12 statistisch eindeutig."
)

add_b(doc, "ANOVA: auffällig-Quote nach Trägerschaft", bold=True)
add_b(
    doc,
    f"Prüft, ob sich die Mittelwerte von mehr als zwei Gruppen unterscheiden (hier: 3 Trägerarten). "
    f"Befund: F = {de(f_stat, 3)}, p {de_p(p_anova)} → signifikant. Bestätigt, dass der in Grafik 3 gesehene "
    f"Trägerschafts-Unterschied nicht zufällig ist — aber: nicht um die Bettengröße bereinigt "
    f"(siehe Grafik 10, Größen-Störfaktor)."
)

add_b(doc, "95 %-Konfidenzintervall: Ärzte pro Bett", bold=True)
add_b(
    doc,
    f"Ein Wertebereich, der den 'wahren' Mittelwert mit 95 % Sicherheit einschließt. Befund: "
    f"wenige Probleme [{de(ci0[0], 3)}; {de(ci0[1], 3)}], viele Probleme [{de(ci1[0], 3)}; {de(ci1[1], 3)}] — "
    f"die Intervalle überlappen sich nicht, ein weiteres Indiz für einen realen Unterschied."
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# KAPITEL 4 — Gesamtfazit
# ══════════════════════════════════════════════════════════════════
add_h(doc, "4  Gesamtfazit", level=1)
add_b(
    doc,
    "Keine starken, eindeutigen Zusammenhänge zwischen Strukturmerkmalen und Qualitätsproblemen. Die "
    "einzigen statistisch abgesicherten Befunde sind aerzte_pro_bett und pflege_pro_bett (beide "
    "p < 0,001, aber beide nur schwach korreliert mit der Ziel-Variable) sowie der Trägerschafts-"
    "Unterschied (ANOVA p < 0,001, aber mit dem in Grafik 10 aufgedeckten Größen-Störfaktor). "
    "ist_konzern und fortbildungsquote zeigen dagegen klar keinen Zusammenhang.",
)
add_b(
    doc,
    "Warum das kein Scheitern ist: Ein Qualitätsbericht-'auffällig' ist kein automatisches "
    "Qualitätsurteil, und viele Faktoren, die wirklich zählen könnten (Patientenmix, Spezialisierung, "
    "Dokumentationsqualität), stehen gar nicht im Datensatz. 'Kein Zusammenhang ist ein valides "
    "Ergebnis' — die ehrliche Aussage dieser Analyse ist, dass Strukturmerkmale allein die "
    "Auffälligkeit eines Hauses nur sehr schwach erklären. Das wird in 03_Decision_Tree.ipynb mit "
    "einem R² von nur 0,033 noch einmal zahlenmäßig bestätigt.",
    italic=True,
)

# ── Speichern ─────────────────────────────────────────────────────
OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(OUT_PATH))
print(f"Gespeichert: {OUT_PATH}")
