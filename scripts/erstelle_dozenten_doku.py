"""
erstelle_dozenten_doku.py
=========================
Erzeugt Doku/Dozent/Fortschrittsbericht_Qualitaets_Muster_Finder.docx
Umfang: max. 10 Seiten · Bausteine 1 & 2 (Wochen 1–2)

Aufruf: python scripts/erstelle_dozenten_doku.py
"""

import datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PROJEKT_ROOT = Path(__file__).resolve().parent.parent
OUT_PATH     = PROJEKT_ROOT / "Doku" / "Dozent" / "Fortschrittsbericht_Qualitaets_Muster_Finder.docx"
GRAFIKEN     = PROJEKT_ROOT / "grafiken"

BLAU   = RGBColor(0x1F, 0x49, 0x7D)
ORANGE = RGBColor(0xBF, 0x5A, 0x00)
GRUEN  = RGBColor(0x37, 0x5E, 0x23)
GRAU   = RGBColor(0x55, 0x55, 0x55)

doc = Document()
doc.core_properties.author   = "Datenanalyse-Projekt"
doc.core_properties.modified = datetime.datetime.now()

for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── Hilfsfunktionen ───────────────────────────────────────────────

def h1(text):
    p = doc.add_heading(level=1)
    p.clear()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.font.size = Pt(14); r.font.bold = True; r.font.color.rgb = BLAU

def h2(text):
    p = doc.add_heading(level=2)
    p.clear()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(text)
    r.font.size = Pt(12); r.font.bold = True; r.font.color.rgb = ORANGE

def body(text, size=11, farbe=None, italic=False, fett=False, abstand_nach=5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(abstand_nach)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.italic = italic
    r.font.bold = fett
    if farbe: r.font.color.rgb = farbe

def code(text):
    # Teilt bei '#' auf: Code dunkelblau, Kommentar grün
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.8)
    if '#' in text:
        idx = text.index('#')
        r_code = p.add_run(text[:idx])
        r_code.font.size = Pt(9.5); r_code.font.name = "Courier New"
        r_code.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
        r_cmt = p.add_run(text[idx:])
        r_cmt.font.size = Pt(9.5); r_cmt.font.name = "Courier New"
        r_cmt.font.color.rgb = RGBColor(0x27, 0x7A, 0x3B)  # grün
    else:
        r = p.add_run(text)
        r.font.size = Pt(9.5); r.font.name = "Courier New"
        r.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

def bullet(text, bold_prefix=""):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    if bold_prefix:
        rb = p.add_run(bold_prefix)
        rb.bold = True; rb.font.size = Pt(11)
    r = p.add_run(text)
    r.font.size = Pt(11)

def trennlinie():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1");   bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom); pPr.append(pBdr)

def body_runs(parts, size=11, abstand_nach=5):
    """parts = list of (text, bold) tuples for mixed bold/normal in one paragraph"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(abstand_nach)
    for text, bold in parts:
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.font.bold = bold

def tabelle(headers, rows, col_widths):
    tbl = doc.add_table(rows=1+len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        r = hdr[i].paragraphs[0].runs[0]
        r.bold = True; r.font.size = Pt(10); r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        tc = hdr[i]._tc; tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "1F497D"); shd.set(qn("w:val"), "clear")
        tcPr.append(shd)
    for ri, row in enumerate(rows):
        cells = tbl.rows[ri+1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            cells[ci].paragraphs[0].runs[0].font.size = Pt(10)
    for row in tbl.rows:
        for i, w in enumerate(col_widths):
            row.cells[i].width = Cm(w)
    return tbl

def bild(pfad, breite_cm=14.5, titel=""):
    if titel:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.keep_with_next = True   # Überschrift bleibt bei Seitenumbruch bei der Grafik
        r = p.add_run(titel)
        r.bold = True; r.font.size = Pt(10); r.font.color.rgb = GRAU
    if Path(pfad).exists():
        doc.add_picture(str(pfad), width=Cm(breite_cm))
        bild_p = doc.paragraphs[-1]
        bild_p.paragraph_format.keep_together = True   # Grafik selbst nicht über Seiten aufteilen
        bild_p.paragraph_format.keep_with_next = True  # Grafik bleibt beim nachfolgenden Erklärtext
        doc.add_paragraph().paragraph_format.space_after = Pt(2)
    else:
        body(f"[Grafik nicht gefunden: {pfad}]", farbe=GRAU, italic=True)

# ═════════════════════════════════════════════════════════════════
# DECKBLATT
# ═════════════════════════════════════════════════════════════════
titel = doc.add_paragraph()
titel.alignment = WD_ALIGN_PARAGRAPH.CENTER
titel.paragraph_format.space_before = Pt(30)
r = titel.add_run("Qualitäts-Muster-Finder")
r.bold = True; r.font.size = Pt(22); r.font.color.rgb = BLAU

untertitel = doc.add_paragraph()
untertitel.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = untertitel.add_run("Fortschrittsbericht — Anfang Woche 3")
r2.font.size = Pt(14); r2.font.color.rgb = GRAU

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = info.add_run(
    f"Datenanalyse-Projekt · {datetime.date.today().strftime('%d.%m.%Y')}\n"
    "Datenbasis: IQTIG Qualitätsberichte 2023 · 1.824 Krankenhäuser"
)
r3.font.size = Pt(11); r3.font.color.rgb = GRAU

doc.add_paragraph()
doc.add_paragraph()
body(
    "Dieser Bericht dokumentiert den Projektfortschritt nach Abschluss der "
    "ersten zwei Arbeitswochen. Gezeigt werden die Aufgaben, die laut Projektplan "
    "bis zu diesem Zeitpunkt erledigt sein sollten: "
    "Datenvorbereitung (Baustein 1) und Deskriptive Analyse (Baustein 2). "
    "Die Umsetzung dieser Schritte erfolgte mit Python und Power BI.",
    italic=True, farbe=GRAU, abstand_nach=0
)
doc.add_page_break()

# ═════════════════════════════════════════════════════════════════
# 1. PROJEKTÜBERSICHT
# ═════════════════════════════════════════════════════════════════
h1("1  Projektübersicht")

h2("1.1  Fragestellung")
body(
    "Welche Krankenhausmerkmale h\u00e4ngen damit zusammen, dass ein Haus "
    "überdurchschnittlich viele Qualit\u00e4tsprobleme aufweist?",
    fett=True
)
body(
    "Kein Zusammenhang ist ebenfalls ein valides Ergebnis — "
    "Ehrlichkeit ist hier wichtiger als positive Befunde.",
    italic=True, farbe=GRAU
)

h2("1.2  Datenbasis")
tabelle(
    ["Kennzahl", "Wert"],
    [
        ("Datenjahr",            "2023 · Quelle: IQTIG / G-BA"),
        ("Krankenhäuser",        "1.824 (mit Qualitätsbewertung)"),
        ("CSV-Dateien gesamt",   "86 · ca. 1,2 GB"),
        ("Ziel-Variable",        "hat_viele_Probleme (0 = wenige, 1 = viele)"),
        ("Analysetabelle",       "analysetabelle.csv · 1.824 Zeilen · 18 Spalten"),
        ("Python-Hauptdatei",    "Notebooks/01_Exploration.ipynb"),
    ],
    col_widths=[5.5, 9.5]
)
body(
    "IQTIG — Institut für Qualitätssicherung und Transparenz im Gesundheitswesen: "
    "Bundesbehörde, die im Auftrag des G-BA die jährlichen Qualitätsberichte aller "
    "deutschen Krankenhäuser erhebt, auswertet und veröffentlicht. "
    "G-BA — Gemeinsamer Bundesausschuss: oberstes Beschlussgremium der gemeinsamen "
    "Selbstverwaltung im deutschen Gesundheitswesen; legt fest, welche Qualitätsindikatoren "
    "erhoben werden müssen.",
    size=10, farbe=GRAU, italic=True
)
doc.add_paragraph()

h2("1.3  Verwendete Python-Bibliotheken")
tabelle(
    ["Bibliothek", "Verwendungszweck"],
    [
        ("pandas",      "Daten laden, filtern, mergen, aggregieren (groupby, merge, pivot_table)"),
        ("numpy",       "Numerische Berechnungen (Histogramm-Bins, Median)"),
        ("matplotlib",  "Alle 12 Grafiken (Histogramm, Boxplot, Scatter, Balken, Heatmap)"),
        ("seaborn",     "Korrelations-Heatmap (Grafik 8), Boxplot Träger×Betten (Grafik 10)"),
        ("scipy.stats", "T-Test, ANOVA, Chi²-Test, Konfidenzintervalle"),
        ("pathlib",     "Plattformunabhängige Dateipfade"),
        ("python-docx", "Word-Dokumentation (dieses Dokument)"),
    ],
    col_widths=[3.5, 11.5]
)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════
# 2. BAUSTEIN 1 — DATENVORBEREITUNG (Woche 1)
# ═════════════════════════════════════════════════════════════════
h1("2  Baustein 1 — Datenvorbereitung  (Woche 1)")
body(
    "Notebook: Notebooks/01_Exploration.ipynb  ·  "
    "Ergebnis: Data/analysetabelle.csv"
)
trennlinie()

h2("2.1  Datensatz erkunden — Join-Schlüssel finden")
body(
    "Von allen 86 CSV-Dateien wurden mit `nrows=0` nur die Spaltennamen geladen "
    "(kein einziger Datensatz), um schnell den universellen Join-Schlüssel "
    "zu finden: SO.QBID erscheint in 60 der 86 Dateien."
)
code("df_header = pd.read_csv(datei, nrows=0)   # nur Spaltenheader, keine Daten")
code("spalten_zaehler = Counter()               # zählt Spalten über alle Dateien")
code("→ SO.QBID: 60 Dateien  |  ABTID: 5 Dateien")
body(
    "SO.QBID — Qualitätsbericht-ID des Standorts: eindeutige Nummer, die jedem "
    "Krankenhaus-Standort zugeteilt ist und in fast allen Tabellen vorkommt. "
    "Sie ist der universelle Verbindungsschlüssel des gesamten Datensatzes. "
    "ABTID — Abteilungs-ID: eindeutige Nummer einer einzelnen Fachabteilung "
    "innerhalb eines Krankenhauses (z. B. Chirurgie oder Innere Medizin). "
    "Sie verbindet FA.csv (Abteilungsliste) mit FA.Personalliste.csv (Personal je Abteilung).",
    size=10, farbe=GRAU, italic=True
)

h2("2.2  Ziel-Variable berechnen  (QS.Qualitätsindikator.csv)")
body_runs([
    ("Warum steckt die Ziel-Variable genau in dieser Datei? ", True),
    ("Beim Durchsuchen der 86 CSV-Header fiel diese Datei durch zwei Eigenschaften auf: "
     "Sie ist mit 911 MB die größte Datei im Datensatz — ein Signal für zentrale Inhaltsdaten. "
     "Außerdem enthält sie die Spalte ", False),
    ("QSErgBewStrukDialog", True),
    (" (= Ergebnis der Bewertung im Strukturierten Dialog), "
     "deren Werte R10 / R20 für auffällig und N01 / N02 für unauffällig stehen. "
     "Damit war klar: Diese Spalte ist die gesuchte Bewertung.", False),
])
body_runs([
    ("Die zweite wichtige Spalte ist ", False),
    ("QSQI.ArtDesWertes", True),
    (". Sie gibt an, ob eine Zeile ein echter Qualitätsindikator (QI) "
     "oder nur eine Fallzählung (EKez, TKez, KKez) ist. "
     "Fallzählungen enthalten keine Bewertung und müssen herausgefiltert werden.", False),
])
body("Ergebnis der Bereinigung — Zeilen nach jedem Schritt:")
tabelle(
    ["Schritt", "Beschreibung", "Entfernt", "Verbleibend"],
    [
        ("Start",     "Ausgangsdatensatz QS.Qualitätsindikator.csv",        "—",        "417.799"),
        ("Schritt 1", "Zählkennzahlen entfernt (EKez 33.557 · TKez 56.977 · KKez 18.539)",
                                                                           "109.073",   "308.726"),
        ("Schritt 2", "N99 entfernt — nicht bewertet ≠ unauffällig",        "36.358",   "272.368"),
        ("Schritt 3", "Duplikate entfernt — je Haus+Indikator eine Zeile (63 % der Zeilen!)",
                                                                           "172.683",    "99.685"),
        ("Ergebnis",  "Einzigartige Häuser: 1.824 · Ø 54,7 Indikatoren/Haus",   "—",    "99.685"),
    ],
    col_widths=[2.5, 7.5, 2.5, 2.5]
)
body("Median auffällig-Quote: 76,92 %  →  Trennwert für hat_viele_Probleme (0/1)",
     italic=True, farbe=GRAU)
body_runs([
    ("Von 99.685 Zeilen zu 1.824 Häusern — wie? ", True),
    ("Jede der 99.685 Zeilen enthält eine SO.QBID — die ID des Krankenhauses, zu dem der Indikator gehört. "
     "Viele Zeilen teilen dieselbe SO.QBID, weil ein Haus viele Indikatoren hat. "
     "Mit nunique() zählen wir, wie viele verschiedene IDs vorkommen: "
     "qi_dedup['SO.QBID'].nunique() → 1.824. "
     "Das bedeutet: In den 99.685 Zeilen gibt es genau 1.824 verschiedene Krankenhaus-IDs.", False),
])
body_runs([
    ("Wichtig — 1.824 ist kein Bereinigungsartefakt: ", True),
    ("Nachgerechnet auf den komplett ungefilterten 417.799 Roh-Zeilen (vor jedem der drei "
     "Bereinigungsschritte) ergibt qi['SO.QBID'].nunique() bereits ebenfalls 1.824. Die Bereinigung "
     "reduziert also nur die Zeilenzahl (417.799 → 99.685), verliert dabei aber kein einziges Haus — "
     "jedes Haus, das überhaupt in QS.Qualitätsindikator.csv vorkommt, hat mindestens einen gültigen, "
     "bewerteten, nicht-doppelten QI-Eintrag. Zum Vergleich: SO.csv (Stammdaten) hat 2.310 eindeutige "
     "SO.QBID — das sind alle Krankenhaus-Standorte, unabhängig davon, ob eine Qualitätsbewertung "
     "vorliegt.", False),
])
body_runs([
    ("Daher der Durchschnitt: ", True),
    ("99.685 Zeilen ÷ 1.824 Häuser = 54,7 Indikatoren pro Haus. "
     "Der groupby()-Schritt fasst dann alle Zeilen desselben Hauses zusammen: "
     "Er zählt, wie viele Indikatoren das Haus hat (total_qi) und wie viele davon auffällig sind (auffaellig_n). "
     "Das Ergebnis: eine Zeile pro Krankenhaus mit der berechneten auffaellig_quote.", False),
])
body_runs([
    ("Wie wurde der Median berechnet? ", True),
    ("Alle 1.824 Häuser haben jetzt eine auffaellig_quote zwischen 0 % und 100 %. "
     "Diese Werte werden der Größe nach sortiert — der Median ist der Wert genau in der Mitte "
     "(Platz 912 von 1.824). Er liegt bei 76,92 %. "
     "Jedes Haus oberhalb dieses Wertes bekommt hat_viele_Probleme = 1, alle anderen = 0. "
     "Der Median wurde als Schwelle gewählt, weil er die Häuser in zwei gleich große Gruppen teilt "
     "und robuster gegen einzelne Extremwerte ist als der Mittelwert.", False),
])
bild(GRAFIKEN / "g1_auffaellig_quote.png", breite_cm=13,
     titel="Ergebnis: Verteilung der auffällig-Quote — graue Linie = Median 76,92 % (Trennwert)")
body("Beispiel — so sieht das Ergebnis für die ersten Häuser aus:",
     fett=True, abstand_nach=3)
tabelle(
    ["SO.QBID", "total_qi", "auffaellig_n", "auffaellig_quote", "hat_viele_Probleme"],
    [
        ("4876", "19",  "18", "94,7 %", "1  (viele Probleme)"),
        ("4878", "2",   "2",  "100,0 %","1  (viele Probleme)"),
        ("4879", "24",  "17", "70,8 %", "0  (wenige Probleme)"),
        ("4880", "16",  "11", "68,8 %", "0  (wenige Probleme)"),
        ("4881", "30",  "20", "66,7 %", "0  (wenige Probleme)"),
        ("4882", "35",  "25", "71,4 %", "0  (wenige Probleme)"),
        ("4886", "101", "79", "78,2 %", "1  (viele Probleme)"),
        ("4887", "99",  "67", "67,7 %", "0  (wenige Probleme)"),
    ],
    col_widths=[2.0, 2.0, 2.5, 3.5, 5.0]
)
body(
    "SO.QBID = Krankenhaus-ID  ·  total_qi = Anzahl bewerteter Indikatoren  ·  "
    "auffaellig_n = davon auffällig  ·  auffaellig_quote = auffaellig_n / total_qi  ·  "
    "hat_viele_Probleme = 1 wenn Quote > Median 76,92 %",
    size=9.5, farbe=GRAU, italic=True
)
doc.add_paragraph()

h2("2.3  Merkmale (Features) aufbereiten")
tabelle(
    ["Merkmal", "Quelle", "Python-Technik"],
    [
        ("SO.Betten",          "SO.csv",              "pd.read_csv → direkt verfügbar"),
        ("KH.Träger.Art",      "SO.csv",              "pd.read_csv → direkt verfügbar"),
        ("SO.Bundesland",      "SO.csv",              "pd.read_csv → direkt verfügbar"),
        ("SO.Uni",             "SO.csv",              "pd.read_csv → direkt verfügbar"),
        ("fortbildungsquote",  "QS.Fortbildung.csv",  "Berechnet: Erbracht / Pflichtige"),
        ("aerzte_pro_bett",    "FA.Personalliste.csv + FA.csv",
         "2× Left Join über ABTID → FA.QBID; str.replace(',','.') → float"),
        ("pflege_pro_bett",    "SO.Personalliste.csv","Filter Bereich=='Pflege', groupby, / SO.Betten"),
        ("ist_konzern",        "Konzern.csv",          "SO.Standortnummer in Konzern? → 0/1"),
    ],
    col_widths=[3.5, 4.5, 7.0]
)
doc.add_paragraph()
body(
    "Technische Besonderheit bei aerzte_pro_bett: "
    "Dezimalzahlen lagen im deutschen Format gespeichert ('13,47' statt '13.47'). "
    "pandas interpretiert das als String — Summe wäre 0. Fix: str.replace + pd.to_numeric.",
    farbe=GRAU, italic=True
)
code("aerzte['anzahl_float'] = aerzte['FA.Personal.Anzahl']")
code("    .astype(str).str.replace(',', '.').pipe(pd.to_numeric, errors='coerce')")

h2("2.4  Analysetabelle zusammenführen")
body("Alle Merkmale + Ziel-Variable werden über SO.QBID per Left Join zusammengeführt:")
code("analyse = auffaellig_quote")
code("    .merge(so_klein,        on='SO.QBID', how='left')")
code("    .merge(fb_quote,        on='SO.QBID', how='left')")
code("    .merge(aerzte_pro_haus, on='SO.QBID', how='left')")
code("    .merge(pflege_haus,     on='SO.QBID', how='left')")
code("    .merge(konzern_flag,    on='SO.QBID', how='left')")
code("→ Ergebnis: Data/analysetabelle.csv  ·  1.824 Zeilen × 18 Spalten")
body("Woher stammt jede Tabelle in dieser Merge-Kette?", fett=True, abstand_nach=3)
tabelle(
    ["Variable", "Quelldatei", "Beigetragene Spalten"],
    [
        ("auffaellig_quote", "QS.Qualitätsindikator.csv (Abschnitt 2.2, Basis der Kette)",
         "total_qi, auffaellig_n, auffaellig_quote, hat_viele_Probleme"),
        ("so_klein",          "SO.csv (Stammdaten)",
         "SO.Name, SO.Betten, SO.Bundesland, SO.Uni, KH.Träger, KH.Träger.Art, "
         "SO.Latitude, SO.Longitude, SO.Standortnummer"),
        ("fb_quote",          "QS.Fortbildung.csv",
         "fortbildungsquote"),
        ("aerzte_pro_haus",   "FA.Personalliste.csv + FA.csv (2 Joins über ABTID)",
         "aerzte_pro_bett"),
        ("pflege_haus",       "SO.Personalliste.csv",
         "pflege_pro_bett"),
        ("konzern_flag",      "Konzern.csv",
         "ist_konzern"),
    ],
    col_widths=[3.0, 6.0, 5.0]
)
body(
    "SO.QBID ist in jeder der sechs Tabellen der gemeinsame Schlüssel — dadurch landet am Ende "
    "jede Spalte in der richtigen Zeile (= richtiges Krankenhaus), obwohl die Spalten aus fünf "
    "völlig unterschiedlichen Rohdateien stammen. how='left' mit auffaellig_quote als Basis "
    "stellt sicher, dass alle 1.824 Häuser mit Ziel-Variable erhalten bleiben, auch wenn ihnen "
    "einzelne Merkmale fehlen (z. B. keine Fortbildungsdaten → NaN bei fortbildungsquote).",
    size=9.5, farbe=GRAU, italic=True
)
body_runs([
    ("Warum 1.824 Zeilen und nicht 2.310? ", True),
    ("Ausgangspunkt des Merges ist auffaellig_quote (die Ziel-Variable, 1.824 Häuser) — so_klein aus "
     "SO.csv (2.310 Häuser) wird per Left Join daran angehängt, nicht umgekehrt. Dadurch bleiben nur "
     "Häuser übrig, die auch in der Ziel-Variable vorkommen. Die übrigen 486 Häuser (2.310 − 1.824) "
     "stehen zwar in SO.csv mit Stammdaten, haben aber keine Zeile in QS.Qualitätsindikator.csv — ohne "
     "Ziel-Variable lässt sich für sie kein hat_viele_Probleme bestimmen, sie fallen beim Merge "
     "automatisch heraus.", False),
])

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════
# 3. BAUSTEIN 2 — DESKRIPTIVE ANALYSE (Woche 2)
# ═════════════════════════════════════════════════════════════════
h1("3  Baustein 2 — Deskriptive Analyse  (Woche 2)")
body(
    "Notebook: Notebooks/02_Analyse.ipynb  ·  "
    "Ergebnis: 12 Grafiken in grafiken/  ·  5 statistische Tests"
)
trennlinie()

h2("3.1  Die 12 Grafiken im Detail")
body(
    "Alle Analysen arbeiten ausschließlich auf analysetabelle.csv. Zu jeder Grafik: was sie zeigt "
    "und warum sie für die Projektfrage wichtig ist. Die Tabelle unten fasst die Kernaussage jeder "
    "Grafik auf einen Blick zusammen — im Anschluss folgt jede Grafik einzeln mit Bild und Erklärung."
)
tabelle(
    ["Grafik", "Kernaussage"],
    [
        ("1 — Verteilung der auffällig-Quote",     "Zielvariable linkssteil verteilt, Median 77 % — Ausgangspunkt aller Vergleiche"),
        ("2 — Bettenzahl",                          "Häuser mit wenigen vs. vielen Qualitätsproblemen unterscheiden sich in der Bettenzahl kaum (Md 214 vs. 170, Bereiche überlappen stark)"),
        ("3 — Trägerschaft",                        "Klarster Unterschied: Häuser mit vielen Qualitätsproblemen — privat 56,5 % vs. öffentlich 46,7 %"),
        ("4 — Uni-Kliniken",                        "Anteil mit vielen Qualitätsproblemen bei Uni-Kliniken (47,3 %) vs. normalen Häusern (49,4 %) fast gleich"),
        ("5+6 — Fortbildung & Ärzte pro Bett",       "Fortbildung ohne Effekt · Ärzte pro Bett bei Häusern mit vielen Problemen niedriger — wichtigstes Merkmal"),
        ("7 — Bundesland",                          "Regionale Unterschiede im Anteil der Häuser mit vielen Qualitätsproblemen sichtbar, aber kleine Stichproben beachten"),
        ("8 — Korrelationsmatrix",                   "Kompakte Übersicht: Ärzte/Bett und Pflege/Bett korrelieren am stärksten mit vielen Qualitätsproblemen"),
        ("9 — Streudiagramm Betten × Ärzte/Bett",    "Häuser mit wenigen vs. vielen Qualitätsproblemen lassen sich anhand dieser zwei Merkmale nicht klar trennen"),
        ("10 — Störfaktor Träger × Bettengröße",     "Private Häuser sind im Schnitt kleiner (Md 125 vs. 233 Betten) — die Bettengröße könnte den Träger-Befund aus Grafik 3 verzerren: kleine Häuser schwanken pro Indikator stärker rein zufällig, was fälschlich wie ein Träger-Effekt aussehen kann"),
        ("11 — Pflegekräfte pro Bett",               "Gleiches Muster wie Ärzte/Bett — zweitwichtigstes Merkmal"),
        ("12 — Konzernvergleich",                    "Kein Unterschied zwischen Konzern- und unabhängigen Häusern"),
    ],
    col_widths=[6.0, 9.0]
)
doc.add_paragraph()

# G1
bild(GRAFIKEN / "g1_auffaellig_quote.png", breite_cm=13,
     titel="Grafik 1 — Verteilung der auffällig-Quote (Median 77 %, linkssteil)")
body(
    "Warum wichtig: Zeigt, wie die Ziel-Variable selbst verteilt ist, bevor überhaupt etwas verglichen "
    "wird — Grundvoraussetzung für alle folgenden Auswertungen. Die auffällige Spitze bei 100 % (21,5 % "
    "der Häuser) kommt von Häusern mit nur sehr wenigen bewerteten Indikatoren und ist bei der "
    "Interpretation aller weiteren Ergebnisse zu berücksichtigen.",
    farbe=GRAU, size=9.5
)

# G2
bild(GRAFIKEN / "g2_bettenzahl.png", breite_cm=13,
     titel="Grafik 2 — Bettenzahl: Wenige Probleme Md=214, Viele Probleme Md=170")
body(
    "Warum wichtig: Prüft die naheliegende Hypothese 'größere Häuser = andere Auffälligkeit'. Die "
    "Wertebereiche beider Gruppen überlappen sich fast vollständig — Bettenzahl allein erklärt kaum "
    "etwas, ein wichtiger Ausschluss-Befund für die weitere Merkmalsauswahl.",
    farbe=GRAU, size=9.5
)

# G3
bild(GRAFIKEN / "g3_traegerschaft.png", breite_cm=13,
     titel="Grafik 3 — Trägerschaft: Private Häuser mit höchstem Auffälligkeitsanteil")
body(
    "Warum wichtig: Der optisch klarste Unterschied der gesamten Analyse (privat 56,5 % vs. öffentlich "
    "46,7 % vs. freigemeinnützig 46,4 %) — später per ANOVA statistisch bestätigt (F=11,3, p<0,001). "
    "Muss aber zusammen mit Grafik 10 gelesen werden, da ein Störfaktor dahinterstecken könnte.",
    farbe=GRAU, size=9.5
)

# G4
bild(GRAFIKEN / "g4_uni.png", breite_cm=13,
     titel="Grafik 4 — Uni-Kliniken (47,3 %) vs. normale Häuser (49,4 %)")
body(
    "Warum wichtig: Testet, ob spezialisierte Häuser mit komplexeren Fällen anders abschneiden. "
    "Praktisch kein Unterschied — ein weiterer Ausschluss-Befund, der zeigt, dass nicht jedes plausible "
    "Merkmal auch tatsächlich einen Zusammenhang zeigt.",
    farbe=GRAU, size=9.5
)

# G5+6
bild(GRAFIKEN / "g5_6_fortbildung_aerzte.png", breite_cm=13,
     titel="Grafik 5+6 — Fortbildungsquote (kein Unterschied) & Ärzte pro Bett (sichtbarer Unterschied)")
body(
    "Warum wichtig: Die entscheidende Weiche der gesamten Analyse. Fortbildungsquote zeigt keinerlei "
    "Unterschied (Md=0,667 in beiden Gruppen) und wird verworfen. Ärzte pro Bett zeigt eine sichtbare "
    "Verschiebung (Md 0,468 vs. 0,390) — bestätigt sich später per T-Test (t=6,002, p<0,001) als "
    "stärkstes Einzelmerkmal und höchste Feature Importance im späteren Decision Tree.",
    farbe=GRAU, size=9.5
)

# G7
bild(GRAFIKEN / "g7_bundesland.png", breite_cm=13,
     titel="Grafik 7 — Bundesland: Saarland 63,2 % (n=19) vs. Berlin 33,3 % (n=54)")
body(
    "Warum wichtig: Prüft, ob Region/Landesvorgaben eine Rolle spielen. Sichtbare Unterschiede, aber "
    "mit Vorsicht zu lesen: kleine Bundesländer haben wenige Häuser, ein einzelnes Haus kann den "
    "Landeswert stark verschieben.",
    farbe=GRAU, size=9.5
)

# G8
bild(GRAFIKEN / "g8_korrelation.png", breite_cm=13,
     titel="Grafik 8 — Korrelationsmatrix: Alle Merkmale auf einen Blick")
body(
    "Warum wichtig: Fasst die gesamte deskriptive Analyse in einer Zahl pro Merkmal zusammen — die "
    "kompakteste Übersicht, welches Merkmal am stärksten mit hat_viele_Probleme zusammenhängt "
    "(aerzte_pro_bett und pflege_pro_bett vorn, ist_konzern und fortbildungsquote praktisch bei null).",
    farbe=GRAU, size=9.5
)

# G9
bild(GRAFIKEN / "g9_scatter_betten_aerzte.png", breite_cm=13,
     titel="Grafik 9 — Streudiagramm: Bettenzahl vs. Ärzte pro Bett")
body(
    "Warum wichtig: Prüft, ob sich die beiden Gruppen anhand von zwei Merkmalen gemeinsam trennen "
    "lassen — mehr Information als eine einzelne Korrelationszahl. Starke Überlappung bestätigt: Mit "
    "diesen beiden Merkmalen allein lässt sich kein Haus zuverlässig zuordnen.",
    farbe=GRAU, size=9.5
)

# G10
bild(GRAFIKEN / "g10_stoerfaktor_traeger.png", breite_cm=13,
     titel="Grafik 10 — Störfaktor: Private Häuser (Md 125 Betten) deutlich kleiner als öffentliche (233)")
body(
    "Warum wichtig: Der kritischste Kontroll-Befund im ganzen Notebook. Er zeigt, dass der Träger-"
    "Effekt aus Grafik 3 teilweise ein versteckter Größen-Effekt sein könnte — kleinere Häuser haben "
    "pro Indikator weniger Fälle und damit mehr statistische Schwankung. Ohne diese Grafik würde man "
    "'privat = schlechter' zu unkritisch stehen lassen.",
    farbe=GRAU, size=9.5
)

# G11
bild(GRAFIKEN / "g11_pflege_pro_bett.png", breite_cm=13,
     titel="Grafik 11 — Pflegekräfte pro Bett: Wenige Probleme Md=1,041, Viele Probleme Md=0,892")
body(
    "Warum wichtig: Zeigt dasselbe Muster wie Ärzte pro Bett und bestätigt sich später ebenfalls per "
    "T-Test (p<0,001) — wird zum zweitwichtigsten Merkmal im Decision Tree. Die Ähnlichkeit zu Grafik 6 "
    "ist zugleich ein Hinweis auf Multikollinearität (beide Merkmale korrelieren auch untereinander).",
    farbe=GRAU, size=9.5
)

# G12
bild(GRAFIKEN / "g12_konzern_vergleich.png", breite_cm=13,
     titel="Grafik 12 — Konzernhaus (49,7 %) vs. unabhängiges Haus (49,2 %)")
body(
    "Warum wichtig: Testet eine von Kollegen vorgeschlagene Hypothese (zentrale Qualitätssicherung im "
    "Konzern). Praktisch identische Werte — später per Chi²-Test bestätigt (p=0,90, klar nicht "
    "signifikant). Bewusst trotzdem im Modell belassen: 'Kein Zusammenhang' ist ein valider, "
    "dokumentierter Befund.",
    farbe=GRAU, size=9.5
)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════
# 4. PROJEKTSTRUKTUR & REPRODUZIERBARKEIT
# ═════════════════════════════════════════════════════════════════
h1("4  Projektstruktur & Reproduzierbarkeit")

h2("4.1  Datei-Übersicht")
tabelle(
    ["Datei / Ordner", "Inhalt"],
    [
        ("Notebooks/01_Exploration.ipynb", "Datenvorbereitung: CSV-Analyse → analysetabelle.csv"),
        ("Notebooks/02_Analyse.ipynb",     "Deskriptive Analyse: 12 Grafiken, 5 Tests"),
        ("Data/analysetabelle.csv",         "Zentrale Analysetabelle: 1.824 × 18"),
        ("grafiken/g1–g12_*.png",           "Alle 12 Grafiken als PNG"),
        ("scripts/Grafiken_Speichern.py",   "PNG-Export aller Grafiken (ohne Notebook)"),
        ("Doku/MD/01_Exploration.md",        "Schritt-für-Schritt-Erklärung Notebook 1"),
        ("Doku/MD/02_Analyse.md",            "Schritt-für-Schritt-Erklärung Notebook 2"),
    ],
    col_widths=[6.5, 8.5]
)

h2("4.2  Reproduzierbarkeit")
body(
    "Der gesamte Analyse-Pipeline ist vollständig reproduzierbar: "
    "Aus den Rohdaten in Data/CSV/ erzeugt das Ausführen von 01_Exploration.ipynb "
    "die analysetabelle.csv, 02_Analyse.ipynb daraus alle 12 Grafiken. "
    "Kein manueller Schritt ist nötig."
)
bullet("End-to-end-Test durchgeführt: ", "2026-07-29 — fehlerfrei von Rohdaten bis Grafiken")
bullet("Versionskontrolle: ", "Git-Repository · Rohdaten per .gitignore ausgeschlossen (zu groß)")
bullet("Abhängigkeiten: ", "requirements.txt enthält alle 7 Pakete mit Versionen")

doc.add_paragraph()
trennlinie()
body(
    f"Stand: {datetime.date.today().strftime('%d.%m.%Y')} · "
    "Nächster Schritt (Woche 3): Streamlit-Dashboard + Decision Tree",
    farbe=GRAU, size=10, italic=True
)

# ── Speichern ─────────────────────────────────────────────────────
OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(OUT_PATH))
print(f"Gespeichert: {OUT_PATH}")
