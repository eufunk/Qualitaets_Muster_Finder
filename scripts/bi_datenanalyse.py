"""
Erstellt BI_Datenanalyse.docx im TMP-Ordner:
Stellungnahme zur Dateiauswahl der Kollegen (BI-Tool-Analyse)
im Vergleich zu unserer Python-Exploration (01_Exploration.ipynb).
"""
import os, tempfile
from datetime import date, datetime
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TMP = tempfile.gettempdir()
OUT = os.path.join(TMP, "BI_Datenanalyse.docx")

# ── Farben ────────────────────────────────────────────────────────
C_BLAU   = RGBColor(0x1F, 0x49, 0x7D)
C_GRUEN  = RGBColor(0x37, 0x5E, 0x23)
C_ORANGE = RGBColor(0xBF, 0x5A, 0x00)
C_ROT    = RGBColor(0x7B, 0x0D, 0x0D)
C_WEISS  = RGBColor(0xFF, 0xFF, 0xFF)
F_GRUEN  = "E2EFDA"
F_GELB   = "FFF2CC"
F_ROT    = "FCE4D6"
F_BLAU   = "DEEAF1"
F_GRAU   = "F2F2F2"
F_ORANGE = "FCE9D9"

def set_bg(cell, fill):
    tc = cell._tc; p = tc.get_or_add_tcPr()
    s = OxmlElement("w:shd"); s.set(qn("w:fill"), fill); s.set(qn("w:val"), "clear")
    p.append(s)

def hdr_cell(cell, fill):
    set_bg(cell, fill)
    for run in cell.paragraphs[0].runs:
        run.font.color.rgb = C_WEISS; run.bold = True; run.font.size = Pt(9.5)

def add_tbl(doc, headers, rows, cw, hdr_color, row_fills):
    t = doc.add_table(rows=1+len(rows), cols=len(headers)); t.style = "Table Grid"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h; hdr_cell(c, hdr_color)
    for ri, row in enumerate(rows):
        fill = row_fills[ri] if ri < len(row_fills) else F_BLAU
        for ci, v in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = str(v)
            c.paragraphs[0].runs[0].font.size = Pt(9)
            set_bg(c, fill)
    for row in t.rows:
        for i, w in enumerate(cw): row.cells[i].width = Cm(w)
    return t

def add_h(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs: run.font.color.rgb = color
    return p

def add_b(doc, text, size=10.5, color=None, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    if bold: r.bold = True
    return p

def add_infobox(doc, text, fill, text_color=None):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.size = Pt(10)
    if text_color: r.font.color.rgb = text_color
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), fill); shd.set(qn("w:val"), "clear")
    pPr.append(shd)
    return p

# ══════════════════════════════════════════════════════════════════
doc = Document()
doc.core_properties.author = "Datenanalyse-Team"
doc.core_properties.created = datetime.now()
doc.core_properties.modified = datetime.now()
doc.core_properties.title = "BI-Datenanalyse — Stellungnahme zur Dateiauswahl"
for s in doc.sections:
    s.top_margin=Cm(2.0); s.bottom_margin=Cm(2.0)
    s.left_margin=Cm(2.5); s.right_margin=Cm(2.0)

# ── Titelseite ────────────────────────────────────────────────────
tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
rt = tp.add_run("BI-Datenanalyse — Stellungnahme zur Dateiauswahl")
rt.bold=True; rt.font.size=Pt(18); rt.font.color.rgb=C_BLAU

doc.add_paragraph()
sp = doc.add_paragraph(); sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
sp.add_run("Projekt: Qualitäts-Muster-Finder  |  Datensatz: 86 CSV-Dateien, Qualitätsberichte 2023 (IQTIG / G-BA)").font.size = Pt(11)

dp = doc.add_paragraph(); dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
dp.add_run(f"Datum: {date.today().strftime('%d.%m.%Y')}  |  Erstellt von: Datenanalyse-Team").font.size = Pt(10)

doc.add_paragraph()
add_infobox(doc,
    "Kontext: Die Kollegen haben den Datensatz mit einem BI-Tool (z.B. Power BI / Tableau) "
    "eigenständig exploriert und 21 Dateien in drei thematischen Gruppen als relevant eingestuft. "
    "Dieses Dokument vergleicht ihre Auswahl mit unserer Python-basierten Exploration "
    "(01_Exploration.ipynb) und gibt zu jeder Datei eine Stellungnahme ab.",
    F_BLAU
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# LEGENDE
# ══════════════════════════════════════════════════════════════════
add_h(doc, "Legende", level=1, color=C_BLAU)
add_tbl(doc,
    ["Symbol / Farbe", "Bedeutung"],
    [
        ("✅  VERW  (grün)",   "Von uns bereits in der Analyse verwendet"),
        ("⚠️  MOEG  (gelb)",   "Von uns als möglicherweise relevant identifiziert — noch nicht eingebunden"),
        ("❌  NEIN  (rot)",    "Von uns als nicht relevant klassifiziert"),
        ("🔵  NEU   (blau)",   "Von Kollegen vorgeschlagen — nicht in unserer ursprünglichen Shortlist"),
        ("⚠️  DSGVO  (orange)","Enthält personenbezogene Daten — besondere Vorsicht erforderlich"),
    ],
    cw=[4.5, 14.0], hdr_color="1F497D",
    row_fills=[F_GRUEN, F_GELB, F_ROT, F_BLAU, F_ORANGE]
)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════
# GRUPPE 1: Psychiatrie & Einrichtungstypen
# ══════════════════════════════════════════════════════════════════
add_h(doc, "1  Psychiatrie & Einrichtungstypen", level=1, color=C_BLAU)
add_infobox(doc,
    "⚠️  Wichtiger technischer Hinweis: Die meisten Dateien dieser Gruppe "
    "(QS.Pso, QS.Psy, QS.Struktur.Station) verwenden QS.Einrichtung.ID als Schlüssel — "
    "NICHT SO.QBID. Das bedeutet: Ein direkter Join mit unserer Haupttabelle SO.csv "
    "ist ohne Brückentabelle nicht möglich. Diese Einschränkung muss bei der "
    "weiteren Analyse berücksichtigt werden.",
    "FFF2CC", C_ORANGE
)
doc.add_paragraph()

add_tbl(doc,
    ["Datei (Kollegen)", "Unser Status", "Stellungnahme"],
    [
        ("QS.Einrichtungstypen\n★ Zentrales Element (Kollegen)",
         "❌ NEIN\n(Lookup-Tabelle)",
         "Unsere Einschätzung: reine Code-Dekodierungstabelle — übersetzt Einrichtungstyp-Codes in Bezeichnungen, enthält keine Messwerte. "
         "ABER: Wenn die Kollegen damit Datensätze nach Einrichtungstyp gefiltert haben (z.B. nur Akutkrankenhäuser), "
         "ist das ein valider Preprocessing-Schritt. Als eigenständige Analysedatei bleibt sie irrelevant. "
         "Empfehlung: Als Filter-Hilfstabelle nutzen, nicht als Feature-Quelle."),
        ("QS.Behandlungsumfang",
         "⚠️ MOEG\n(noch nicht gesichtet)",
         "Deckt sich mit unserer Einschätzung. Enthält möglicherweise Informationen zum Umfang der QS-Leistungen "
         "eines Hauses — potenziell interessant als Strukturmerkmal. Bisher nicht gesichtet. "
         "Empfehlung: Gemeinsam sichten und prüfen, ob QSLB-Informationen ergänzend zur QS.Leistungsbereich.csv sind."),
        ("QS.Pso\n(Personalstunden Psychiatrie)",
         "⚠️ MOEG\n(zu spezifisch)",
         "Schlüssel: QS.Einrichtung.ID — kein direkter Join mit SO.csv ohne Brücke. "
         "Sinnvoll NUR für eine auf Psychiatrie-Häuser beschränkte Teilanalyse. "
         "Für die allgemeine Projektfragestellung (alle 1.824 Häuser mit Bewertung) eingeschränkt verwendbar."),
        ("QS.Psy\n(Psychiatrie-Qualitätsdaten)",
         "⚠️ MOEG\n(zu spezifisch)",
         "Gleiche technische Einschränkung wie QS.Pso. Enthält psychiatriespezifische Qualitätsindikatoren. "
         "Wenn das Ziel eine allgemeine Analyse ist: nicht einbinden. Wenn Psychiatrie-Fokus: prüfen."),
        ("QS.Struktur.Station\n(Stationsstruktur Psychiatrie)",
         "⚠️ MOEG\n(zu spezifisch)",
         "Stationsstruktur (Planbetten, Stationstyp) für psychiatrische Stationen. "
         "Gleiche Einschränkung. Für allgemeine Analyse nicht geeignet."),
        ("QS.Berufsgruppen",
         "❌ NEIN\n(Lookup-Tabelle)",
         "Reine Lookup-Tabelle: übersetzt Berufsgruppen-Codes in Bezeichnungen. "
         "Kann als Legende beim Interpretieren von FA.Personalliste.csv nützlich sein, "
         "enthält aber selbst keine Analysedaten. Empfehlung: nur als Nachschlagewerk."),
    ],
    cw=[3.8, 2.5, 12.2], hdr_color="1F497D",
    row_fills=[F_GELB, F_GELB, F_GELB, F_GELB, F_GELB, F_ROT]
)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════
# GRUPPE 2: Standort-Stammdaten & Allgemeine QS
# ══════════════════════════════════════════════════════════════════
add_h(doc, "2  Standort-Stammdaten & Allgemeine Qualitätssicherung", level=1, color=C_BLAU)
add_infobox(doc,
    "✅  Diese Gruppe deckt sich gut mit unserer Analyse. Die Kerndateien "
    "(SO, QS, QS.Qualitätsindikator, QS.Fortbildung) sind identisch mit unserer Auswahl. "
    "Ergänzend haben die Kollegen SO.Personalliste, QS.Landesrecht, Konzern und "
    "Akademische_Lehre hinzugefügt — dazu unsere Stellungnahme unten.",
    F_GRUEN
)
doc.add_paragraph()

add_tbl(doc,
    ["Datei (Kollegen)", "Unser Status", "Stellungnahme"],
    [
        ("SO\n★ Zentrale Haupttabelle",
         "✅ VERW\n(verwendet)",
         "Vollständige Übereinstimmung. SO.csv ist unsere Ankertabelle: alle Strukturmerkmale "
         "(Betten, Träger, Bundesland, Uni), alle Geo-Koordinaten, universeller Schlüssel SO.QBID."),
        ("SO.Personalliste",
         "✅ VERW\n(seit 2026-07-29)",
         "Ursprünglich als Alternativquelle eingestuft — FA.Personalliste.csv wurde für "
         "aerzte_pro_bett bevorzugt, weil sie feinere Aufschlüsselung nach Berufsgruppe ermöglicht. "
         "SO.Personalliste.csv wurde inzwischen aber selbst eingebunden: direkte Quelle für "
         "pflege_pro_bett (Filter SO.Personal.Bereich == 'Pflege'), da SO.QBID + Anzahl direkt "
         "vorhanden sind — kein Umweg über FA.csv nötig. 2. wichtigstes Merkmal (FI 23,8 %)."),
        ("QS.Leistungsbereich",
         "⚠️ MOEG\n(weiterhin offen)",
         "Übereinstimmung. Wir haben diese Datei als relevant identifiziert (QSLB.Dokumentationsrate "
         "= potenzielle Qualitätskennzahl), aber bis heute (Stand 2026-07-30) nicht eingebunden. "
         "Unterstützt die Einschätzung der Kollegen — sollte eingebunden werden."),
        ("QS.Fortbildung",
         "✅ VERW\n(verwendet)",
         "Vollständige Übereinstimmung. Fortbildungsquote ist explizit in der Aufgabenstellung "
         "als zu untersuchendes Merkmal genannt."),
        ("QS.Landesrecht",
         "❌ NEIN\n(nicht vergleichbar)",
         "Hier weichen wir ab. QS.Landesrecht enthält länderspezifische QS-Anforderungen — "
         "diese sind nicht bundeseinheitlich und damit schlecht vergleichbar über alle 1.824 Häuser. "
         "Risiko: Häuser in Bundesländern mit strengeren Landesregeln erscheinen systematisch 'schlechter', "
         "obwohl sie nur andere Anforderungen erfüllen müssen. "
         "Empfehlung: Nur einbinden, wenn länderspezifische Analyse gewünscht ist."),
        ("QS.Qualitätsindikator",
         "✅ VERW\n(Ziel-Variable)",
         "Vollständige Übereinstimmung. Dies ist die wichtigste Datei des gesamten Projekts — "
         "Grundlage der Ziel-Variable hat_viele_Probleme. Technischer Hinweis: 911 MB, "
         "nur per Python ladbar."),
        ("QS",
         "⚠️ MOEG\n(nicht benötigt)",
         "Korrektur (2026-07-30): QS.csv wurde ursprünglich als notwendige Brückentabelle "
         "QS-Berichtsbasis ↔ Standort eingeschätzt. Tatsächlich wird sie nie geladen — "
         "QS.Qualitätsindikator.csv trägt SO.QBID bereits selbst, ein Join über QS.csv ist "
         "für die Zusammenführung nicht nötig. QS.Typ (bund/land) könnte für eine Filterung "
         "genutzt werden, wurde aber bisher nicht geprüft."),
        ("Konzern",
         "✅ VERW\n(seit 2026-07-29)",
         "Gute Ergänzung — inzwischen eingebunden als Merkmal ist_konzern. Dabei wurde ein "
         "Join-Bug gefunden und behoben: Konzern.csv nutzt SO.Standortnummer als Schlüssel, "
         "nicht SO.QBID; der erste Versuch verglich fälschlich gegen SO.QBID (0 Treffer). "
         "Nach Korrektur: 358 von 1.824 Häusern (19,6 %) sind Konzernhäuser. Chi²-Test zeigt "
         "aber keinen signifikanten Zusammenhang mit Qualitätsproblemen (p=0,90) — der Decision "
         "Tree bestätigt das mit 0 % Feature Importance."),
        ("Akademische_Lehre",
         "❌ NEIN\n(in SO.Uni enthalten)",
         "Hier weichen wir ab. Der Lehrstatus eines Krankenhauses ist bereits in SO.csv "
         "über die Spalte SO.Uni (1 = Universitätsklinik, 0 = sonstiges) abgebildet. "
         "Akademische_Lehre könnte Detailinformationen enthalten, die SO.Uni ergänzen — "
         "das sollte geprüft werden. Als Standalone-Datei ohne diesen Check: redundant."),
    ],
    cw=[3.8, 2.5, 12.2], hdr_color="1F497D",
    row_fills=[F_GRUEN, F_GRUEN, F_GELB, F_GRUEN, F_ROT, F_GRUEN, F_GELB, F_GRUEN, F_ROT]
)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════
# GRUPPE 3: Fachabteilungen & Abteilungs-Personal
# ══════════════════════════════════════════════════════════════════
add_h(doc, "3  Fachabteilungen & Abteilungs-Personal", level=1, color=C_BLAU)
add_infobox(doc,
    "✅  Gute Übereinstimmung bei den Kerndateien (FA, FA.Personalliste). "
    "Wichtiger Hinweis zu FA.Personen: Diese Datei enthält personenbezogene Daten "
    "(Namen, E-Mail-Adressen von Ärzten) und ist daher DSGVO-relevant.",
    F_BLAU
)
doc.add_paragraph()

add_tbl(doc,
    ["Datei (Kollegen)", "Unser Status", "Stellungnahme"],
    [
        ("FA\n★ Zentrale Abteilungstabelle",
         "✅ VERW\n(verwendet)",
         "Vollständige Übereinstimmung. FA.csv ist die notwendige Brückentabelle: "
         "verbindet ABTID (aus FA.Personalliste) mit FA.QBID = SO.QBID. "
         "Ohne FA.csv können Personaldaten nicht dem richtigen Haus zugeordnet werden."),
        ("FA.Personen",
         "⚠️ MOEG\n(DSGVO!)",
         "⚠️ DSGVO-Hinweis: FA.Personen enthält personenbezogene Daten — "
         "Vorname, Nachname, E-Mail und Telefon der ärztlichen Leitungen. "
         "Diese Daten dürfen nur im Rahmen des zulässigen Zwecks verarbeitet werden. "
         "Der Data/-Ordner ist per .gitignore vom Repository ausgeschlossen. "
         "Für Analysezwecke (Merkmale aggregieren) ist FA.Personalliste.csv ausreichend — "
         "FA.Personen liefert keine aggregierten Kennzahlen, sondern Einzelpersonen-Daten. "
         "Empfehlung: Nicht einbinden; FA.Personalliste.csv verwenden."),
        ("FA.Personalliste",
         "✅ VERW\n(verwendet)",
         "Vollständige Übereinstimmung. FA.Personalliste.csv ist die Hauptquelle für "
         "aerzte_pro_bett — das Merkmal mit der höchsten Feature Importance (53,6 %) "
         "im Decision Tree, seit der Ergänzung um pflege_pro_bett und ist_konzern am 2026-07-29. "
         "Wichtig: FA.Personal.Anzahl ist Komma-Dezimal ('13,47') "
         "und muss vor der Aggregation zu float konvertiert werden."),
        ("AQ.Ärzte",
         "⚠️ MOEG\n(Qualifikation, keine Anzahlen)",
         "Interessante Ergänzung — aber andere Aussage als FA.Personalliste. "
         "AQ.Ärzte enthält Qualifikationsmerkmale (Facharzt-Bezeichnungen, Weiterbildungen), "
         "aber keine Vollzeit-Äquivalente. Für aerzte_pro_bett ist FA.Personalliste besser. "
         "Sinnvoll wenn Qualifikationsniveau (z.B. Anteil Fachärzte) als Merkmal untersucht werden soll. "
         "Empfehlung: Einbinden wenn Qualifikationsfrage relevant ist."),
        ("AQ.Pflege",
         "❌ NEIN\n(Ersatz gewählt)",
         "Pflegekräfte pro Bett ist laut Aufgabenstellung (Fragestellung.docx) ein explizit "
         "zu untersuchendes Merkmal und ist seit 2026-07-29 in der Analysetabelle (pflege_pro_bett). "
         "AQ.Pflege.csv selbst wurde aber bewusst NICHT verwendet: sie enthält nur "
         "Qualifikationsnachweise, keine Personal-Anzahlen. SO.Personalliste.csv liefert "
         "SO.QBID + SO.Personal.Anzahl direkt — kein Umweg über FA.csv nötig und detaillierter "
         "als AQ.Pflege.csv. Ergebnis: 2. wichtigstes Merkmal im Decision Tree (FI 23,8 %)."),
    ],
    cw=[3.8, 2.5, 12.2], hdr_color="1F497D",
    row_fills=[F_GRUEN, F_ORANGE, F_GRUEN, F_GELB, F_ROT]
)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════
# GESAMTBEWERTUNG
# ══════════════════════════════════════════════════════════════════
add_h(doc, "4  Gesamtbewertung & Empfehlungen", level=1, color=C_BLAU)

add_infobox(doc,
    "Stand 2026-07-30: Diese Tabelle wurde aktualisiert, nachdem Konzern und SO.Personalliste "
    "tatsächlich eingebunden wurden und ein Join-Bug bei Konzern.csv sowie eine Fehleinschätzung "
    "zu QS.csv aufgefallen sind (siehe Gruppe 2 oben).",
    F_BLAU
)
doc.add_paragraph()

add_tbl(doc,
    ["Kategorie", "Anzahl Dateien", "Bewertung"],
    [
        ("Tatsächlich in Analysetabelle eingebunden (✅)",
         "7",
         "SO, QS.Qualitätsindikator, QS.Fortbildung, FA, FA.Personalliste, SO.Personalliste, Konzern\n"
         "→ Kernauswahl bestätigt; SO.Personalliste und Konzern kamen am 2026-07-29 hinzu."),
        ("Identifiziert, aber weiterhin nicht eingebunden (⚠️)",
         "7",
         "QS (nicht benötigt — SO.QBID bereits in QS.Qualitätsindikator enthalten), "
         "QS.Leistungsbereich, QS.Behandlungsumfang, AQ.Ärzte, QS.Pso, QS.Psy, QS.Struktur.Station\n"
         "→ QS.Leistungsbereich bleibt der aussichtsreichste Kandidat für eine künftige Ergänzung."),
        ("Lookup-Tabellen — kein Analysewert als Feature (❌)",
         "2",
         "QS.Einrichtungstypen, QS.Berufsgruppen\n"
         "→ Als Filter/Legende nützlich, nicht als Feature-Quelle."),
        ("Abweichung / bewusster Ersatz (❌)",
         "3",
         "QS.Landesrecht (nicht vergleichbar über Bundesländer), "
         "Akademische_Lehre (in SO.Uni enthalten), "
         "AQ.Pflege (durch SO.Personalliste ersetzt, da detaillierter)\n"
         "→ Bewusste, begründete Entscheidungen — kein offener Punkt."),
        ("DSGVO-Hinweis",
         "1",
         "FA.Personen → personenbezogene Daten (Namen, E-Mails).\n"
         "→ Nicht für Feature-Engineering verwenden."),
    ],
    cw=[5.5, 2.0, 11.0], hdr_color="1F497D",
    row_fills=[F_GRUEN, F_GELB, F_GRAU, F_ROT, F_ORANGE]
)
doc.add_paragraph()

add_b(doc,
    "Fazit: Die BI-Tool-Analyse der Kollegen war insgesamt gut und hat unsere Kernauswahl "
    "bestätigt sowie zwei echte Lücken aufgedeckt. Konzern und SO.Personalliste (für "
    "pflege_pro_bett) wurden am 2026-07-29 eingebunden — bei Konzern.csv fiel dabei zusätzlich "
    "ein Join-Bug auf (falscher Schlüssel, 0 Treffer), der erst durch die erneute Prüfung dieser "
    "Empfehlung entdeckt wurde. AQ.Pflege wurde bewusst NICHT übernommen, da SO.Personalliste.csv "
    "denselben Zweck detaillierter erfüllt. Bei der Psychiatrie-Gruppe bleibt der abweichende "
    "Join-Schlüssel (QS.Einrichtung.ID) weiterhin ungelöst — nur relevant für eine "
    "Psychiatrie-Teilanalyse.",
    size=10.5, color=C_BLAU, bold=True
)

# ── Speichern ─────────────────────────────────────────────────────
doc.save(OUT)
print(f"✅ Gespeichert: {OUT}")
