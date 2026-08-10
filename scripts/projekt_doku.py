"""
word_dokumentation.py
=====================
Erzeugt die Word-Projektdokumentation "Dokumentation_Qualitaets_Muster_Finder.docx"
aus den Daten der analysetabelle.csv.

Ursprung:
    Teil "Word-Dokumentation generieren" aus 01_Exploration.ipynb —
    in diese eigenstaendige Skriptdatei ausgelagert.

Aufruf:
    python scripts/word_dokumentation.py
    (oder aus dem Projekt-Root:  python scripts/word_dokumentation.py)

Abhaengigkeiten:
    pandas, python-docx  (python-docx wird bei Bedarf automatisch installiert)
"""

import os
import subprocess
import sys
from datetime import date, datetime
from pathlib import Path

import pandas as pd

# ── python-docx sicherstellen ─────────────────────────────────────
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

# ── Projekt-Root (eine Ebene ueber /scripts) ──────────────────────
PROJEKT_ROOT = Path(__file__).resolve().parent.parent
os.chdir(PROJEKT_ROOT)

DATA_PATH  = PROJEKT_ROOT / "Data" / "analysetabelle.csv"
OUT_PATH   = PROJEKT_ROOT / "Doku" / "Word" / "Dokumentation_Qualitaets_Muster_Finder.docx"


# ══════════════════════════════════════════════════════════════════
# DATEN LADEN & KENNGROESSEN BERECHNEN
# ══════════════════════════════════════════════════════════════════
def lade_daten() -> dict:
    """
    Laedt analysetabelle.csv und berechnet alle Kenngroessen,
    die fuer die Word-Dokumentation benoetigt werden.

    Returns:
        dict mit den Schluesseln:
            analyse, auffaellig_quote, median_quote, n_haeuser,
            n_auffaellig, stats, total_qi_sum
    """
    analyse = pd.read_csv(DATA_PATH, low_memory=False)

    # auffaellig_quote als eigenstaendiger DataFrame (wie im Notebook)
    quote_cols = ["SO.QBID", "total_qi", "auffaellig_n",
                  "auffaellig_quote", "hat_viele_Probleme"]
    auffaellig_quote = analyse[[c for c in quote_cols if c in analyse.columns]].copy()

    median_quote   = auffaellig_quote["auffaellig_quote"].median()
    n_haeuser      = len(auffaellig_quote)
    n_auffaellig   = int(auffaellig_quote["hat_viele_Probleme"].sum())
    stats          = auffaellig_quote["auffaellig_quote"].describe()
    total_qi_sum   = int(analyse["total_qi"].sum())  # = len(qi_dedup) im Notebook

    return {
        "analyse":          analyse,
        "auffaellig_quote": auffaellig_quote,
        "median_quote":     median_quote,
        "n_haeuser":        n_haeuser,
        "n_auffaellig":     n_auffaellig,
        "stats":            stats,
        "total_qi_sum":     total_qi_sum,
    }


# ══════════════════════════════════════════════════════════════════
# HILFSFUNKTIONEN FUER DAS WORD-DOKUMENT
# ══════════════════════════════════════════════════════════════════
def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.style.font.size = Pt(11)
    return p


def add_body_bold(doc, parts):
    """parts = list of (text, bold) tuples"""
    p = doc.add_paragraph()
    for text, bold in parts:
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(11)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    p.style.font.size = Pt(11)
    return p


def add_table(doc, headers, rows, col_widths=None, bold_cols=None, skip_bold_rows=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    # Header-Zeile
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        run = hdr[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)
        tc = hdr[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "4472C4")
        shd.set(qn("w:color"), "FFFFFF")
        shd.set(qn("w:val"), "clear")
        tcPr.append(shd)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # Daten-Zeilen
    for r_i, row in enumerate(rows):
        cells = table.rows[r_i + 1].cells
        for c_i, val in enumerate(row):
            cells[c_i].text = str(val)
            run = cells[c_i].paragraphs[0].runs[0]
            run.font.size = Pt(10)
            if bold_cols and c_i in bold_cols:
                if not (skip_bold_rows and r_i in skip_bold_rows):
                    run.bold = True
    # Spaltenbreiten
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return table


# ══════════════════════════════════════════════════════════════════
# DOKUMENT ERZEUGEN
# ══════════════════════════════════════════════════════════════════
def erzeuge_dokument(daten: dict) -> Document:
    """Baut das komplette Word-Dokument auf und gibt es zurueck."""
    analyse          = daten["analyse"]
    auffaellig_quote = daten["auffaellig_quote"]
    median_quote     = daten["median_quote"]
    n_haeuser        = daten["n_haeuser"]
    n_auffaellig     = daten["n_auffaellig"]
    stats            = daten["stats"]
    total_qi_sum     = daten["total_qi_sum"]

    doc = Document()
    doc.core_properties.author = "Datenanalyse-Team"
    doc.core_properties.created = datetime.now()
    doc.core_properties.modified = datetime.now()

    # ── Seitenraender ──────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # ══════════════════════════════════════════════════════════════
    # TITELSEITE
    # ══════════════════════════════════════════════════════════════
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Qualitäts-Muster-Finder")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("Projektdokumentation — Stand der Datenaufbereitung").font.size = Pt(14)

    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run(
        f"Erstellt: {date.today().strftime('%d.%m.%Y')}    |    "
        "Datenbasis: Qualitätsberichte deutscher Krankenhäuser 2023"
    ).font.size = Pt(11)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # INHALTSVERZEICHNIS (manuell)
    # ══════════════════════════════════════════════════════════════
    add_heading(doc, "Inhaltsverzeichnis", level=1)
    toc_entries = [
        ("1",   "Projektübersicht",                         "3"),
        ("1.1", "Fragestellung",                            "3"),
        ("1.2", "Projektziel & Bausteine",                  "3"),
        ("2",   "Datensatz",                                "4"),
        ("2.1", "Überblick",                                "4"),
        ("2.2", "Schlüssel-ID: SO.QBID",                    "4"),
        ("2.3", "Relevante Tabellen",                       "4"),
        ("2.4", "Möglicherweise relevante Tabellen",        "5"),
        ("2.5", "Nicht relevante Tabellen",                 "6"),
        ("3",   "Datenaufbereitung",                        "7"),
        ("3.1", "Vorgehen & Kriterien",                     "7"),
        ("3.2", "Ziel-Variable",                            "7"),
        ("3.3", "Merkmale (Features)",                      "8"),
        ("4",   "Ergebnisse",                               "9"),
        ("4.1", "Ziel-Variable — Statistiken",              "9"),
        ("4.2", "Analysetabelle",                           "9"),
        ("4.3", "Wozu wird die Analysetabelle genutzt?",    "10"),
        ("4.4", "Fehlende Werte",                           "11"),
        ("5",   "Deskriptive Analyse — Befunde",            "12"),
        ("5.1", "Übersicht der Befunde",                    "12"),
        ("5.2", "Grafiken",                                 "13"),
        ("5.3", "Gesamteinschätzung",                       "14"),
        ("6",   "Projektstand & Nächste Schritte",       "15"),
        ("6.1", "Abgeschlossene Bausteine",               "15"),
        ("6.2", "Noch offen",                             "16"),
        ("6.3", "Mögliche Erweiterungen",                 "16"),
        ("7",   "Fragestellung & Antworten",              "17"),
        ("7.1", "Die zentrale Fragestellung",             "17"),
        ("7.2", "Aufgabe 1: Daten vorbereiten",           "17"),
        ("7.3", "Aufgabe 2: Deskriptive Analyse",         "18"),
        ("7.4", "Aufgabe 3: Dashboard bauen",             "18"),
        ("7.5", "Bonus: Entscheidungsbaum",               "19"),
        ("7.6", "Gesamtfazit zur Fragestellung",          "19"),
    ]
    for nr, title_text, page in toc_entries:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5 * (title_text.count(".") + 0))
        run_nr = p.add_run(f"{nr}  {title_text}")
        run_nr.font.size = Pt(11)
        if "." not in nr:
            run_nr.bold = True

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 1. PROJEKTÜBERSICHT
    # ══════════════════════════════════════════════════════════════
    add_heading(doc, "1  Projektübersicht", level=1)

    add_heading(doc, "1.1  Fragestellung", level=2)
    add_body_bold(doc, [
        ("Welche Krankenhausmerkmale hängen damit zusammen, dass ein Haus "
         "überdurchschnittlich viele Qualitätsprobleme aufweist?", True)])
    add_body(doc,
        "Grundlage sind die jährlichen Qualitätsberichte aller deutschen Krankenhäuser. "
        "Jedes Haus berichtet über ~150 Qualitätsindikatoren. Bei manchen Indikatoren "
        "werden Häuser als 'rechnerisch auffällig' bewertet. Ziel ist es, strukturelle "
        "Merkmale (Größe, Personal, Träger, Region) zu identifizieren, die mit einer "
        "erhöhten Auffälligkeitsquote zusammenhängen.")
    add_body_bold(doc, [("Wichtiger Hinweis: ", True), ("Kein Zusammenhang ist auch ein valides Ergebnis.", False)])
    add_body(doc,
        "Hintergrund: Die Qualitätsindikatoren markieren Häuser als 'rechnerisch auffällig', "
        "wenn ihr Wert außerhalb eines Referenzbereichs liegt. Das ist aber nur ein statistisches Signal "
        "— kein Qualitätsurteil. Ob wirklich ein Problem dahintersteckt, klärt ein separates Prüfverfahren "
        "(Strukturierter Dialog). Außerdem gibt es viele Einflussfaktoren (Patientenmix, Größe, Spezialisierung), "
        "die wir nicht alle kontrollieren können. Wenn die Analyse also keinen klaren Zusammenhang "
        "zwischen Strukturmerkmalen und Auffälligkeit zeigt, ist das ein ehrliches und wissenschaftlich "
        "korrektes Ergebnis — kein Scheitern. Quelle: Aufgabenstellung/Text_Presentation.docx, Folie 7.")

    add_heading(doc, "1.2  Projektziel & Bausteine", level=2)
    add_body(doc, "Das Projekt ist in fünf Bausteine gegliedert:")
    bausteine = [
        ("Baustein 1", "Daten vorbereiten", "✅ Abgeschlossen"),
        ("Baustein 2", "Deskriptive Analyse", "✅ Abgeschlossen"),
        ("Baustein 3", "Streamlit-Dashboard / Power BI Dashboard", "✅ Live"),
        ("Baustein 4", "Entscheidungsbaum (Bonus)", "✅ Abgeschlossen"),
        ("Baustein 5", "Abschluss & Präsentation", "✅ Abgeschlossen"),
    ]
    add_table(doc, ["Baustein", "Beschreibung", "Status"], bausteine,
              col_widths=[3.5, 8.5, 3.0])
    doc.add_paragraph()

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 2. DATENSATZ
    # ══════════════════════════════════════════════════════════════
    add_heading(doc, "2  Datensatz", level=1)

    add_heading(doc, "2.1  Überblick", level=2)
    add_body(doc,
        "Der Datensatz besteht aus 86 CSV-Dateien im Data/-Ordner. "
        "Alle Daten stammen aus den offiziellen Qualitätsberichten "
        "deutscher Krankenhäuser (Berichtsjahr 2023) und werden vom IQTIG veröffentlicht. "
        "IQTIG steht für Institut für Qualitätssicherung und Transparenz im Gesundheitswesen — "
        "eine vom G-BA (Gemeinsamer Bundesausschuss) beauftragte Einrichtung, "
        "die jährlich die Qualitätsdaten aller deutschen Krankenhäuser erhebt, "
        "auswertet und veröffentlicht.")
    add_bullet(doc, "2.310 Krankenhäuser (Standorte) in SO.csv — davon 1.824 mit Qualitätsbewertung")
    add_bullet(doc, "~150 Qualitätsindikatoren pro Haus")
    add_bullet(doc, "Strukturdaten: Betten, Personal, Träger, Standort, Geo-Koordinaten")
    doc.add_paragraph()
    klassifikation = [
        ("✅  Relevant — aktiv verwendet",        "7",  "In analysetabelle.csv eingeflossen"),
        ("⚠️  Möglicherweise relevant",           "33", "Identifiziert, aber nicht eingebunden — Kandidaten für Folgeanalysen"),
        ("❌  Nicht relevant — ausgeschlossen",   "29", "Kein Analysebezug, DSGVO, Lookup-Tabellen, Verwaltungsdaten"),
        ("🔑  *.Key.csv (Lookup-Schlüssel)",      "17", "Technische Dekodierungstabellen — keine Analysedaten"),
        ("Gesamt", "86", ""),
    ]
    add_table(doc, ["Kategorie", "Anzahl Dateien", "Erläuterung"], klassifikation,
              col_widths=[5.5, 2.5, 8.0], bold_cols=[0, 1])

    add_heading(doc, "2.2  Schlüssel-ID: SO.QBID", level=2)
    add_body(doc,
        "Fast alle Tabellen sind über die Spalte SO.QBID miteinander verknüpft. "
        "SO.QBID ist die eindeutige ID eines Krankenhaus-Standorts und dient als "
        "primärer Join-Key für alle Merge-Operationen.")

    add_heading(doc, "2.3  Relevante Tabellen", level=2)
    rel_rows = [
        ("SO.csv",                    "Stammdaten aller Krankenhäuser (Haupttabelle)",
         "SO.QBID, SO.Betten, SO.Bundesland, SO.Uni, KH.Träger.Art, SO.Standortnummer, Koordinaten",
         "Kern-Merkmale"),
        ("QS.Qualitätsindikator.csv", "Qualitätsindikatoren mit Bewertungen (>50 MB)",
         "SO.QBID, QSErgBewStrukDialog, QSQI.Indikator, QSQI.ArtDesWertes",
         "Ziel-Variable"),
        ("QS.Fortbildung.csv",        "Fortbildungsnachweise der Ärzte",
         "SO.QBID, QS.Fortbildungspflichtige, QS.Fortbildungsnachweis_Erbracht_Habende",
         "Merkmal: Fortbildungsquote"),
        ("FA.csv",                    "Fachabteilungen der Krankenhäuser (Brückentabelle)",
         "ABTID, FA.QBID (= SO.QBID)",
         "Verknüpfungstabelle für Personaldaten"),
        ("FA.Personalliste.csv",      "Personal pro Fachabteilung nach Berufsgruppe",
         "ABTID, FA.Personal.Bereich, FA.Personal.Anzahl",
         "Merkmal: aerzte_pro_bett (stärkster Prädiktor)"),
        ("SO.Personalliste.csv",      "Personal auf Standortebene nach Berufsgruppe *(2026-07-29)*",
         "SO.QBID, SO.Personal.Bereich, SO.Personal.Anzahl",
         "Merkmal: pflege_pro_bett"),
        ("Konzern.csv",               "Konzernzugehörigkeit *(2026-07-29)*",
         "SO.Standortnummer, Konzern",
         "Merkmal: ist_konzern"),
    ]
    add_table(doc,
              ["Datei", "Inhalt", "Wichtige Spalten", "Rolle"],
              rel_rows,
              col_widths=[4.5, 4.5, 5.5, 3.5],
              bold_cols=[0, 3],
              skip_bold_rows=[3])  # FA.csv = Brückentabelle, kein Merkmal
    doc.add_paragraph()
    add_body(doc,
        "Hinweis: QS.csv enthält Verwaltungsmetadaten zur Qualitätsberichterstattung "
        "(IK-Nummer, Standortnummer, Berichtstyp, SO.QBID) — eine Registrierungstabelle, "
        "welche Häuser am QS-Verfahren teilnehmen. Sie wird nicht geladen, da "
        "QS.Qualitätsindikator.csv die SO.QBID bereits selbst trägt und "
        "ein zusätzlicher Join nicht nötig ist.")

    add_heading(doc, "2.4  Möglicherweise relevante Tabellen (identifiziert, nicht eingebunden)", level=2)
    add_body(doc,
        "Diese 33 Tabellen wurden gesichtet und als potenziell nützlich eingestuft, "
        "aber bewusst nicht in die aktuelle Analysetabelle aufgenommen. "
        "Sie sind kein Teil der Präsentation und des Modells dieser Version.")
    add_body_bold(doc, [
        ("Die drei wertvollsten Kandidaten für eine Erweiterung:", True)])
    add_bullet(doc,
        "QS.Leistungsbereich.csv (12 MB) — enthält QSLB.Dokumentationsrate je Haus und Leistungsbereich. "
        "Häuser mit lückenhafter Dokumentation fallen häufiger rechnerisch auffällig — "
        "dieser Störfaktor lässt sich mit dieser Datei teilweise kontrollieren. "
        "Join über SO.QBID direkt möglich.")
    add_bullet(doc,
        "Notfallversorgung.csv (363 KB) — enthält die Notfallversorgungsstufe (1/2/3) je Haus. "
        "Häuser der höchsten Stufe behandeln die schwersten Fälle — "
        "ein wichtiger Confounder, der im aktuellen Modell fehlt. "
        "Join über SO.QBID direkt möglich.")
    add_bullet(doc,
        "MM.csv (236 KB) — Mindestmengen-Compliance: Hat das Haus die gesetzlich "
        "vorgeschriebene Mindestfallzahl (z. B. 50 Knie-TEP/Jahr) erreicht? "
        "Ein binäres Strukturmerkmal mit möglichem Qualitätsbezug.")
    add_body_bold(doc, [
        ("Explorations-Notebook:", True),
        (" Notebooks/04_Potenzielle_Erweiterungen.ipynb — analysiert diese drei Dateien "
         "und bewertet, ob sie die Modellgüte verbessern würden.", False)])
    add_body(doc, "Alle weiteren Einträge dieser Kategorie im Detail:")
    moeglich_rows = [
        ("QS.Leistungsbereich.csv", "QSLB.Dokumentationsrate = potenzielle Qualitätskennzahl. Noch nicht gesichtet/eingebunden."),
        ("AQ.Pflege.csv", "Enthält nur Pflege-Qualifikationsnachweise, keine Personal-Anzahlen. SO.Personalliste.csv liefert pflege_pro_bett direkter — deshalb hier statt AQ.Pflege.csv verwendet."),
        ("AQ.Ärzte.csv", "Qualifikationsmerkmale der Ärzte (Facharzt-Bezeichnungen), aber keine Vollzeit-Anzahlen. FA.Personalliste.csv ist für aerzte_pro_bett besser geeignet."),
        ("QS.Behandlungsumfang.csv", "Möglicherweise Ergänzung zu QS.Leistungsbereich.csv. Noch nicht gesichtet."),
        ("QS.Extern.Sonstige.csv", "Alternative Ziel-Variable denkbar, aber keine einheitliche auffällig-Bewertung wie QS.Qualitätsindikator.csv."),
        ("QS.Pso.csv, QS.Psy.csv, QS.Struktur.Station.csv", "Psychiatrie-spezifische Qualitätsdaten. Schlüssel QS.Einrichtung.ID — kein direkter Join mit SO.QBID ohne Brückentabelle. Nur für Psychiatrie-Teilanalyse geeignet."),
        ("MM.csv, MM.Ausnahme.csv, MM.Leistungsberechtigung.Prognose.csv", "Mindestmengen — ob ein Haus gesetzliche Mindestfallzahlen erfüllt. Strukturmerkmal, nicht gesichtet."),
        ("AM.csv, AM.Leistung.csv, AM.VAVU.csv", "Ausstattungsmerkmale (medizinisch-technische Geräte). Nicht gesichtet."),
        ("CQ.csv", "Zertifizierungen/Strukturqualitätsvereinbarungen. Nicht verwendet."),
        ("BF.csv, BM.csv", "Behandlungsfelder/-möglichkeiten. Bedeutung für Analyse unklar."),
        ("AMTS.csv, AMTS_Massnahme.csv", "Arzneimitteltherapiesicherheit. Bedeutung für Auffälligkeitsquote unklar."),
        ("RM.csv, RM.Fallbesprechung.csv", "Risikomanagement-Systeme. Nicht gesichtet."),
        ("Notfallversorgung.csv, MP.csv", "Notfallversorgungsstufe / Mindestpersonalbedarf. Strukturmerkmale, nicht gesichtet."),
        ("HB.csv, HD.csv, HM.csv, WeitereHygiene.csv, KISS.csv", "Hygiene- und Infektionsdaten. HD.csv ist >50 MB. Potenzielles Qualitätsmerkmal, nicht gesichtet."),
        ("AA.csv", "Apparative Ausstattung (Geräte-Verfügbarkeit) — gehört zur AM-Familie, nicht gesichtet."),
        ("AMTS_InstrumentMassnahme.csv", "Detail-/Lookup-Tabelle zu AMTS_Massnahme.csv."),
        ("DMP.csv", "Teilnahme an Disease-Management-Programmen — potenzielles Strukturmerkmal, nie geprüft."),
        ("EF.csv, IF.csv", "Externe/interne Fachabteilungen — nicht gesichtet."),
        ("Mitbewerber_Betten.csv", "Bettenzahl benachbarter Wettbewerber — Marktdichte-/Konkurrenzmerkmal, im gesamten Projekt bisher nicht in Betracht gezogen."),
    ]
    add_table(doc, ["Datei(en)", "Warum (noch) nicht eingebunden"], moeglich_rows,
              col_widths=[6.0, 10.0])
    doc.add_paragraph()

    add_heading(doc, "2.5  Nicht relevante Tabellen", level=2)
    add_body(doc, "Folgende Tabellen wurden bewusst ausgeschlossen:")
    nicht_rel = [
        ("NM.csv",              "Nicht-medizinische Angebote (Parkplatz, Telefon, WLAN) — kein Analysebezug"),
        ("Personen.csv, FA.Personen.csv", "⚠️ DSGVO — personenbezogene Daten (Name, E-Mail, Telefon). Nicht in Analyse; Data/ per .gitignore ausgeschlossen"),
        ("QS.Einrichtungstypen.csv, QS.Berufsgruppen.csv", "Reine Lookup-/Dekodierungstabellen — kein eigener Analysewert"),
        ("ICD.Code.csv, OPS.csv, OPS.Code.csv", "Diagnose-/Prozedurenschlüssel (reine Lookup-Tabellen)"),
        ("Alle *.Key.csv (16 Dateien)", "Technische Schlüssel-/Lookup-Tabellen ohne eigene Analysedaten (z.B. AA.Key.csv, FA-verwandte Key-Dateien)"),
        ("Link.csv, LinkVersorgunggebieteSO.csv, Weiterführender_Link.csv", "Nur URLs zu externen Webseiten — keine Analysedaten"),
        ("QS.Nachweis.csv, BewertungStrukDialog.csv", "Technische Meta-Daten (Nachweiszeiträume, Erläuterung der Bewertungscodes)"),
        ("QS.Landesrecht.csv", "Länderspezifische QS-Anforderungen — nicht bundeseinheitlich vergleichbar, würde Häuser in strengeren Bundesländern systematisch benachteiligen"),
        ("GIQI.csv", "Geriatrische Indikatoren — fachspezifisch, keine allgemeine Relevanz"),
        ("VAVU.csv", "Versorgungsstruktur (14 MB) — nicht gesichtet, keine erkennbare Relevanz"),
        ("Schutzkonzept.csv, Praevention_Missbrauch_und_Gewalt.csv, Neuartige_Therapien.csv", "Zu spezifisch für die allgemeine Fragestellung"),
        ("Sicherstellungszuschlaege.csv, Pflegepersonalregelung.csv, ErfPersVorgaben.csv", "Verwaltungsdaten ohne direkten Qualitätsbezug"),
        ("Akademische_Lehre.csv, Lenkungsgremium.csv, ZV.csv", "Metadaten/Organisationsstruktur — Lehrstatus bereits über SO.Uni abgedeckt"),
        ("Abt.Zugang.csv", "Adress-/Kontaktdaten einer Abteilung, keine Analysedaten"),
        ("Abt301.csv", "Lookup-Tabelle (amtlicher Fachabteilungsschlüssel 301)"),
        ("Error.csv", "Technisches Fehlerprotokoll der Berichtserstellung, keine Inhaltsdaten"),
        ("Sicherstellungszuschlaege_Fachabteilungen.csv", "Abteilungsdetail zum bereits ausgeschlossenen Sicherstellungszuschlaege.csv"),
    ]
    for datei, grund in nicht_rel:
        add_bullet(doc, f"{datei} — {grund}")
    add_body(doc,
        "Vollständige Klassifikation aller 86 Dateien: siehe Doku/MD/Daten_Inhaltsverzeichnis.md.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 3. DATENAUFBEREITUNG
    # ══════════════════════════════════════════════════════════════
    add_heading(doc, "3  Datenaufbereitung", level=1)

    add_heading(doc, "3.1  Vorgehen & Kriterien", level=2)
    add_body(doc,
        "Die Datenaufbereitung erfolgte vollständig reproduzierbar per Python-Skript "
        "(01_Exploration.ipynb). Kein manuelles Zusammenklicken — alle Schritte können "
        "durch erneutes Ausführen des Notebooks repliziert werden.")
    add_body(doc,
        "Sichtung der CSV-Dateien: Erste 3–5 Zeilen jeder Datei wurden gelesen "
        "(kein Python nötig, da CSV = Textdatei). Kriterium für Relevanz:")
    add_bullet(doc, "✅ Relevant: Enthält Strukturmerkmal aus der Aufgabenstellung ODER Qualitätsindikator-Bewertungen")
    add_bullet(doc, "⚠️ Möglicherweise: QS-relevante Infos, aber Bedeutung noch unklar")
    add_bullet(doc, "❌ Nicht relevant: Lookup-Tabellen, nicht-medizinische Angebote, reine Link-/Verwaltungsdaten")

    add_heading(doc, "3.2  Ziel-Variable", level=2)
    add_body(doc,
        "Die Ziel-Variable wurde aus QS.Qualitätsindikator.csv berechnet. "
        "Die Bewertungsspalte QSErgBewStrukDialog enthält den Bewertungscode "
        "des Strukturierten Dialogs:")
    add_bullet(doc, "R* (R10, R20, ...) = rechnerisch auffällig")
    add_bullet(doc, "N01, N02 = nicht auffällig")
    add_bullet(doc, "N99 = nicht bewertet → wird ausgeschlossen (nicht bewertet ≠ unauffällig!)")

    add_body(doc, "Berechnungsschritte:")
    schritte = [
        ("1", "Filter: QSQI.ArtDesWertes == 'QI'",
         "Nur echte Qualitätsindikatoren, keine Zählkennzahlen (EKez, TKez, ...)"),
        ("2", "Filter: QSErgBewStrukDialog != 'N99'",
         "Nur bewertete Indikatoren berücksichtigen"),
        ("3", "Deduplizierung: drop_duplicates(['SO.QBID', 'QSQI.Indikator'])",
         "Je Haus+Indikator nur eine Zeile behalten"),
        ("4", "Flag: ist_auffaellig = QSErgBewStrukDialog.str.startswith('R')",
         "Binäre Markierung pro Indikator"),
        ("5", "Aggregation: groupby('SO.QBID').agg(count, sum)",
         "Anzahl bewerteter QI und auffälliger QI pro Haus"),
        ("6", "Quote: auffaellig_n / total_qi",
         "Anteil auffälliger Indikatoren pro Haus"),
        ("7", "Target: quote > Median → hat_viele_Probleme = 1",
         "Binäre Ziel-Variable (0/1)"),
    ]
    add_table(doc, ["Schritt", "Code", "Erklärung"], schritte,
              col_widths=[1.5, 6.5, 6.0])
    doc.add_paragraph()

    add_heading(doc, "3.3  Merkmale (Features)", level=2)
    add_body(doc, "Folgende Merkmale wurden für die Analysetabelle ausgewählt:")
    merkmale_rows = [
        ("SO.Betten",          "SO.csv",              "Direkt verfügbar",     "Numerisch"),
        ("KH.Träger.Art",      "SO.csv",              "Direkt verfügbar",     "Kategorial: privat / freigemeinnützig / öffentlich"),
        ("SO.Bundesland",      "SO.csv",              "Direkt verfügbar",     "Kategorial: 16 Bundesländer"),
        ("SO.Uni",             "SO.csv",              "Direkt verfügbar",     "Binär: 0/1"),
        ("SO.Latitude/Long.",  "SO.csv",              "Direkt verfügbar",     "Numerisch (für Karte)"),
        ("fortbildungsquote",  "QS.Fortbildung.csv",  "Berechnet: Erbracht / Pflichtige", "Numerisch 0–1"),
        ("aerzte_pro_bett",    "FA.Personalliste.csv + FA.csv", "Σ Ärzte / SO.Betten", "Numerisch — stärkster Prädiktor (FI 53,6 %)"),
        ("pflege_pro_bett",    "SO.Personalliste.csv", "Σ Pflegekräfte / SO.Betten", "Numerisch — 2. stärkster Prädiktor (FI 23,8 %), ergänzt 2026-07-29"),
        ("ist_konzern",        "Konzern.csv",         "SO.Standortnummer in Konzern.csv? → 0/1", "Binär — kein Zusammenhang gefunden (Chi² p=0,90), ergänzt 2026-07-29"),
    ]
    add_table(doc,
              ["Merkmal", "Quelle", "Berechnung", "Typ"],
              merkmale_rows,
              col_widths=[3.5, 3.5, 5.0, 4.0])
    doc.add_paragraph()

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 4. ERGEBNISSE
    # ══════════════════════════════════════════════════════════════
    add_heading(doc, "4  Ergebnisse", level=1)

    add_heading(doc, "4.1  Ziel-Variable — Statistiken", level=2)
    add_body_bold(doc, [("Datenbasis nach Aufbereitung: ", False), (f"{n_haeuser:,} Krankenhäuser", True)])
    stat_rows = [
        ("Anzahl Krankenhäuser",              f"{n_haeuser:,}"),
        ("Ø Indikatoren pro Haus",            f"{(total_qi_sum / n_haeuser):.1f}"),
        ("Median auffällig-Quote",            f"{median_quote:.2%}"),
        ("Mittelwert auffällig-Quote",        f"{stats['mean']:.2%}"),
        ("Min. auffällig-Quote",              f"{stats['min']:.2%}"),
        ("Max. auffällig-Quote",              f"{stats['max']:.2%}"),
        ("Häuser mit vielen Problemen (=1)",  f"{n_auffaellig:,} ({n_auffaellig / n_haeuser:.1%})"),
        ("Häuser ohne viele Probleme (=0)",   f"{n_haeuser - n_auffaellig:,} ({(n_haeuser - n_auffaellig) / n_haeuser:.1%})"),
    ]
    add_table(doc, ["Kennzahl", "Wert"], stat_rows, col_widths=[8.0, 4.0])
    doc.add_paragraph()
    add_body_bold(doc, [
        ("Die Ziel-Variable ist nahezu ", False),
        ("ausgewogen verteilt (ca. 49 % vs. 51 %)", True),
        (", was für Machine-Learning-Modelle optimal ist.", False),
    ])

    add_heading(doc, "4.2  Analysetabelle", level=2)
    add_body(doc,
        f"Die finale Analysetabelle (analysetabelle.csv) enthält {analyse.shape[0]:,} Zeilen "
        f"und {analyse.shape[1]} Spalten. Jede Zeile repräsentiert genau ein Krankenhaus.")
    col_rows = [(c, str(analyse[c].dtype)) for c in analyse.columns]
    add_table(doc, ["Spalte", "Datentyp"], col_rows, col_widths=[8.0, 4.0])
    doc.add_paragraph()

    add_heading(doc, "4.3  Wozu wird die Analysetabelle genutzt?", level=2)
    add_body(doc,
        "Die Analysetabelle ist die einzige Datengrundlage für alle weiteren Projektschritte. "
        "Das Prinzip: Rohdaten → Analysetabelle → alles andere.")
    nutzung_rows = [
        ("Baustein 2\nDeskriptive Analyse",
         "Grafiken (Box-Plots, Scatter-Plots, Balkendiagramme) werden direkt aus der Tabelle erzeugt. "
         "Beispiel: \"Unterscheiden sich Bettenzahl zwischen Häusern mit/ohne viele Probleme?\""),
        ("Baustein 3\nDashboard Seite 1",
         "hat_viele_Probleme + Koordinaten → Deutschland-Karte mit regionaler Verteilung"),
        ("Baustein 3\nDashboard Seite 2",
         "Merkmale gruppiert nach hat_viele_Probleme → Vergleichsdiagramme mit Dropdown"),
        ("Baustein 3\nDashboard Seite 3",
         "Filter nach Betten / Region / Träger → ähnliche Häuser finden und deren Qualität zeigen"),
        ("Baustein 4\nDecision Tree",
         "Feature Matrix X = SO.Betten, SO.Uni, fortbildungsquote, aerzte_pro_bett, pflege_pro_bett, "
         "ist_konzern, traeger_enc; Zielvariable y = hat_viele_Probleme; direkt für train_test_split "
         "und DecisionTreeClassifier nutzbar. Ergebnis: Accuracy 63,6 % (Basislinie 50,7 %), R²=0,033"),
    ]
    add_table(doc, ["Baustein", "Nutzung der Analysetabelle"], nutzung_rows,
              col_widths=[4.0, 12.0])
    doc.add_paragraph()

    add_heading(doc, "4.4  Fehlende Werte", level=2)
    missing = analyse.isnull().sum()
    missing_rows = [(col, int(missing[col]), f"{missing[col] / len(analyse):.1%}")
                    for col in analyse.columns if missing[col] > 0]
    if missing_rows:
        add_table(doc, ["Spalte", "Fehlende Werte (n)", "Anteil"], missing_rows,
                  col_widths=[6.0, 4.0, 4.0])
        doc.add_paragraph()
        add_body(doc,
            "KH.Träger.Art (28 fehlend) und fortbildungsquote (33 fehlend) haben sehr "
            "geringe Fehlquoten (<2%) und können für die Analyse entweder mit dem Modus "
            "imputiert oder als eigene Kategorie 'unbekannt' behandelt werden.")
    else:
        add_body(doc, "Keine fehlenden Werte.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 5. DESKRIPTIVE ANALYSE — BEFUNDE
    # ══════════════════════════════════════════════════════════════
    add_heading(doc, "5  Deskriptive Analyse — Befunde", level=1)
    add_body(doc,
        "Die Analyse wurde in 02_Analyse.ipynb durchgeführt. Grundlage: Data/analysetabelle.csv. "
        "12 Grafiken wurden erstellt, jede mit automatisch berechnetem Befundsatz. "
        "Farbschema einheitlich: grün = wenige Probleme, rot = viele Probleme.")

    add_heading(doc, "5.1  Übersicht der Befunde", level=2)
    df_analyse = analyse.copy()
    m_betten_0 = int(df_analyse[df_analyse["hat_viele_Probleme"] == 0]["SO.Betten"].median())
    m_betten_1 = int(df_analyse[df_analyse["hat_viele_Probleme"] == 1]["SO.Betten"].median())
    pct_privat   = df_analyse[df_analyse["KH.Träger.Art"] == "privat"]["hat_viele_Probleme"].mean()
    pct_frei     = df_analyse[df_analyse["KH.Träger.Art"] == "freigemeinnützig"]["hat_viele_Probleme"].mean()
    pct_oeffentl = df_analyse[df_analyse["KH.Träger.Art"] == "öffentlich"]["hat_viele_Probleme"].mean()
    pct_uni      = df_analyse[df_analyse["SO.Uni"] == 1]["hat_viele_Probleme"].mean()
    pct_normal   = df_analyse[df_analyse["SO.Uni"] == 0]["hat_viele_Probleme"].mean()
    fb_0 = df_analyse[df_analyse["hat_viele_Probleme"] == 0]["fortbildungsquote"].median()
    fb_1 = df_analyse[df_analyse["hat_viele_Probleme"] == 1]["fortbildungsquote"].median()
    ab_0 = df_analyse[df_analyse["hat_viele_Probleme"] == 0]["aerzte_pro_bett"].median()
    ab_1 = df_analyse[df_analyse["hat_viele_Probleme"] == 1]["aerzte_pro_bett"].median()
    pb_0 = df_analyse[df_analyse["hat_viele_Probleme"] == 0]["pflege_pro_bett"].median()
    pb_1 = df_analyse[df_analyse["hat_viele_Probleme"] == 1]["pflege_pro_bett"].median()
    n_konzern = int(df_analyse["ist_konzern"].sum())
    quote_konzern = df_analyse[df_analyse["ist_konzern"] == 1]["hat_viele_Probleme"].mean()
    quote_unabh   = df_analyse[df_analyse["ist_konzern"] == 0]["hat_viele_Probleme"].mean()

    befund_rows = [
        ("Grafik 1\nauffällig-Quote",
         f"Median {df_analyse['auffaellig_quote'].median():.0%}, linkssteil verteilt — "
         f"die meisten Häuser liegen zwischen 60–90%."),
        ("Grafik 2\nBettenzahl",
         f"Median: Wenige Probleme = {m_betten_0} Betten, Viele Probleme = {m_betten_1} Betten. "
         f"Kein klarer Größenunterschied."),
        ("Grafik 3\nTrägerschaft",
         f"Privat: {pct_privat:.1%} | Freigemeinnützig: {pct_frei:.1%} | Öffentlich: {pct_oeffentl:.1%} "
         f"haben viele Probleme. Private Häuser auffällig höher."),
        ("Grafik 4\nUni-Kliniken",
         f"Uni-Kliniken: {pct_uni:.1%} vs. normale Häuser: {pct_normal:.1%} — kaum Unterschied."),
        ("Grafik 5+6\nFortbildung & Ärzte/Bett",
         f"Fortbildungsquote: kein Unterschied (Md={fb_0:.3f} vs. {fb_1:.3f}). "
         f"Ärzte/Bett: Wenige={ab_0:.3f}, Viele={ab_1:.3f} — leichter Unterschied sichtbar."),
        ("Grafik 7\nBundesland",
         "Saarland höchster Anteil (63,2%), Berlin niedrigster (33,3%). "
         "Regionale Unterschiede sichtbar — kleine n beachten."),
        ("Grafik 8\nKorrelationsmatrix",
         "Stärkste Korrelation mit Ziel-Variable: total_qi (r=−0,28), "
         "dann aerzte_pro_bett (r=−0,14). Fortbildungsquote praktisch keine Korrelation."),
        ("Grafik 9\nScatter Betten/Ärzte",
         "Kein klares Trennmuster zwischen den Gruppen — Überlappung stark."),
        ("Grafik 10\nStörfaktor Träger×Betten",
         f"Private Häuser sind kleiner (Md={int(df_analyse[df_analyse['KH.Träger.Art'] == 'privat']['SO.Betten'].median())} Betten). "
         "Der Trägereffekt muss daher mit Vorsicht interpretiert werden."),
        ("Grafik 11\nPflegekräfte pro Bett *(2026-07-29)*",
         f"Wenige={pb_0:.3f}, Viele={pb_1:.3f} — ähnliches Muster wie Ärzte/Bett (T-Test p<0,001, signifikant)."),
        ("Grafik 12\nKonzernvergleich *(2026-07-29)*",
         f"Konzernhäuser: {n_konzern:,} von {len(df_analyse):,} ({n_konzern/len(df_analyse):.1%}). "
         f"Anteil viele Probleme: Konzern={quote_konzern:.1%} vs. unabhängig={quote_unabh:.1%} — "
         "praktisch kein Unterschied (Chi² p=0,90, nicht signifikant)."),
    ]
    add_table(doc, ["Grafik", "Befund"], befund_rows, col_widths=[3.5, 12.5])
    doc.add_paragraph()

    add_heading(doc, "5.2  Grafiken", level=2)
    grafik_dateien = [
        ("grafiken/g1_auffaellig_quote.png",       "Grafik 1: Verteilung der auffällig-Quote"),
        ("grafiken/g2_bettenzahl.png",              "Grafik 2: Bettenzahl"),
        ("grafiken/g3_traegerschaft.png",           "Grafik 3: Trägerschaft"),
        ("grafiken/g4_uni.png",                     "Grafik 4: Uni-Kliniken vs. normale Häuser"),
        ("grafiken/g5_6_fortbildung_aerzte.png",    "Grafik 5+6: Fortbildungsquote & Ärzte pro Bett"),
        ("grafiken/g7_bundesland.png",              "Grafik 7: Anteil Häuser mit vielen Problemen je Bundesland"),
        ("grafiken/g8_korrelation.png",             "Grafik 8: Korrelationsmatrix"),
        ("grafiken/g9_scatter_betten_aerzte.png",   "Grafik 9: Scatter — Bettenzahl vs. Ärzte pro Bett"),
        ("grafiken/g10_stoerfaktor_traeger.png",    "Grafik 10: Störfaktor — Bettengröße je Trägerschaft"),
        ("grafiken/g11_pflege_pro_bett.png",        "Grafik 11: Pflegekräfte pro Bett — MIT vs. OHNE viele Probleme"),
        ("grafiken/g12_konzern_vergleich.png",      "Grafik 12: Konzernhaus vs. unabhängiges Haus"),
    ]
    for pfad, titel in grafik_dateien:
        if os.path.exists(pfad):
            p = doc.add_paragraph()
            run = p.add_run(titel)
            run.bold = True
            run.font.size = Pt(11)
            doc.add_picture(pfad, width=Cm(15))
            doc.add_paragraph()
        else:
            add_body(doc, f"[Grafik nicht gefunden: {pfad}]")

    add_heading(doc, "5.3  Gesamteinschätzung", level=2)
    add_body(doc,
        "Die Analyse zeigt keine starken, eindeutigen Zusammenhänge zwischen den untersuchten "
        "Strukturmerkmalen und der Ziel-Variable. Der stärkste Prädiktor ist total_qi "
        "(Anzahl bewerteter Indikatoren) — ein strukturelles Merkmal des Hauses, kein "
        "Qualitätsmerkmal: Häuser mit mehr bewerteten Indikatoren haben tendenziell niedrigere "
        "Auffälligkeitsquoten.")
    add_body(doc,
        "Einziger klarer inhaltlicher Befund: Private Häuser haben mit 56,5% einen höheren "
        "Anteil als freigemeinnützige (46,4%) und öffentliche (46,7%) Träger. Allerdings sind "
        "private Häuser im Median deutlich kleiner — der Trägereffekt könnte durch "
        "Größenunterschiede mitverursacht sein (Störfaktor).")
    add_body(doc,
        "Wichtig: Das ist ein valides Ergebnis. Schwache Zusammenhänge sind in echten "
        "Gesundheitsdaten normal, da viele weitere Faktoren (Patientenmix, Spezialisierung, "
        "Dokumentationsqualität) eine Rolle spielen, die nicht im Datensatz enthalten sind.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 6. PROJEKTSTAND & NÄCHSTE SCHRITTE
    # ══════════════════════════════════════════════════════════════
    add_heading(doc, "6  Projektstand & Nächste Schritte", level=1)

    add_heading(doc, "6.1  Abgeschlossene Bausteine", level=2)
    abgeschlossen = [
        ("Baustein 1", "Daten vorbereiten",
         "analysetabelle.csv — 1.824 Zeilen, 18 Spalten. End-to-end getestet 2026-07-29."),
        ("Baustein 2", "Deskriptive Analyse",
         "12 Grafiken, T-Test, ANOVA, Chi²-Test, Korrelationsmatrix. 02_Analyse.ipynb."),
        ("Baustein 3", "Streamlit-Dashboard",
         "4 Seiten: Gesamtüberblick, Einflussfaktoren, Häuser vergleichen, Qualitäts-Vorhersage. "
         "Deployed auf Streamlit Community Cloud via GitHub."),
        ("Baustein 4", "Decision Tree",
         "Accuracy 63,6 % (Basislinie 50,7 %), R²=0,033, Feature Importance: "
         "aerzte_pro_bett 53,6 %, pflege_pro_bett 23,8 %, SO.Betten 22,6 %."),
        ("Baustein 5", "Abschluss & Präsentation",
         "Folie 13 (PPTX) auf Streamlit aktualisiert. Streamlit-Präsentationsfolie erstellt "
         "(folie13_praesentation.py). Vollständiges Sprechertext-Dokument erstellt "
         "(Praesentationsskript_Qualitaets_Muster_Finder.docx, 30 Min, 15 Folien, Wir-Perspektive). "
         "Praesentation_Folien_Beschreibung.md mit PPTX abgeglichen."),
    ]
    add_table(doc, ["Baustein", "Titel", "Ergebnis"], abgeschlossen,
              col_widths=[2.5, 3.5, 10.0])
    doc.add_paragraph()

    add_heading(doc, "6.2  Noch offen", level=2)
    offen = [
        ("Generalprobe mit Stoppuhr",
         "Präsentation einmal komplett durchlaufen — Ziel: 30 Min ± 2 Min"),
        ("Live-Demo Dashboard",
         "Dashboard auf Streamlit Cloud vor Präsentation testen (URL prüfen, Ladezeit)"),
        ("scripts/powerbi_anleitung.py",
         "Hartkodierter os.chdir()-Pfad eines fremden Rechners noch nicht repariert"),
    ]
    add_table(doc, ["Aufgabe", "Details"], offen, col_widths=[5.5, 10.5])
    doc.add_paragraph()

    add_heading(doc, "6.3  Mögliche Erweiterungen", level=2)
    erweiterungen = [
        ("Dokumentationsrate",
         "QS.Leistungsbereich.csv → QSLB.Dokumentationsrate als weiteres Merkmal einbinden"),
        ("NLP / Sentiment-Analyse",
         "QSQI.KommentarKrankenhaus — individuelle Freitextkommentare der Häuser, NLP-fähig"),
        ("Mehrstufige Regression",
         "Patientenstruktur als Kontrollvariable → echten Träger-/Personaleffekt isolieren"),
    ]
    add_table(doc, ["Ansatz", "Beschreibung"], erweiterungen, col_widths=[4.5, 11.5])
    doc.add_paragraph()

    add_body_bold(doc, [
        ("Stand 2026-08-10: ", True),
        ("Alle 5 Bausteine abgeschlossen. Bausteine 1–4 end-to-end getestet (2026-07-29). "
         "Baustein 5 (Präsentation): Unterlagen vollständig erstellt, Generalprobe steht noch aus.",
         False),
    ])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 7. FRAGESTELLUNG & ANTWORTEN
    # ══════════════════════════════════════════════════════════════
    add_heading(doc, "7  Fragestellung & Antworten", level=1)
    add_body(doc,
        "Dieses Kapitel gegenüberstellt die originale Aufgabenstellung "
        "(aus Aufgabenstellung/Fragestellung.docx) und die konkreten Ergebnisse "
        "aus unserer Projektarbeit.")

    # ── 7.1 Zentrale Frage ────────────────────────────────────────
    add_heading(doc, "7.1  Die zentrale Fragestellung", level=2)
    add_body_bold(doc, [("Frage: ", True),
        ("Welche Krankenhausmerkmale hängen damit zusammen, dass ein Haus "
         "überdurchschnittlich viele Qualitätsprobleme hat?", True)])
    add_body_bold(doc, [("Antwort: ", True),
        ("Kein einzelnes Merkmal ist ein starker Prädiktor. "
         "Der stärkste inhaltliche Zusammenhang liegt bei ", False),
        ("Ärzte pro Bett (r = −0,14, T-Test p < 0,001)", True),
        (" und ", False),
        ("Pflegekräfte pro Bett (r = −0,14, p < 0,001)", True),
        (": Häuser mit mehr Personal pro Bett haben tendenziell weniger Qualitätsprobleme. "
         "Trägerschaft zeigt einen sichtbaren Unterschied (privat 56,5 % vs. öffentlich 46,7 %), "
         "ist aber durch die kleinere durchschnittliche Bettenanzahl privater Häuser beeinflusst. "
         "Fortbildungsquote, Uni-Status und Konzernzugehörigkeit zeigen keinen messbaren Zusammenhang. "
         "Das Strukturmodell erklärt insgesamt nur 3,3 % der Varianz (R² = 0,033) — "
         "ein valides Ergebnis, das zeigt, wo die Grenzen strukturbasierter Erklärungsansätze liegen.", False)])

    doc.add_paragraph()

    # ── 7.2 Aufgabe 1: Daten vorbereiten ─────────────────────────
    add_heading(doc, "7.2  Aufgabe 1: Daten vorbereiten", level=2)
    aufgabe1 = [
        ("Ziel-Variable erstellen:\nAnteil auffälliger QI pro Krankenhaus berechnen",
         "✅ Erledigt. auffaellig_quote = auffaellig_n / total_qi, berechnet aus "
         "QS.Qualitätsindikator.csv. Fallstricke behoben: N99-Zeilen ausgeschlossen "
         "(nicht bewertet ≠ unauffällig), Zählkennzahlen (EKez/TKez) gefiltert, "
         "Duplikate per drop_duplicates(['SO.QBID', 'QSQI.Indikator']) entfernt. "
         "Datei: 01_Exploration.ipynb Kap. 3."),
        ("Target-Variable:\nHat ueberdurchschnittlich viele Probleme = ueber Median",
         "✅ Erledigt. Median der auffaellig_quote = 76,92 %. "
         "Spalte hat_viele_Probleme: 899 Häuser = 1, 925 Häuser = 0. "
         "Nahezu ausgewogen (49 % vs. 51 %) — ideal für ML."),
        ("8 Merkmale auswählen & zusammenführen",
         "✅ Erledigt. Merkmale: SO.Betten, KH.Träger.Art, SO.Bundesland, SO.Uni, "
         "fortbildungsquote, aerzte_pro_bett, pflege_pro_bett, ist_konzern. "
         "Ergebnis: analysetabelle.csv — 1.824 Zeilen, 18 Spalten. "
         "Technisch aufwändig: aerzte_pro_bett über 2 Left Joins (FA.Personalliste + FA.csv), "
         "Komma-Dezimal-Bug behoben."),
    ]
    add_table(doc, ["Teilaufgabe", "Ergebnis aus dem Projekt"], aufgabe1,
              col_widths=[5.0, 11.0])
    doc.add_paragraph()

    # ── 7.3 Aufgabe 2: Deskriptive Analyse ───────────────────────
    add_heading(doc, "7.3  Aufgabe 2: Deskriptive Analyse", level=2)
    aufgabe2 = [
        ("Wie unterscheiden sich Häuser MIT vs. OHNE viele Probleme?",
         "Ärzte/Bett: Median 0,468 (wenige) vs. 0,390 (viele) — T-Test t=6,002, p<0,001. "
         "Pflegekräfte/Bett: ähnliches Muster, ebenfalls signifikant. "
         "Trägerschaft: privat 56,5 % vs. freigemeinnützig 46,4 % / öffentlich 46,7 %. "
         "Fortbildungsquote: kein Unterschied (r ≈ 0,01)."),
        ("Korrelationen berechnen: Welche Merkmale hängen zusammen?",
         "Korrelationsmatrix berechnet (Pearson r). Stärkste Korrelation mit hat_viele_Probleme: "
         "total_qi r=−0,28 (Artefakt), aerzte_pro_bett r=−0,14, pflege_pro_bett r=−0,14, "
         "SO.Betten r=−0,08, fortbildungsquote r≈0,01, ist_konzern r≈0,00."),
        ("Gruppenvergleiche: Uni-Kliniken vs. normal, groß vs. klein, öffentlich vs. privat",
         "Uni vs. normal: 47 % vs. 49 % — kein Unterschied. "
         "Groß vs. klein: Bettenzahl zeigt schwachen Zusammenhang (r=−0,08). "
         "Öffentlich vs. privat: sichtbarer Unterschied (46,7 % vs. 56,5 %), "
         "aber private Häuser sind im Median kleiner — Störfaktor!"),
        ("Visualisierungen: Box-Plots, Scatter-Plots, Balkendiagramme",
         "12 Grafiken erstellt und in grafiken/ gespeichert. "
         "Alle reproduzierbar über scripts/Grafiken_Speichern.py und 02_Analyse.ipynb. "
         "Farbschema einheitlich: grün = wenige Probleme, rot = viele Probleme."),
    ]
    add_table(doc, ["Teilaufgabe", "Ergebnis aus dem Projekt"], aufgabe2,
              col_widths=[5.0, 11.0])
    doc.add_paragraph()

    # ── 7.4 Aufgabe 3: Dashboard ──────────────────────────────────
    add_heading(doc, "7.4  Aufgabe 3: Dashboard bauen", level=2)
    aufgabe3 = [
        ("Seite 1 — Gesamtüberblick:\nKennzahlen, Karte, Verteilung",
         "✅ Erledigt. KPI-Tabelle (n Häuser, Anteil viele Probleme, Ø-Quote, Ø-Ärzte/Bett). "
         "Interaktive Deutschlandkarte (Plotly, grün/rot). Histogramm der auffaellig_quote. "
         "Filter: Bundesland, Träger, Klinik-Typ."),
        ("Seite 2 — Einflussfaktoren:\nDropdown → Verteilung je Merkmal",
         "✅ Erledigt. 4 Tabs: Trägervergleich, Personal/Bett Boxplot, Streudiagramm, "
         "Pivot-Tabelle. Direkte visuelle Antwort auf die Projektfrage."),
        ("Seite 3 — Häuser vergleichen:\nÄhnliche Häuser finden + Steckbrief",
         "✅ Erledigt. Filter nach Bundesland, Träger, Bettengröße. "
         "Ergebnistabelle mit auffaellig_quote je Haus. "
         "Einzelhaus-Steckbrief: Haus vs. Ø ähnlicher Häuser."),
        ("Seite 4 — Qualitäts-Vorhersage (Bonus):\nDecision Tree Risiko-Rechner",
         "✅ Erledigt. Merkmal-Eingaben → scikit-learn Decision Tree gibt Risikoeinschätzung. "
         "Baumvisualisierung + Feature Importance angezeigt. "
         "Deployment: Streamlit Community Cloud via GitHub."),
    ]
    add_table(doc, ["Teilaufgabe", "Ergebnis aus dem Projekt"], aufgabe3,
              col_widths=[5.0, 11.0])
    doc.add_paragraph()

    # ── 7.5 Bonus: Entscheidungsbaum ─────────────────────────────
    add_heading(doc, "7.5  Bonus: Einfacher Entscheidungsbaum", level=2)
    aufgabe4 = [
        ("Decision Tree trainieren (max_depth=3)",
         "✅ Erledigt. DecisionTreeClassifier, stratifizierter 80/20-Split, 5-Fold CV. "
         "Klasse KrankenhausModell in model/modell_klasse.py. "
         "Modell gespeichert als Data/modell_krankenhaus.pkl."),
        ("Vorhersage: Hat das Haus überdurchschnittlich viele Probleme?",
         f"✅ Accuracy: 63,6 % (Basislinie 50,7 %). "
         "Besser als Zufall, aber kein starkes Modell — R² = 0,033. "
         "Erste Frage des Baums: aerzte_pro_bett < 0,271? "
         "Feature Importance: aerzte_pro_bett 53,6 %, pflege_pro_bett 23,8 %, SO.Betten 22,6 %."),
        ("Metriken & Evaluation",
         "Accuracy, Precision, Recall, F1-Score, Confusion Matrix berechnet. "
         "5-Fold Cross-Validation: 59,7 % ± 4,2 % — kein Overfitting. "
         "R² = 0,033: Strukturmerkmale erklären nur 3,3 % der Varianz."),
    ]
    add_table(doc, ["Teilaufgabe", "Ergebnis aus dem Projekt"], aufgabe4,
              col_widths=[5.0, 11.0])
    doc.add_paragraph()

    # ── 7.6 Gesamtantwort ─────────────────────────────────────────
    add_heading(doc, "7.6  Gesamtfazit zur Fragestellung", level=2)
    fazit_rows = [
        ("Ärzte pro Bett",        "r = −0,14", "T-Test p < 0,001", "🟢 Schwach, aber stärkster inhaltlicher Prädiktor"),
        ("Pflegekräfte pro Bett", "r = −0,14", "T-Test p < 0,001", "🟢 Gleichstarker Effekt wie Ärzte/Bett"),
        ("Trägerschaft",          "privat +10 PP", "ANOVA p < 0,001", "🟡 Sichtbar — Störfaktor Hausgröße beachten"),
        ("Bundesland",            "variabel",   "—",                "🟡 Sichtbar — kleine Stichproben vorsichtig"),
        ("Bettenzahl",            "r = −0,08",  "schwach",          "🟡 Sehr schwacher Zusammenhang"),
        ("Uni-Status",            "47 % vs. 49 %", "kein Unterschied", "🔴 Kein Zusammenhang"),
        ("Konzernzugehörigkeit",  "r ≈ 0,00",   "Chi² p = 0,90",    "🔴 Kein Zusammenhang"),
        ("Fortbildungsquote",     "r ≈ 0,01",   "kein Unterschied", "🔴 Kein Zusammenhang"),
    ]
    add_table(doc,
              ["Merkmal", "Korrelation / Effekt", "Statistik", "Bewertung"],
              fazit_rows,
              col_widths=[4.0, 3.5, 3.5, 5.0])
    doc.add_paragraph()
    add_body_bold(doc, [
        ("Gesamtantwort: ", True),
        ("Die Strukturmerkmale eines Krankenhauses erklären seine Qualitätsprobleme nur sehr begrenzt. "
         "Personalintensität (Ärzte und Pflegekräfte pro Bett) ist der stärkste, aber immer noch "
         "schwache Faktor. Andere Einflüsse — Patientenmix, Dokumentationsverhalten, regionale "
         "Besonderheiten — spielen wahrscheinlich eine größere Rolle. "
         "Kein Zusammenhang ist ein valides wissenschaftliches Ergebnis.", False),
    ])

    return doc



# ══════════════════════════════════════════════════════════════════
# MAIN
# ══════════════════════════════════════════════════════════════════
if __name__ == "__main__":
    print("Lade Daten aus analysetabelle.csv ...")
    daten = lade_daten()
    print(f"  Häuser: {daten['n_haeuser']:,}  |  Median-Quote: {daten['median_quote']:.2%}")

    print("Erzeuge Word-Dokument ...")
    doc = erzeuge_dokument(daten)

    doc.save(str(OUT_PATH))
    print(f"✅ Dokument gespeichert: {OUT_PATH}")