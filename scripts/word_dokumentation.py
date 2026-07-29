"""
word_dokumentation.py
=====================
Erzeugt die Word-Projektdokumentation "Dokumentation_Qualitaets_Muster_Finder.docx"
aus den Daten der analysetabelle.csv.

Ursprung:
    Teil "Word-Dokumentation generieren" aus 01_Exploration.ipynb —
    in diese eigenstaendige Skriptdatei ausgelagert.

Aufruf:
    python scripts/word_dokumentation.py
    (oder aus dem Projekt-Root:  python scripts/word_dokumentation.py)

Abhaengigkeiten:
    pandas, python-docx  (python-docx wird bei Bedarf automatisch installiert)
"""

import os
import subprocess
import sys
from datetime import date
from pathlib import Path

import pandas as pd

# ── python-docx sicherstellen ─────────────────────────────────────
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

# ── Projekt-Root (eine Ebene ueber /scripts) ──────────────────────
PROJEKT_ROOT = Path(__file__).resolve().parent.parent
os.chdir(PROJEKT_ROOT)

DATA_PATH  = PROJEKT_ROOT / "Data" / "analysetabelle.csv"
OUT_PATH   = PROJEKT_ROOT / "Dokumentation_Qualitaets_Muster_Finder.docx"


# ══════════════════════════════════════════════════════════════════
# DATEN LADEN & KENNGROESSEN BERECHNEN
# ══════════════════════════════════════════════════════════════════
def lade_daten() -> dict:
    """
    Laedt analysetabelle.csv und berechnet alle Kenngroessen,
    die fuer die Word-Dokumentation benoetigt werden.

    Returns:
        dict mit den Schluesseln:
            analyse, auffaellig_quote, median_quote, n_haeuser,
            n_auffaellig, stats, total_qi_sum
    """
    analyse = pd.read_csv(DATA_PATH, low_memory=False)

    # auffaellig_quote als eigenstaendiger DataFrame (wie im Notebook)
    quote_cols = ["SO.QBID", "total_qi", "auffaellig_n",
                  "auffaellig_quote", "hat_viele_Probleme"]
    auffaellig_quote = analyse[[c for c in quote_cols if c in analyse.columns]].copy()

    median_quote   = auffaellig_quote["auffaellig_quote"].median()
    n_haeuser      = len(auffaellig_quote)
    n_auffaellig   = int(auffaellig_quote["hat_viele_Probleme"].sum())
    stats          = auffaellig_quote["auffaellig_quote"].describe()
    total_qi_sum   = int(analyse["total_qi"].sum())  # = len(qi_dedup) im Notebook

    return {
        "analyse":          analyse,
        "auffaellig_quote": auffaellig_quote,
        "median_quote":     median_quote,
        "n_haeuser":        n_haeuser,
        "n_auffaellig":     n_auffaellig,
        "stats":            stats,
        "total_qi_sum":     total_qi_sum,
    }


# ══════════════════════════════════════════════════════════════════
# HILFSFUNKTIONEN FUER DAS WORD-DOKUMENT
# ══════════════════════════════════════════════════════════════════
def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.style.font.size = Pt(11)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    p.style.font.size = Pt(11)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    # Header-Zeile
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        run = hdr[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)
        tc = hdr[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "4472C4")
        shd.set(qn("w:color"), "FFFFFF")
        shd.set(qn("w:val"), "clear")
        tcPr.append(shd)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # Daten-Zeilen
    for r_i, row in enumerate(rows):
        cells = table.rows[r_i + 1].cells
        for c_i, val in enumerate(row):
            cells[c_i].text = str(val)
            cells[c_i].paragraphs[0].runs[0].font.size = Pt(10)
    # Spaltenbreiten
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return table


# ══════════════════════════════════════════════════════════════════
# DOKUMENT ERZEUGEN
# ══════════════════════════════════════════════════════════════════
def erzeuge_dokument(daten: dict) -> Document:
    """Baut das komplette Word-Dokument auf und gibt es zurueck."""
    analyse          = daten["analyse"]
    auffaellig_quote = daten["auffaellig_quote"]
    median_quote     = daten["median_quote"]
    n_haeuser        = daten["n_haeuser"]
    n_auffaellig     = daten["n_auffaellig"]
    stats            = daten["stats"]
    total_qi_sum     = daten["total_qi_sum"]

    doc = Document()

    # ── Seitenraender ──────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # ══════════════════════════════════════════════════════════════
    # TITELSEITE
    # ══════════════════════════════════════════════════════════════
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Qualitäts-Muster-Finder")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("Projektdokumentation — Stand der Datenaufbereitung").font.size = Pt(14)

    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run(
        f"Erstellt: {date.today().strftime('%d.%m.%Y')}    |    "
        "Datenbasis: Qualitätsberichte deutscher Krankenhäuser 2023"
    ).font.size = Pt(11)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # INHALTSVERZEICHNIS (manuell)
    # ══════════════════════════════════════════════════════════════
    add_heading(doc, "Inhaltsverzeichnis", level=1)
    toc_entries = [
        ("1",   "Projektübersicht",                         "3"),
        ("1.1", "Fragestellung",                            "3"),
        ("1.2", "Projektziel & Bausteine",                  "3"),
        ("2",   "Datensatz",                                "4"),
        ("2.1", "Überblick",                                "4"),
        ("2.2", "Schlüssel-ID: SO.QBID",                    "4"),
        ("2.3", "Relevante Tabellen",                       "4"),
        ("2.4", "Nicht relevante Tabellen",                 "5"),
        ("3",   "Datenaufbereitung",                        "6"),
        ("3.1", "Vorgehen & Kriterien",                     "6"),
        ("3.2", "Ziel-Variable",                            "6"),
        ("3.3", "Merkmale (Features)",                      "7"),
        ("4",   "Ergebnisse",                               "8"),
        ("4.1", "Ziel-Variable — Statistiken",              "8"),
        ("4.2", "Analysetabelle",                           "8"),
        ("4.3", "Wozu wird die Analysetabelle genutzt?",    "9"),
        ("4.4", "Fehlende Werte",                           "10"),
        ("5",   "Deskriptive Analyse — Befunde",            "11"),
        ("5.1", "Übersicht der Befunde",                    "11"),
        ("5.2", "Gesamteinschätzung",                       "12"),
        ("6",   "Offene Punkte & Nächste Schritte",         "13"),
    ]
    for nr, title_text, page in toc_entries:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5 * (title_text.count(".") + 0))
        run_nr = p.add_run(f"{nr}  {title_text}")
        run_nr.font.size = Pt(11)
        if "." not in nr:
            run_nr.bold = True

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 1. PROJEKTÜBERSICHT
    # ══════════════════════════════════════════════════════════════
    add_heading(doc, "1  Projektübersicht", level=1)

    add_heading(doc, "1.1  Fragestellung", level=2)
    add_body(doc,
        "Welche Krankenhausmerkmale hängen damit zusammen, dass ein Haus "
        "überdurchschnittlich viele Qualitätsprobleme aufweist?")
    add_body(doc,
        "Grundlage sind die jährlichen Qualitätsberichte aller deutschen Krankenhäuser. "
        "Jedes Haus berichtet über ~150 Qualitätsindikatoren. Bei manchen Indikatoren "
        "werden Häuser als 'rechnerisch auffällig' bewertet. Ziel ist es, strukturelle "
        "Merkmale (Größe, Personal, Träger, Region) zu identifizieren, die mit einer "
        "erhöhten Auffälligkeitsquote zusammenhängen.")
    add_body(doc, "Wichtiger Hinweis: Kein Zusammenhang ist ein valides Ergebnis.")
    add_body(doc,
        "Hintergrund: Die Qualitätsindikatoren markieren Häuser als 'rechnerisch auffällig', "
        "wenn ihr Wert außerhalb eines Referenzbereichs liegt. Das ist aber nur ein statistisches Signal "
        "— kein Qualitätsurteil. Ob wirklich ein Problem dahintersteckt, klärt ein separates Prüfverfahren "
        "(Strukturierter Dialog). Außerdem gibt es viele Einflussfaktoren (Patientenmix, Größe, Spezialisierung), "
        "die wir nicht alle kontrollieren können. Wenn die Analyse also keinen klaren Zusammenhang "
        "zwischen Strukturmerkmalen und Auffälligkeit zeigt, ist das ein ehrliches und wissenschaftlich "
        "korrektes Ergebnis — kein Scheitern. Quelle: Aufgabenstellung/Text_Presentation.docx, Folie 7.")

    add_heading(doc, "1.2  Projektziel & Bausteine", level=2)
    add_body(doc, "Das Projekt ist in fünf Bausteine gegliedert:")
    bausteine = [
        ("Baustein 1", "Daten vorbereiten", "✅ Abgeschlossen"),
        ("Baustein 2", "Deskriptive Analyse", "✅ Abgeschlossen"),
        ("Baustein 3", "Streamlit-Dashboard (3 Seiten)", "⬜ Offen"),
        ("Baustein 4", "Entscheidungsbaum (Bonus)", "⬜ Offen"),
        ("Baustein 5", "Abschluss & Präsentation", "⬜ Offen"),
    ]
    add_table(doc, ["Baustein", "Beschreibung", "Status"], bausteine,
              col_widths=[3.5, 8.5, 3.0])
    doc.add_paragraph()

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 2. DATENSATZ
    # ══════════════════════════════════════════════════════════════
    add_heading(doc, "2  Datensatz", level=1)

    add_heading(doc, "2.1  Überblick", level=2)
    add_body(doc,
        "Der Datensatz besteht aus 86 CSV-Dateien im Data/-Ordner. "
        "Alle Daten stammen aus den offiziellen Qualitätsberichten deutscher "
        "Krankenhäuser (Berichtsjahr 2023) und werden vom IQTIG veröffentlicht.")
    add_bullet(doc, "~1.900 Krankenhäuser (Standorte)")
    add_bullet(doc, "~150 Qualitätsindikatoren pro Haus")
    add_bullet(doc, "Strukturdaten: Betten, Personal, Träger, Standort, Geo-Koordinaten")

    add_heading(doc, "2.2  Schlüssel-ID: SO.QBID", level=2)
    add_body(doc,
        "Fast alle Tabellen sind über die Spalte SO.QBID miteinander verknüpft. "
        "SO.QBID ist die eindeutige ID eines Krankenhaus-Standorts und dient als "
        "primärer Join-Key für alle Merge-Operationen.")

    add_heading(doc, "2.3  Relevante Tabellen", level=2)
    rel_rows = [
        ("SO.csv",                    "Stammdaten aller Krankenhäuser (Haupttabelle)",
         "SO.QBID, SO.Betten, SO.Bundesland, SO.Uni, KH.Träger.Art, Koordinaten",
         "Kern-Merkmale"),
        ("QS.Qualitätsindikator.csv", "Qualitätsindikatoren mit Bewertungen (>50 MB)",
         "SO.QBID, QSErgBewStrukDialog, QSQI.Indikator, QSQI.ArtDesWertes",
         "Ziel-Variable"),
        ("QS.Fortbildung.csv",        "Fortbildungsnachweise der Ärzte",
         "SO.QBID, QS.Fortbildungspflichtige, QS.Fortbildungsnachweis_Erbracht_Habende",
         "Merkmal: Fortbildungsquote"),
        ("FA.csv",                    "Fachabteilungen der Krankenhäuser",
         "FA.QBID, FA.FZ.Voll, FA.FZ.Teil, FA.Key301",
         "Merkmal: Ärzte pro Bett"),
        ("QS.csv",                    "QS-Berichtsbasis pro Standort",
         "QS.ID, SO.QBID, QS.Typ",
         "Verknüpfungstabelle"),
        ("QS.Leistungsbereich.csv",   "Leistungsbereiche mit Dokumentationsraten",
         "SO.QBID, QSLB.Dokumentationsrate, QSLB.Fallzahl",
         "Ergänzung"),
    ]
    add_table(doc,
              ["Datei", "Inhalt", "Wichtige Spalten", "Rolle"],
              rel_rows,
              col_widths=[4.5, 4.5, 5.5, 3.5])
    doc.add_paragraph()

    add_heading(doc, "2.4  Nicht relevante Tabellen", level=2)
    add_body(doc, "Folgende Tabellen wurden bewusst ausgeschlossen:")
    nicht_rel = [
        ("NM.csv",              "Nicht-medizinische Angebote (Parkplatz, Telefon)"),
        ("ICD.Code.csv",        "ICD-Diagnoseschlüssel (reine Lookup-Tabelle)"),
        ("OPS.csv",             "Operationsschlüssel (reine Lookup-Tabelle)"),
        ("Link.csv",            "URL-Links ohne Analysewert"),
        ("QS.Nachweis.csv",     "Technische Meta-Daten (Nachweiszeiträume)"),
    ]
    for datei, grund in nicht_rel:
        add_bullet(doc, f"{datei} — {grund}")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 3. DATENAUFBEREITUNG
    # ══════════════════════════════════════════════════════════════
    add_heading(doc, "3  Datenaufbereitung", level=1)

    add_heading(doc, "3.1  Vorgehen & Kriterien", level=2)
    add_body(doc,
        "Die Datenaufbereitung erfolgte vollständig reproduzierbar per Python-Skript "
        "(01_Exploration.ipynb). Kein manuelles Zusammenklicken — alle Schritte können "
        "durch erneutes Ausführen des Notebooks repliziert werden.")
    add_body(doc,
        "Sichtung der CSV-Dateien: Erste 3–5 Zeilen jeder Datei wurden gelesen "
        "(kein Python nötig, da CSV = Textdatei). Kriterium für Relevanz:")
    add_bullet(doc, "✅ Relevant: Enthält Strukturmerkmal aus der Aufgabenstellung ODER Qualitätsindikator-Bewertungen")
    add_bullet(doc, "⚠️ Möglicherweise: QS-relevante Infos, aber Bedeutung noch unklar")
    add_bullet(doc, "❌ Nicht relevant: Lookup-Tabellen, nicht-medizinische Angebote, reine Link-/Verwaltungsdaten")

    add_heading(doc, "3.2  Ziel-Variable", level=2)
    add_body(doc,
        "Die Ziel-Variable wurde aus QS.Qualitätsindikator.csv berechnet. "
        "Die Bewertungsspalte QSErgBewStrukDialog enthält den Bewertungscode "
        "des Strukturierten Dialogs:")
    add_bullet(doc, "R* (R10, R20, ...) = rechnerisch auffällig")
    add_bullet(doc, "N01, N02 = nicht auffällig")
    add_bullet(doc, "N99 = nicht bewertet → wird ausgeschlossen (nicht bewertet ≠ unauffällig!)")

    add_body(doc, "Berechnungsschritte:")
    schritte = [
        ("1", "Filter: QSQI.ArtDesWertes == 'QI'",
         "Nur echte Qualitätsindikatoren, keine Zählkennzahlen (EKez, TKez, ...)"),
        ("2", "Filter: QSErgBewStrukDialog != 'N99'",
         "Nur bewertete Indikatoren berücksichtigen"),
        ("3", "Deduplizierung: drop_duplicates(['SO.QBID', 'QSQI.Indikator'])",
         "Je Haus+Indikator nur eine Zeile behalten"),
        ("4", "Flag: ist_auffaellig = QSErgBewStrukDialog.str.startswith('R')",
         "Binäre Markierung pro Indikator"),
        ("5", "Aggregation: groupby('SO.QBID').agg(count, sum)",
         "Anzahl bewerteter QI und auffälliger QI pro Haus"),
        ("6", "Quote: auffaellig_n / total_qi",
         "Anteil auffälliger Indikatoren pro Haus"),
        ("7", "Target: quote > Median → hat_viele_Probleme = 1",
         "Binäre Ziel-Variable (0/1)"),
    ]
    add_table(doc, ["Schritt", "Code", "Erklärung"], schritte,
              col_widths=[1.5, 6.5, 6.0])
    doc.add_paragraph()

    add_heading(doc, "3.3  Merkmale (Features)", level=2)
    add_body(doc, "Folgende Merkmale wurden für die Analysetabelle ausgewählt:")
    merkmale_rows = [
        ("SO.Betten",          "SO.csv",              "Direkt verfügbar",     "Numerisch"),
        ("KH.Träger.Art",      "SO.csv",              "Direkt verfügbar",     "Kategorial: privat / freigemeinnützig / öffentlich"),
        ("SO.Bundesland",      "SO.csv",              "Direkt verfügbar",     "Kategorial: 16 Bundesländer"),
        ("SO.Uni",             "SO.csv",              "Direkt verfügbar",     "Binär: 0/1"),
        ("SO.Latitude/Long.",  "SO.csv",              "Direkt verfügbar",     "Numerisch (für Karte)"),
        ("fortbildungsquote",  "QS.Fortbildung.csv",  "Berechnet: Erbracht / Pflichtige", "Numerisch 0–1"),
        ("Ärzte pro Bett",     "FA.csv",              "Noch zu berechnen",    "Numerisch"),
    ]
    add_table(doc,
              ["Merkmal", "Quelle", "Berechnung", "Typ"],
              merkmale_rows,
              col_widths=[3.5, 3.5, 5.0, 4.0])
    doc.add_paragraph()

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 4. ERGEBNISSE
    # ══════════════════════════════════════════════════════════════
    add_heading(doc, "4  Ergebnisse", level=1)

    add_heading(doc, "4.1  Ziel-Variable — Statistiken", level=2)
    add_body(doc, f"Datenbasis nach Aufbereitung: {n_haeuser:,} Krankenhäuser")
    stat_rows = [
        ("Anzahl Krankenhäuser",              f"{n_haeuser:,}"),
        ("Ø Indikatoren pro Haus",            f"{(total_qi_sum / n_haeuser):.1f}"),
        ("Median auffällig-Quote",            f"{median_quote:.2%}"),
        ("Mittelwert auffällig-Quote",        f"{stats['mean']:.2%}"),
        ("Min. auffällig-Quote",              f"{stats['min']:.2%}"),
        ("Max. auffällig-Quote",              f"{stats['max']:.2%}"),
        ("Häuser mit vielen Problemen (=1)",  f"{n_auffaellig:,} ({n_auffaellig / n_haeuser:.1%})"),
        ("Häuser ohne viele Probleme (=0)",   f"{n_haeuser - n_auffaellig:,} ({(n_haeuser - n_auffaellig) / n_haeuser:.1%})"),
    ]
    add_table(doc, ["Kennzahl", "Wert"], stat_rows, col_widths=[8.0, 4.0])
    doc.add_paragraph()
    add_body(doc,
        "Die Ziel-Variable ist nahezu ausgewogen verteilt (ca. 49% vs. 51%), "
        "was für Machine-Learning-Modelle optimal ist.")

    add_heading(doc, "4.2  Analysetabelle", level=2)
    add_body(doc,
        f"Die finale Analysetabelle (analysetabelle.csv) enthält {analyse.shape[0]:,} Zeilen "
        f"und {analyse.shape[1]} Spalten. Jede Zeile repräsentiert genau ein Krankenhaus.")
    col_rows = [(c, str(analyse[c].dtype)) for c in analyse.columns]
    add_table(doc, ["Spalte", "Datentyp"], col_rows, col_widths=[8.0, 4.0])
    doc.add_paragraph()

    add_heading(doc, "4.3  Wozu wird die Analysetabelle genutzt?", level=2)
    add_body(doc,
        "Die Analysetabelle ist die einzige Datengrundlage für alle weiteren Projektschritte. "
        "Das Prinzip: Rohdaten → Analysetabelle → alles andere.")
    nutzung_rows = [
        ("Baustein 2\nDeskriptive Analyse",
         "Grafiken (Box-Plots, Scatter-Plots, Balkendiagramme) werden direkt aus der Tabelle erzeugt. "
         "Beispiel: \"Unterscheiden sich Bettenzahl zwischen Häusern mit/ohne viele Probleme?\""),
        ("Baustein 3\nDashboard Seite 1",
         "hat_viele_Probleme + Koordinaten → Deutschland-Karte mit regionaler Verteilung"),
        ("Baustein 3\nDashboard Seite 2",
         "Merkmale gruppiert nach hat_viele_Probleme → Vergleichsdiagramme mit Dropdown"),
        ("Baustein 3\nDashboard Seite 3",
         "Filter nach Betten / Region / Träger → ähnliche Häuser finden und deren Qualität zeigen"),
        ("Baustein 4\nDecision Tree",
         "Feature Matrix X = Betten, Träger, Bundesland, Uni, Fortbildungsquote; "
         "Zielvariable y = hat_viele_Probleme; direkt für train_test_split und DecisionTreeClassifier nutzbar"),
    ]
    add_table(doc, ["Baustein", "Nutzung der Analysetabelle"], nutzung_rows,
              col_widths=[4.0, 12.0])
    doc.add_paragraph()

    add_heading(doc, "4.4  Fehlende Werte", level=2)
    missing = analyse.isnull().sum()
    missing_rows = [(col, int(missing[col]), f"{missing[col] / len(analyse):.1%}")
                    for col in analyse.columns if missing[col] > 0]
    if missing_rows:
        add_table(doc, ["Spalte", "Fehlende Werte (n)", "Anteil"], missing_rows,
                  col_widths=[6.0, 4.0, 4.0])
        doc.add_paragraph()
        add_body(doc,
            "KH.Träger.Art (28 fehlend) und fortbildungsquote (33 fehlend) haben sehr "
            "geringe Fehlquoten (<2%) und können für die Analyse entweder mit dem Modus "
            "imputiert oder als eigene Kategorie 'unbekannt' behandelt werden.")
    else:
        add_body(doc, "Keine fehlenden Werte.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 5. DESKRIPTIVE ANALYSE — BEFUNDE
    # ══════════════════════════════════════════════════════════════
    add_heading(doc, "5  Deskriptive Analyse — Befunde", level=1)
    add_body(doc,
        "Die Analyse wurde in 02_Analyse.ipynb durchgeführt. Grundlage: analysetabelle.csv. "
        "10 Grafiken wurden erstellt, jede mit automatisch berechnetem Befundsatz. "
        "Farbschema einheitlich: grün = wenige Probleme, rot = viele Probleme.")

    add_heading(doc, "5.1  Übersicht der Befunde", level=2)
    df_analyse = analyse.copy()
    m_betten_0 = int(df_analyse[df_analyse["hat_viele_Probleme"] == 0]["SO.Betten"].median())
    m_betten_1 = int(df_analyse[df_analyse["hat_viele_Probleme"] == 1]["SO.Betten"].median())
    pct_privat   = df_analyse[df_analyse["KH.Träger.Art"] == "privat"]["hat_viele_Probleme"].mean()
    pct_frei     = df_analyse[df_analyse["KH.Träger.Art"] == "freigemeinnützig"]["hat_viele_Probleme"].mean()
    pct_oeffentl = df_analyse[df_analyse["KH.Träger.Art"] == "öffentlich"]["hat_viele_Probleme"].mean()
    pct_uni      = df_analyse[df_analyse["SO.Uni"] == 1]["hat_viele_Probleme"].mean()
    pct_normal   = df_analyse[df_analyse["SO.Uni"] == 0]["hat_viele_Probleme"].mean()
    fb_0 = df_analyse[df_analyse["hat_viele_Probleme"] == 0]["fortbildungsquote"].median()
    fb_1 = df_analyse[df_analyse["hat_viele_Probleme"] == 1]["fortbildungsquote"].median()
    ab_0 = df_analyse[df_analyse["hat_viele_Probleme"] == 0]["aerzte_pro_bett"].median()
    ab_1 = df_analyse[df_analyse["hat_viele_Probleme"] == 1]["aerzte_pro_bett"].median()

    befund_rows = [
        ("Grafik 1\nauffällig-Quote",
         f"Median {df_analyse['auffaellig_quote'].median():.0%}, linkssteil verteilt — "
         f"die meisten Häuser liegen zwischen 60–90%."),
        ("Grafik 2\nBettenzahl",
         f"Median: Wenige Probleme = {m_betten_0} Betten, Viele Probleme = {m_betten_1} Betten. "
         f"Kein klarer Größenunterschied."),
        ("Grafik 3\nTrägerschaft",
         f"Privat: {pct_privat:.1%} | Freigemeinnützig: {pct_frei:.1%} | Öffentlich: {pct_oeffentl:.1%} "
         f"haben viele Probleme. Private Häuser auffällig höher."),
        ("Grafik 4\nUni-Kliniken",
         f"Uni-Kliniken: {pct_uni:.1%} vs. normale Häuser: {pct_normal:.1%} — kaum Unterschied."),
        ("Grafik 5+6\nFortbildung & Ärzte/Bett",
         f"Fortbildungsquote: kein Unterschied (Md={fb_0:.3f} vs. {fb_1:.3f}). "
         f"Ärzte/Bett: Wenige={ab_0:.3f}, Viele={ab_1:.3f} — leichter Unterschied sichtbar."),
        ("Grafik 7\nBundesland",
         "Saarland höchster Anteil (63,2%), Berlin niedrigster (33,3%). "
         "Regionale Unterschiede sichtbar — kleine n beachten."),
        ("Grafik 8\nKorrelationsmatrix",
         "Stärkste Korrelation mit Ziel-Variable: total_qi (r=−0,28), "
         "dann aerzte_pro_bett (r=−0,14). Fortbildungsquote praktisch keine Korrelation."),
        ("Grafik 9\nScatter Betten/Ärzte",
         "Kein klares Trennmuster zwischen den Gruppen — Überlappung stark."),
        ("Grafik 10\nStörfaktor Träger×Betten",
         f"Private Häuser sind kleiner (Md={int(df_analyse[df_analyse['KH.Träger.Art'] == 'privat']['SO.Betten'].median())} Betten). "
         "Der Trägereffekt muss daher mit Vorsicht interpretiert werden."),
    ]
    add_table(doc, ["Grafik", "Befund"], befund_rows, col_widths=[3.5, 12.5])
    doc.add_paragraph()

    add_heading(doc, "5.2  Grafiken", level=2)
    grafik_dateien = [
        ("grafiken/g1_auffaellig_quote.png",       "Grafik 1: Verteilung der auffällig-Quote"),
        ("grafiken/g2_bettenzahl.png",              "Grafik 2: Bettenzahl"),
        ("grafiken/g3_traegerschaft.png",           "Grafik 3: Trägerschaft"),
        ("grafiken/g4_uni.png",                     "Grafik 4: Uni-Kliniken vs. normale Häuser"),
        ("grafiken/g5_6_fortbildung_aerzte.png",    "Grafik 5+6: Fortbildungsquote & Ärzte pro Bett"),
        ("grafiken/g7_bundesland.png",              "Grafik 7: Anteil Häuser mit vielen Problemen je Bundesland"),
        ("grafiken/g8_korrelation.png",             "Grafik 8: Korrelationsmatrix"),
        ("grafiken/g9_scatter_betten_aerzte.png",   "Grafik 9: Scatter — Bettenzahl vs. Ärzte pro Bett"),
        ("grafiken/g10_stoerfaktor_traeger.png",    "Grafik 10: Störfaktor — Bettengröße je Trägerschaft"),
    ]
    for pfad, titel in grafik_dateien:
        if os.path.exists(pfad):
            p = doc.add_paragraph()
            run = p.add_run(titel)
            run.bold = True
            run.font.size = Pt(11)
            doc.add_picture(pfad, width=Cm(15))
            doc.add_paragraph()
        else:
            add_body(doc, f"[Grafik nicht gefunden: {pfad}]")

    add_heading(doc, "5.2  Gesamteinschätzung", level=2)
    add_body(doc,
        "Die Analyse zeigt keine starken, eindeutigen Zusammenhänge zwischen den untersuchten "
        "Strukturmerkmalen und der Ziel-Variable. Der stärkste Prädiktor ist total_qi "
        "(Anzahl bewerteter Indikatoren) — ein strukturelles Merkmal des Hauses, kein "
        "Qualitätsmerkmal: Häuser mit mehr bewerteten Indikatoren haben tendenziell niedrigere "
        "Auffälligkeitsquoten.")
    add_body(doc,
        "Einziger klarer inhaltlicher Befund: Private Häuser haben mit 56,5% einen höheren "
        "Anteil als freigemeinnützige (46,4%) und öffentliche (46,7%) Träger. Allerdings sind "
        "private Häuser im Median deutlich kleiner — der Trägereffekt könnte durch "
        "Größenunterschiede mitverursacht sein (Störfaktor).")
    add_body(doc,
        "Wichtig: Das ist ein valides Ergebnis. Schwache Zusammenhänge sind in echten "
        "Gesundheitsdaten normal, da viele weitere Faktoren (Patientenmix, Spezialisierung, "
        "Dokumentationsqualität) eine Rolle spielen, die nicht im Datensatz enthalten sind.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 6. OFFENE PUNKTE & NÄCHSTE SCHRITTE
    # ══════════════════════════════════════════════════════════════
    add_heading(doc, "6  Offene Punkte & Nächste Schritte", level=1)

    offen = [
        ("Baustein 3: Streamlit-Dashboard",
         "Seite 1: Übersicht + Karte | Seite 2: Vergleiche | Seite 3: Ähnliche Häuser"),
        ("Baustein 4: Decision Tree (Bonus)",
         "max_depth=3, Train-Test-Split, Modell speichern (joblib)"),
        ("Baustein 5: Präsentation & Dokumentation",
         "Startanleitung, Entscheidungsbegründungen, Live-Demo"),
        ("Pflegekräfte pro Bett",
         "Noch nicht berechnet — aus FA.Personalliste.csv extrahieren"),
    ]
    add_table(doc, ["Aufgabe", "Details"], offen, col_widths=[5.5, 10.5])
    doc.add_paragraph()

    add_body(doc,
        "Hinweis: Ein Entscheidungsbaum auf Basis der bereits vorliegenden Analysetabelle "
        "kann jederzeit trainiert werden. Die Datengrundlage ist vollständig.")

    return doc


# ══════════════════════════════════════════════════════════════════
# MAIN
# ══════════════════════════════════════════════════════════════════
if __name__ == "__main__":
    print("Lade Daten aus analysetabelle.csv ...")
    daten = lade_daten()
    print(f"  Häuser: {daten['n_haeuser']:,}  |  Median-Quote: {daten['median_quote']:.2%}")

    print("Erzeuge Word-Dokument ...")
    doc = erzeuge_dokument(daten)

    doc.save(str(OUT_PATH))
    print(f"✅ Dokument gespeichert: {OUT_PATH}")