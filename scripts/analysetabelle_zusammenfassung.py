"""
analysetabelle_zusammenfassung.py
==================================
Erzeugt eine kompakte Word-Zusammenfassung fuer Kollegen:
Welche Spalten/Merkmale wurden gewaehlt, welche Ziel-Variable(n),
aus welchen Rohdaten-Tabellen, wie wurde gemerged (Schluessel/Kriterium),
und wie gross ist die fertige Analysetabelle geworden.

Aufruf:
    python scripts/analysetabelle_zusammenfassung.py
"""
import os
from datetime import date, datetime
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PROJEKT_ROOT = Path(__file__).resolve().parent.parent
os.chdir(PROJEKT_ROOT)

DATA_PATH = PROJEKT_ROOT / "Data" / "analysetabelle.csv"
OUT_PATH = PROJEKT_ROOT / "Doku" / "Word" / "Analysetabelle_Zusammenfassung.docx"

C_BLAU = RGBColor(0x1F, 0x49, 0x7D)
C_WEISS = RGBColor(0xFF, 0xFF, 0xFF)
F_BLAU = "DEEAF1"
F_GRAU = "F2F2F2"


def set_bg(cell, fill):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def hdr_cell(cell):
    set_bg(cell, "1F497D")
    for run in cell.paragraphs[0].runs:
        run.font.color.rgb = C_WEISS
        run.bold = True
        run.font.size = Pt(9.5)


def add_tbl(doc, headers, rows, cw=None, zebra=True):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        hdr_cell(c)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri + 1].cells[ci]
            c.text = str(val)
            c.paragraphs[0].runs[0].font.size = Pt(9)
            if zebra and ri % 2 == 1:
                set_bg(c, F_GRAU)
    if cw:
        for row in t.rows:
            for i, w in enumerate(cw):
                row.cells[i].width = Cm(w)
    return t


def add_h(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = C_BLAU
    return p


def add_b(doc, text, size=10.5, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    return p


def de(x, d=0):
    """Zahl im deutschen Format: Punkt=Tausendertrenner, Komma=Dezimaltrenner."""
    s = f"{x:,.{d}f}"
    return s.replace(",", "§").replace(".", ",").replace("§", ".")


def depct(x, d=1):
    """Anteil (0-1) als deutsch formatierte Prozentzahl, z. B. 0.565 -> '56,5 %'."""
    return de(x * 100, d) + " %"


# ── Daten laden (fuer echte Kennzahlen, keine hartkodierten Zahlen) ──
analyse = pd.read_csv(DATA_PATH, low_memory=False)
n_rows, n_cols = analyse.shape
n_probleme = int(analyse["hat_viele_Probleme"].sum())
n_konzern = int(analyse["ist_konzern"].sum())
missing = analyse.isnull().sum()

# ══════════════════════════════════════════════════════════════════
doc = Document()
doc.core_properties.author = "Datenanalyse-Team"
doc.core_properties.created = datetime.now()
doc.core_properties.modified = datetime.now()
doc.core_properties.title = "Analysetabelle — Zusammenfassung fuer Kollegen"

for s in doc.sections:
    s.top_margin = Cm(2.2)
    s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.5)
    s.right_margin = Cm(2.0)

# ── Titelseite ────────────────────────────────────────────────────
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Analysetabelle — Zusammenfassung für Kollegen")
r.bold = True
r.font.size = Pt(20)
r.font.color.rgb = C_BLAU

doc.add_paragraph()
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.add_run(
    "Welche Spalten wurden gewählt, aus welchen Tabellen, wie wurde gemerged, "
    "und wie groß ist das Ergebnis?"
).font.size = Pt(12)

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run(
    f"Erstellt: {date.today().strftime('%d.%m.%Y')}  |  Quelle: Notebooks/01_Exploration.ipynb  |  "
    f"Datei: Data/analysetabelle.csv"
).font.size = Pt(10)
doc.add_paragraph()

add_b(
    doc,
    f"Kurzfassung: Die fertige Analysetabelle hat {de(n_rows)} Zeilen (1 Zeile = 1 Krankenhaus) "
    f"und {n_cols} Spalten. Davon sind 7 Spalten echte Analysemerkmale, der Rest sind "
    f"Schlüssel-/Hilfsspalten oder die Ziel-Variable(n). Alle Zusammenführungen erfolgten "
    f"als Left Join über SO.QBID (Ausnahme: Konzern.csv über SO.Standortnummer).",
    bold=True,
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 1. ZIEL-VARIABLE(N)
# ══════════════════════════════════════════════════════════════════
add_h(doc, "1  Ziel-Variable(n)", level=1)
add_b(
    doc,
    "Es gibt zwei Ziel-Variablen: eine kontinuierliche Zwischengröße und die eigentliche, "
    "binäre Ziel-Variable, die daraus abgeleitet wird."
)

add_tbl(
    doc,
    ["Spalte", "Typ", "Quelle", "Bedeutung"],
    [
        ("auffaellig_quote", "Kontinuierlich (Anteil, z. B. 0,77 = 77 %)", "QS.Qualitätsindikator.csv",
         "Anteil der Qualitätsindikatoren eines Hauses, die als 'auffällig' (R*) bewertet wurden — "
         "kein 0/1-Wert, sondern ein Dezimalwert zwischen 0,0 und 1,0 mit 768 verschiedenen "
         "Ausprägungen im Datensatz. Zwischengröße, aus der die eigentliche Ziel-Variable abgeleitet wird."),
        ("hat_viele_Probleme", "Binär, nimmt nur 0 oder 1 an", "berechnet aus auffaellig_quote",
         f"1 = auffaellig_quote liegt über dem Median (76,92 %) → 'hat überdurchschnittlich viele "
         f"Probleme'. {de(n_probleme)} von {de(n_rows)} Häusern ({depct(n_probleme/n_rows)}) haben Wert 1. "
         "Das ist die eigentliche Ziel-Variable (y) für Baustein 2 und 4."),
    ],
    cw=[3.5, 2.5, 4.0, 8.0],
)
add_b(
    doc,
    "Warum Median als Schwelle: teilt die Häuser automatisch in zwei etwa gleich große Gruppen "
    "— eine ausgewogene Klassenverteilung ist für Gruppenvergleiche und ein Klassifikationsmodell "
    "wichtiger als ein willkürlicher fester Wert (z. B. 80 %)."
)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════
# 2. MERKMALE (FEATURES)
# ══════════════════════════════════════════════════════════════════
add_h(doc, "2  Merkmale (X) — 7 Stück", level=1)
add_b(
    doc,
    "Das sind die Spalten, die tatsächlich als Merkmale in Vergleiche (Baustein 2) und das "
    "Modell (Baustein 4) einfließen:"
)

add_tbl(
    doc,
    ["#", "Merkmal", "Quelltabelle(n)", "Berechnung"],
    [
        ("1", "SO.Betten", "SO.csv", "Direkt übernommen — Bettenzahl"),
        ("2", "SO.Bundesland", "SO.csv", "Direkt übernommen — Region (16 Bundesländer)"),
        ("3", "SO.Uni", "SO.csv", "Direkt übernommen — Universitätsklinikum ja/nein (0/1)"),
        ("4", "KH.Träger.Art", "SO.csv", "Direkt übernommen — privat/freigemeinnützig/öffentlich"),
        ("5", "fortbildungsquote", "QS.Fortbildung.csv",
         "Berechnet: Fortbildungsnachweis_Erbracht_Habende ÷ Fortbildungspflichtige"),
        ("6", "aerzte_pro_bett", "FA.Personalliste.csv × FA.csv",
         "Berechnet: Summe Ärzte pro Haus (Filter Bereich='Ärzte', über ABTID→SO.QBID aggregiert) "
         "÷ SO.Betten"),
        ("7", "pflege_pro_bett", "SO.Personalliste.csv",
         "Berechnet: Summe Pflegekräfte pro Haus (Filter Bereich='Pflege') ÷ SO.Betten"),
    ],
    cw=[0.8, 3.5, 4.5, 9.2],
)
doc.add_paragraph()

add_b(doc, "Zusätzlich, aber mit anderer Rolle:", bold=True)
add_tbl(
    doc,
    ["Spalte", "Quelle", "Rolle"],
    [
        ("ist_konzern", "Konzern.csv × SO.csv",
         f"Binäres Merkmal (0/1), {de(n_konzern)} von {de(n_rows)} Häusern ({depct(n_konzern/n_rows)}) sind "
         "Konzernhäuser. Wird im Modell mitgeführt, zeigt aber in Baustein 2/4 keinen "
         "signifikanten Zusammenhang mit der Ziel-Variable (Chi²-Test p=0,90)."),
        ("SO.Latitude / SO.Longitude", "SO.csv",
         "Geo-Koordinaten — kein Merkmal für das Modell, sondern Grundlage der Deutschlandkarte "
         "im Dashboard."),
    ],
    cw=[4.0, 4.0, 9.5],
)
doc.add_paragraph()

add_b(doc, "Reine Schlüssel- und Anzeige-Spalten (keine Merkmale):", bold=True)
add_tbl(
    doc,
    ["Spalte", "Rolle"],
    [
        ("SO.QBID", "Primärschlüssel — eindeutige Krankenhaus-ID, verbindet alle Quelltabellen"),
        ("SO.Name", "Nur für Anzeige (z. B. Dashboard-Steckbrief), kein Analysemerkmal"),
        ("KH.Träger", "Trägername im Klartext — unbereinigt, KH.Träger.Art ist die auswertbare Version"),
        ("SO.Standortnummer", "Reine Join-Hilfsspalte, nur für den Konzern.csv-Abgleich benötigt"),
        ("total_qi, auffaellig_n", "Zwischenwerte zur Berechnung von auffaellig_quote (Zähler/Nenner)"),
    ],
    cw=[4.0, 13.5],
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 3. QUELLTABELLEN
# ══════════════════════════════════════════════════════════════════
add_h(doc, "3  Quelltabellen — 7 von 86 CSV-Dateien", level=1)
add_b(
    doc,
    "Von den 86 Rohdaten-Dateien fließen genau 7 tatsächlich in die Analysetabelle ein "
    "(vollständige Klassifikation aller 86 Dateien: siehe Doku/MD/Daten_Inhaltsverzeichnis.md):"
)

add_tbl(
    doc,
    ["Datei", "Liefert", "Rolle im Merge"],
    [
        ("QS.Qualitätsindikator.csv", "auffaellig_quote, hat_viele_Probleme",
         "Ausgangspunkt — hier beginnt die Analysetabelle (Ziel-Variable)"),
        ("SO.csv", "SO.Betten, SO.Bundesland, SO.Uni, KH.Träger.Art, Koordinaten, SO.Standortnummer",
         "1. Anbindung — Stammdaten je Haus"),
        ("QS.Fortbildung.csv", "fortbildungsquote", "2. Anbindung"),
        ("FA.csv", "Brückentabelle (ABTID ↔ SO.QBID)", "Nötig, um FA.Personalliste.csv anzubinden"),
        ("FA.Personalliste.csv", "aerzte_pro_bett", "3. Anbindung (zweistufig über FA.csv)"),
        ("SO.Personalliste.csv", "pflege_pro_bett", "4. Anbindung (ergänzt 2026-07-29)"),
        ("Konzern.csv", "ist_konzern", "5. Anbindung (ergänzt 2026-07-29, über SO.Standortnummer)"),
    ],
    cw=[4.5, 6.0, 7.0],
)
doc.add_paragraph()
add_b(
    doc,
    "Nicht verwendet, obwohl thematisch naheliegend: QS.csv (vermeintliche Brückentabelle — "
    "tatsächlich nie geladen, da QS.Qualitätsindikator.csv SO.QBID bereits selbst trägt) und "
    "AQ.Pflege.csv (bewusst durch SO.Personalliste.csv ersetzt, da diese direkt Anzahlen statt "
    "nur Qualifikationsnachweise liefert)."
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 4. MERGE-KRITERIUM
# ══════════════════════════════════════════════════════════════════
add_h(doc, "4  Wie wurde gemerged?", level=1)
add_b(
    doc,
    "Alle sieben Zusammenführungen sind Left Joins (pandas merge(..., how='left')). "
    "Anker ist immer die bereits im Aufbau befindliche Analysetabelle (beginnend bei der "
    "Ziel-Variable) — jede weitere Tabelle wird von links angehängt, ohne dass Zeilen verloren "
    "gehen, wenn eine Datei zu einem Haus keine Daten hat (dann entsteht dort NaN, siehe Kapitel 5)."
)

add_tbl(
    doc,
    ["Merge", "Schlüssel", "Join-Typ", "Besonderheit"],
    [
        ("Analysetabelle ← SO.csv", "SO.QBID", "Left", "Standardfall"),
        ("Analysetabelle ← QS.Fortbildung.csv", "SO.QBID", "Left", "Standardfall"),
        ("FA.Personalliste.csv ← FA.csv", "ABTID → SO.QBID", "Left",
         "Zweistufig: FA.Personalliste kennt nur ABTID, FA.csv liefert den Übergang zu SO.QBID"),
        ("Analysetabelle ← FA.Personalliste.csv (aggregiert)", "SO.QBID", "Left", "Standardfall"),
        ("Analysetabelle ← SO.Personalliste.csv (aggregiert)", "SO.QBID", "Left", "Standardfall"),
        ("Analysetabelle ← Konzern.csv", "SO.Standortnummer", "Left",
         "⚠️ NICHT SO.QBID! Ursprünglicher Join-Versuch verglich fälschlich gegen SO.QBID "
         "(0 Treffer) — nach Korrektur 358 von 1.824 Treffern"),
    ],
    cw=[5.5, 3.5, 2.0, 6.5],
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 5. ENDERGEBNIS
# ══════════════════════════════════════════════════════════════════
add_h(doc, "5  Endergebnis — wie groß ist die fertige Tabelle?", level=1)

add_tbl(
    doc,
    ["Kennzahl", "Wert"],
    [
        ("Datei", "Data/analysetabelle.csv"),
        ("Zeilen (Krankenhäuser)", de(n_rows)),
        ("Spalten insgesamt", f"{n_cols}"),
        ("davon echte Merkmale (Kapitel 2)", "7"),
        ("Ziel-Variable(n)", "2 (auffaellig_quote, hat_viele_Probleme)"),
        ("Schlüssel-/Hilfsspalten", f"{n_cols - 7 - 2}"),
        ("hat_viele_Probleme = 1", f"{de(n_probleme)} ({depct(n_probleme/n_rows)})"),
        ("hat_viele_Probleme = 0", f"{de(n_rows - n_probleme)} ({depct((n_rows - n_probleme)/n_rows)})"),
    ],
    cw=[9.0, 8.0],
)
doc.add_paragraph()

add_b(doc, "Fehlende Werte je Spalte (nur Spalten mit Lücken):", bold=True)
missing_rows = [
    (col, int(missing[col]), depct(missing[col] / n_rows))
    for col in analyse.columns
    if missing[col] > 0
]
if missing_rows:
    add_tbl(doc, ["Spalte", "Fehlende Werte (n)", "Anteil"], missing_rows, cw=[7.0, 5.0, 5.0])
else:
    add_b(doc, "Keine fehlenden Werte.")
doc.add_paragraph()

add_b(
    doc,
    "Hinweis: Kein Haus geht beim Merge verloren — QS.Qualitätsindikator.csv (der Ausgangspunkt) "
    "hat bereits genau 1.824 eindeutige Häuser, dieselbe Zahl wie die fertige Tabelle. Häuser, "
    "die in einer der ergänzenden Dateien fehlen (z. B. keine Fortbildungsdaten), bekommen dort "
    "NaN statt aus der Tabelle zu verschwinden.",
    size=10,
)

# ── Speichern ─────────────────────────────────────────────────────
doc.save(str(OUT_PATH))
print(f"✅ Gespeichert: {OUT_PATH}")
