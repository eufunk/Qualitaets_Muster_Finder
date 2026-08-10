"""
erstelle_praesentationsskript.py
=================================
Erzeugt Doku/Praesentationsskript_Qualitaets_Muster_Finder.docx —
vollständiges Präsentationsskript mit erzählendem Sprechertext je Folie
und fließenden Übergängen.

Aufruf: python scripts/erstelle_praesentationsskript.py
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

PROJEKT_ROOT = Path(__file__).resolve().parent.parent
ZIEL         = PROJEKT_ROOT / "Doku" / "Praesentationsskript_Qualitaets_Muster_Finder.docx"

BLAU   = RGBColor(0x1F, 0x49, 0x7D)
ORANGE = RGBColor(0xBF, 0x5A, 0x00)
GRAU   = RGBColor(0x55, 0x55, 0x55)
GRUEN  = RGBColor(0x37, 0x5E, 0x23)

doc = Document()

# ── Seitenränder ─────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── Hilfsfunktionen ───────────────────────────────────────────────

def absatz(text, fett=False, farbe=None, size=11, align=None, italic=False, abstand_vor=0, abstand_nach=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(abstand_vor)
    p.paragraph_format.space_after  = Pt(abstand_nach)
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold   = fett
    run.italic = italic
    run.font.size = Pt(size)
    if farbe:
        run.font.color.rgb = farbe
    return p

def folie_heading(nummer, titel, zeitangabe=""):
    doc.add_paragraph()
    p = doc.add_heading(level=2)
    p.clear()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(f"Folie {nummer}  —  {titel}")
    run.font.size  = Pt(13)
    run.font.bold  = True
    run.font.color.rgb = BLAU
    if zeitangabe:
        run2 = p.add_run(f"  [{zeitangabe}]")
        run2.font.size  = Pt(10)
        run2.font.bold  = False
        run2.font.color.rgb = GRAU

def block_heading(text):
    p = doc.add_heading(level=1)
    p.clear()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.font.size  = Pt(14)
    run.font.bold  = True
    run.font.color.rgb = BLAU

def uebergang(text):
    """Grau kursiver Übergangshinweis"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(f"→ Übergang: {text}")
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = ORANGE

def trennlinie():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)

# ═════════════════════════════════════════════════════════════════
# DECKBLATT
# ═════════════════════════════════════════════════════════════════
h = doc.add_heading(level=0)
h.clear()
h.alignment = WD_ALIGN_PARAGRAPH.CENTER
h.paragraph_format.space_before = Pt(40)
r = h.add_run("Qualitäts-Muster-Finder")
r.font.size  = Pt(22)
r.font.bold  = True
r.font.color.rgb = BLAU

absatz("Präsentationsskript — Sprechertext je Folie", fett=False, farbe=GRAU,
       size=13, align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, abstand_vor=6, abstand_nach=4)
absatz("~32 Minuten · 17 Folien · Einzelpräsentation",
       farbe=GRAU, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, abstand_vor=0, abstand_nach=4)
absatz(f"Stand: {datetime.date.today().strftime('%d.%m.%Y')}",
       farbe=GRAU, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, abstand_vor=0, abstand_nach=40)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════
# VORBEMERKUNG
# ═════════════════════════════════════════════════════════════════
absatz("Hinweis zur Nutzung", fett=True, farbe=BLAU, size=12, abstand_vor=0, abstand_nach=4)
absatz(
    "Dieses Skript enthält den vollständigen Sprechertext für alle 17 Folien. "
    "Der Text ist erzählend formuliert und kann direkt vorgelesen oder als Gedächtnisstütze "
    "für freies Sprechen genutzt werden. Übergänge zwischen den Folien sind kursiv hervorgehoben.",
    farbe=GRAU, size=10, abstand_vor=0, abstand_nach=20
)
trennlinie()

# ═════════════════════════════════════════════════════════════════
# BLOCK 1 — EINSTIEG & PROJEKTFRAGE  (Folien 1–3)
# ═════════════════════════════════════════════════════════════════
block_heading("Block 1 — Einstieg & Projektfrage  (Folien 1–3 · ca. 4 Min)")

# ── Folie 1 ──────────────────────────────────────────────────────
folie_heading(1, "Titelfolie", "ca. 1 Min")
absatz(
    "Herzlich willkommen. Unser Projekt trägt den Namen Qualitäts-Muster-Finder. "
    "Wie der Name schon andeutet, geht es um Muster in den Daten: Wir haben konkrete "
    "Datensätze aus dem deutschen Gesundheitswesen untersucht und uns gefragt, welche "
    "Krankenhausmerkmale damit zusammenhängen, dass ein Haus überdurchschnittlich viele "
    "Qualitätsprobleme aufweist. Dabei zeigen wir auch, wenn sich kein klarer Zusammenhang "
    "ergeben hat — denn auch das ist ein valides und interessantes Ergebnis."
)
uebergang("Bevor wir ins Detail gehen — hier kurz der Fahrplan für die nächsten 30 Minuten.")

# ── Folie 2 ──────────────────────────────────────────────────────
folie_heading(2, "Agenda", "ca. 1 Min")
absatz(
    "Wir laufen heute den kompletten Weg einer Datenanalyse ab — von den Rohdaten "
    "bis zum interaktiven Dashboard. Zuerst klären wir die Projektfrage und stellen "
    "den Datensatz vor. Dann zeigen wir, wie wir aus 86 CSV-Dateien eine saubere "
    "Analysetabelle gebaut haben. Im dritten Block kommen die eigentlichen Ergebnisse: "
    "Welche Merkmale hängen mit Qualitätsproblemen zusammen — und welche nicht? "
    "Im vierten Block folgen Machine Learning und das interaktive Dashboard. "
    "Und zum Abschluss ziehen wir ein Fazit und zeigen, wo die Grenzen dieser Analyse liegen."
)
uebergang("Fangen wir mit der zentralen Frage an, um die es in unserem Projekt geht.")

# ── Folie 3 ──────────────────────────────────────────────────────
folie_heading(3, "Projektfrage & Datensatz", "ca. 2 Min")
absatz(
    "Die Frage lautet: Gibt es Zusammenhänge zwischen den Strukturmerkmalen eines "
    "Krankenhauses und der Häufigkeit seiner auffälligen Qualitätsindikatoren? "
    "Qualitätsindikatoren sind dabei messbare Kennzahlen, mit denen die Qualität der "
    "medizinischen Versorgung bewertet wird — z. B. der Anteil der Patientinnen, bei denen "
    "während einer gynäkologischen Bauchspiegelung ein umliegendes Organ verletzt wurde, "
    "die Rate der Patientinnen und Patienten, die nach einem Herzkathetereingriff keine "
    "Beschwerden wie Schmerzen oder Taubheitsgefühle aufweisen, oder der Anteil der "
    "Patientinnen und Patienten, bei denen während des Krankenhausaufenthalts ein "
    "Druckgeschwür — Dekubitus — entstanden ist. "
    "Alle drei Beispiele stammen direkt aus unserem Datensatz, aus QS.Qualitätsindikator.csv, "
    "Leistungsbereiche GYN-OP, PCI und DEK."
)
absatz(
    "Der Hintergrund: Jedes deutsche Krankenhaus ist gesetzlich verpflichtet, jährlich "
    "einen Qualitätsbericht zu veröffentlichen. Das macht diesen Datensatz besonders wertvoll — "
    "alle Häuser werden nach denselben Regeln bewertet. Das ist die Grundvoraussetzung "
    "für einen fairen Vergleich."
)
absatz(
    "Zur Datenbasis: Wir arbeiten mit Qualitätsberichten aus dem Jahr 2023, "
    "herausgegeben vom Institut für Qualitätssicherung und Transparenz im Gesundheitswesen — "
    "kurz IQTIG. Der Datensatz umfasst 1.824 Krankenhäuser, die tatsächlich bewertet wurden, "
    "verteilt auf 86 CSV-Dateien mit insgesamt rund 1,2 Gigabyte. Pro Haus gibt es "
    "etwa 150 Qualitätsindikatoren. Dabei kann das Ergebnis in beide Richtungen gehen: "
    "Kein Zusammenhang ist genauso ein valides Ergebnis wie ein starker Zusammenhang."
)
uebergang("Bevor wir tiefer einsteigen — kurz ein Überblick, wie wir uns durch diesen Datensatz durchgearbeitet haben.")
trennlinie()

# ═════════════════════════════════════════════════════════════════
# BLOCK 2 — DATENVORBEREITUNG  (Folien 4–7 · ca. 7 Min)
# ═════════════════════════════════════════════════════════════════
block_heading("Block 2 — Datenvorbereitung  (Folien 4–8 · ca. 8 Min)")

# ── Folie 4 ──────────────────────────────────────────────────────
folie_heading(4, "Datensatz: 86 CSV-Dateien klassifiziert", "ca. 1 Min")
absatz(
    "Der Datensatz besteht aus 86 CSV-Dateien mit insgesamt rund 1,2 Gigabyte. "
    "Nicht jede Datei ist für unsere Fragestellung relevant — deshalb haben wir "
    "zunächst jede Datei gesichtet und einer von drei Kategorien zugewiesen."
)
absatz(
    "Sieben Dateien haben wir aktiv verwendet und in die Analysetabelle eingebunden — "
    "darunter SO.csv für Stammdaten, QS.Qualitätsindikator.csv als Quelle der Ziel-Variable "
    "und FA.Personalliste.csv für die Ärztedichte. "
    "33 weitere Dateien haben wir als potenziell relevant identifiziert, aber noch nicht eingebunden — "
    "sie wären Kandidaten für Folgeanalysen. "
    "29 Dateien wurden bewusst ausgeschlossen: keine Analyserelevanz, DSGVO-Daten oder reine Lookup-Tabellen. "
    "17 Dateien sind technische Schlüsseltabellen ohne eigene Analysedaten."
)
uebergang("Jetzt schauen wir, welche dieser 7 Dateien wir konkret verwendet haben — und was in ihnen steckt.")

# ── Folie 5 ──────────────────────────────────────────────────────
folie_heading(5, "Die 7 verwendeten Quelldateien", "ca. 1 Min")
absatz(
    "Von den 86 CSV-Dateien haben wir 7 aktiv eingebunden. "
    "Die Tabelle zeigt, welche Datei welche Rolle spielt: drei Dateien liefern strukturelle Merkmale "
    "(SO.csv, QS.Fortbildung.csv, Konzern.csv), zwei liefern die Personalmerkmale "
    "(FA.Personalliste.csv und SO.Personalliste.csv), eine dient als technische Brückentabelle "
    "(FA.csv) — und QS.Qualitätsindikator.csv liefert als einzige die Ziel-Variable."
)
uebergang("Jetzt schauen wir, wie diese Daten inhaltlich zusammenhängen.")

# ── Folie 6 ──────────────────────────────────────────────────────
folie_heading(6, "Zwei Datenarten: Merkmale & Ziel-Variable", "ca. 2 Min")
absatz(
    "Das Herzstück des Projekts ist diese Unterscheidung: Auf der einen Seite stehen die "
    "Strukturdaten — das sind unsere Eingabevariablen. Sie beschreiben, wie ein Krankenhaus "
    "aufgebaut ist: Wie viele Betten hat es? Wer trägt es — privat, öffentlich oder "
    "freigemeinnützig? Ist es eine Uni-Klinik? Wie ist die Personalausstattung? "
    "Diese Daten stammen aus dem sogenannten A-Teil des Qualitätsberichts: SO.csv für "
    "Stammdaten, FA.csv für Fachabteilungen, Konzern.csv für Konzernzugehörigkeit."
)
absatz(
    "Auf der anderen Seite stehen die Qualitätsdaten — das ist unsere Zielvariable. "
    "Sie beschreiben, wie gut die Versorgung ist: Für jeden der rund 150 Qualitätsindikatoren "
    "bekommt ein Haus entweder die Bewertung 'rechnerisch auffällig' oder 'nicht auffällig'. "
    "Diese Daten stecken in QS.Qualitätsindikator.csv — allein diese eine Datei ist 911 Megabyte groß."
)
absatz(
    "Diese Unterscheidung wird später für unser Machine Learning wichtig: "
    "Die Strukturdaten dienen als Eingaben für das Modell, während die Qualitätsdaten "
    "das Ergebnis darstellen, das wir erklären beziehungsweise vorhersagen möchten."
)
absatz(
    "Das Verbindungsstück zwischen beiden Welten ist der Schlüssel SO.QBID — "
    "er erscheint in rund 60 der 86 Dateien und macht es erst möglich, "
    "Strukturdaten und Qualitätsdaten für jedes Krankenhaus zusammenzuführen."
)
uebergang("Aber wie wird aus 150 Einzelbewertungen eine einzige Kennzahl pro Haus? Das zeigt die nächste Folie.")

# ── Folie 5 ──────────────────────────────────────────────────────
folie_heading(7, "Wie entsteht die Ziel-Variable?", "ca. 2 Min")
absatz(
    "Die Berechnung der Ziel-Variable war technisch aufwändig — und es gab drei "
    "klassische Fallstricke, in die man leicht hineintappt."
)
absatz(
    "Erstens: Nicht alle Zeilen in QS.Qualitätsindikator.csv sind echte Qualitätsindikatoren. "
    "Es gibt auch sogenannte Zählkennzahlen — EKez, TKez, KKez — die nur Fallzahlen erfassen "
    "und keine inhaltliche Bewertung darstellen. Diese müssen vor der Analyse herausgefiltert werden."
)
absatz(
    "Zweitens: Der Bewertungscode N99 bedeutet 'nicht bewertet' — nicht etwa 'unauffällig'. "
    "Diese Zeilen werden ausgeschlossen, weil ein fehlendes Urteil kein positives Urteil ist."
)
absatz(
    "Drittens: Ohne Deduplizierung wird jeder Indikator mehrfach gezählt, "
    "weil manche Häuser mehrere Einträge pro Indikator haben. "
    "Die Lösung: Je Krankenhaus und Indikator exakt eine Zeile behalten."
)
absatz(
    "Nach diesen Schritten gilt: R* bedeutet rechnerisch auffällig, N* bedeutet nicht auffällig. "
    "Wir berechnen dann pro Haus die auffaellig_quote — also den Anteil auffälliger Indikatoren. "
    "Der Median dieser Quote liegt bei 76,92 Prozent. Alles oberhalb dieses Medians "
    "gilt als 'viele Probleme', alles darunter als 'wenige Probleme'. "
    "Damit werden die Krankenhäuser in zwei ungefähr gleich große Gruppen aufgeteilt — "
    "eine mit niedrigeren und eine mit höheren Anteilen auffälliger Qualitätsindikatoren."
)
uebergang("Jetzt wissen wir, was wir erklären wollen. Aber was sind die erklärenden Merkmale?")

# ── Folie 6 ──────────────────────────────────────────────────────
folie_heading(8, "Die 8 Merkmale der Analysetabelle", "ca. 2 Min")
absatz(
    "Wir haben aus den 86 Quelldateien acht Merkmale ausgewählt, die wir für besonders "
    "aussagekräftig halten. Vier davon sind direkt strukturell: Bettenzahl als Maß für "
    "die Größe des Hauses, Trägerschaft in drei Kategorien — privat, freigemeinnützig, öffentlich — "
    "Bundesland als regionale Komponente, und der Uni-Status als Indikator für akademische Kliniken."
)
absatz(
    "Dazu kommen vier personalbasierte und organisatorische Merkmale: "
    "Ärzte pro Bett und Pflegekräfte pro Bett als Maße für die Personalintensität, "
    "Fortbildungsquote als Qualitätsmerkmal der Weiterbildung, "
    "und Konzernzugehörigkeit als Hinweis auf die Einbettung in größere Strukturen."
)
absatz(
    "Technisch besonders anspruchsvoll war die Berechnung von Ärzte pro Bett: "
    "Die Personaldaten lagen in zwei verschiedenen Tabellen — FA.Personalliste.csv enthält "
    "die einzelnen Mitarbeiter mit Stellenanteil, kennt aber nur die Fachabteilungs-ID, nicht "
    "das Krankenhaus. FA.csv verbindet Fachabteilung und Krankenhaus. Deshalb waren zwei "
    "Left Joins nötig: erst Personalliste an FA.csv, dann das Ergebnis an SO.csv — "
    "Left Joins, weil sonst Tageskliniken ohne Bettenzahl still aus der Analyse gefallen wären. "
    "Dazu kam ein zweites Problem: Die Stellenanteile waren als Text mit Komma gespeichert, "
    "zum Beispiel '1,75' statt '1.75'. pandas interpretiert das als Zeichenkette, nicht als Zahl — "
    "die Summierung liefert dann lautlos nur NaN-Werte statt echter Zahlen, ohne jede Fehlermeldung. "
    "Erst nach str.replace(',', '.') und pd.to_numeric() stimmten die Ergebnisse."
)
absatz(
    "Das Ergebnis all dieser Arbeit ist eine einzige Datei: analysetabelle.csv — "
    "1.824 Zeilen, 18 Spalten, eine Zeile pro Krankenhaus. Von dieser Tabelle aus "
    "läuft alles andere: die Analyse, das Dashboard, das Machine-Learning-Modell."
)
uebergang("Mit dieser Grundlage können wir jetzt schauen, was die Daten uns tatsächlich sagen.")
trennlinie()

# ═════════════════════════════════════════════════════════════════
# BLOCK 3 — ERGEBNISSE DER ANALYSE  (Folien 7–11)
# ═════════════════════════════════════════════════════════════════
block_heading("Block 3 — Ergebnisse der Analyse  (Folien 9–13 · ca. 10 Min)")

# ── Folie 9 ──────────────────────────────────────────────────────
folie_heading(9, "Befund 1: Verteilung der Auffälligkeitsquote", "ca. 2 Min")
absatz(
    "Bevor wir einzelne Merkmale vergleichen, lohnt sich ein Blick auf die Ziel-Variable selbst: "
    "Wie ist die Auffälligkeitsquote über alle 1.824 Häuser verteilt?"
)
absatz(
    "Das Histogramm zeigt: Die meisten Häuser häufen sich im rechten Bereich — "
    "also bei hohen Auffälligkeitsquoten zwischen 60 und 90 Prozent. "
    "Es gibt kaum Häuser mit sehr niedrigen Werten, dafür eine längere Ausläufer der Verteilung nach links. "
    "Der Median liegt bei 76,92 Prozent — "
    "das klingt auf den ersten Blick erschreckend hoch, ist aber erklärbar: "
    "Die Qualitätsindikatoren des IQTIG sind bewusst anspruchsvoll gesetzt, "
    "um Verbesserungspotenziale sichtbar zu machen. "
    "Wichtig für die Analyse: Es gibt kein extremes Ungleichgewicht — "
    "rund 925 Häuser liegen unterhalb des Medians, rund 899 darüber. "
    "Das ist eine gute Ausgangslage für Gruppenvergleiche."
)
uebergang("Schauen wir uns jetzt an, ob die Trägerschaft eines Hauses mit seiner Auffälligkeitsquote zusammenhängt.")

# ── Folie 8 ──────────────────────────────────────────────────────
folie_heading(10, "Befund 2: Trägerschaft", "ca. 2 Min")
absatz(
    "Das ist unser klarster inhaltlicher Befund: Private Krankenhäuser haben mit 56,5 Prozent "
    "deutlich häufiger viele Qualitätsprobleme als freigemeinnützige mit 46,4 Prozent "
    "oder öffentliche mit 46,7 Prozent."
)
absatz(
    "Um zu prüfen, ob dieser Unterschied statistisch belastbar ist oder nur zufällig, "
    "haben wir eine ANOVA durchgeführt — eine Varianzanalyse. "
    "Die ANOVA beantwortet genau diese Frage: Sind die Unterschiede zwischen drei oder mehr Gruppen "
    "groß genug, um nicht durch puren Zufall erklärbar zu sein? "
    "Das Ergebnis ist ein p-Wert — der gibt an, wie wahrscheinlich es wäre, "
    "diesen Unterschied zu sehen, wenn in Wirklichkeit gar kein Unterschied existiert. "
    "Unser p-Wert liegt unter 0,001 — das bedeutet: die Wahrscheinlichkeit, "
    "dass dieser Unterschied reiner Zufall ist, liegt unter 0,1 Prozent. "
    "Das nennt man statistisch hochsignifikant."
)
absatz(
    "Aber hier ist Vorsicht geboten: Private Häuser sind im Median "
    "deutlich kleiner — rund 90 Betten gegenüber etwa 260 Betten bei öffentlichen Häusern. "
    "Kleinere Häuser behandeln weniger Fälle pro Indikator, das macht die Ergebnisse "
    "schwankungsanfälliger — ein Haus mit 10 Fällen kann leichter 100 Prozent Auffälligkeit "
    "erreichen als ein Haus mit 500 Fällen. Das bedeutet: Was wir als Trägerschaftseffekt sehen, "
    "könnte in Wirklichkeit ein Größeneffekt sein. "
    "Korrelation ist nicht Kausalität — das gilt auch hier."
)
uebergang("Das führt uns direkt zur nächsten Frage: Welches Merkmal hängt am stärksten mit Qualitätsproblemen zusammen?")

# ── Folie 9 ──────────────────────────────────────────────────────
folie_heading(11, "Befund 3: Personal ist der stärkste Prädiktor", "ca. 2 Min")
absatz(
    "Mehr Personal pro Bett geht mit weniger Qualitätsproblemen einher — "
    "das ist der stärkste und robusteste Befund in diesem Projekt. "
    "Häuser mit wenigen Problemen haben im Median 0,468 Ärzte pro Bett, "
    "Häuser mit vielen Problemen nur 0,390 — ein Unterschied von rund 17 Prozent."
)
absatz(
    "Wir haben das mit einem T-Test überprüft. Der T-Test vergleicht zwei Gruppen — hier: "
    "Häuser mit wenigen vs. vielen Problemen — und testet, ob der Unterschied im Mittelwert "
    "groß genug ist, um nicht durch Zufall erklärbar zu sein. "
    "Der t-Wert von 6,002 beschreibt dabei, wie viele Standardabweichungen der Gruppenunterschied "
    "vom Zufall entfernt ist — je größer, desto klarer das Signal. "
    "Der p-Wert liegt auch hier unter 0,001: statistisch hochsignifikant. "
    "Das gleiche Muster zeigt sich bei den Pflegekräften pro Bett mit ähnlicher Stärke."
)
absatz(
    "Umso überraschender ist ein anderer Befund: Die Fortbildungsquote — also der Anteil "
    "der Mitarbeiter, die an qualitätssichernden Fortbildungen teilgenommen haben — "
    "zeigt praktisch keinen Zusammenhang mit der Auffälligkeitsquote. "
    "Die Korrelation liegt bei etwa 0,01 — das ist so nah an null, dass man von keinem "
    "messbaren Zusammenhang sprechen kann. Das ist kein Messfehler, das ist ein Ergebnis. "
    "Manche Erwartungen bestätigen sich, manche nicht — beides gehört zu den Ergebnissen unserer Analyse."
)
uebergang("Neben Trägerschaft und Personal haben wir auch regionale Unterschiede und den Uni-Status untersucht.")

# ── Folie 10 ──────────────────────────────────────────────────────
folie_heading(12, "Befund 4: Bundesland & Uni-Status", "ca. 2 Min")
absatz(
    "Wir haben zwei weitere Merkmale untersucht: Bundesland und Uni-Status — "
    "beides Merkmale, bei denen wir eigentlich einen Zusammenhang erwartet hätten."
)
absatz(
    "Beim Bundesland sehen wir auf der Grafik große Unterschiede: Im Saarland haben 63 Prozent "
    "der dortigen Krankenhäuser viele Qualitätsprobleme, in Berlin nur 33 Prozent. "
    "Aber: Im Saarland haben wir nur 19 Häuser. "
    "Wenn dort eines mehr auffällig ist, verschiebt sich der Balken um 5 Prozentpunkte. "
    "Das heißt: Ein einzelnes Krankenhaus mehr oder weniger in der roten Gruppe "
    "reicht aus, um den Prozentwert des gesamten Bundeslandes um 5 Punkte zu verändern. "
    "Das ist kein belastbares Ergebnis — das ist Zufall durch kleine Stichprobe."
)
absatz(
    "Beim Uni-Status haben wir erwartet, dass Unikliniken besser abschneiden — "
    "mehr Forschung, mehr Fachärzte. Tatsächlich: 47 Prozent Unikliniken auffällig, "
    "49 Prozent normale Häuser. Praktisch identisch. "
    "Der Grund: Unikliniken behandeln die schwersten Fälle. "
    "Die Grenzwerte der Qualitätsindikatoren wurden aber anhand typischer Standardeingriffe festgelegt — "
    "also für einen durchschnittlichen Patienten. "
    "Wer ausschließlich Hochrisikopatienten behandelt, überschreitet diese Grenzwerte öfter, "
    "nicht weil die Qualität schlechter ist, sondern weil die Ausgangslage schwieriger ist."
)
absatz(
    "Für Konzernhäuser haben wir den Chi-Quadrat-Test genutzt. "
    "Der prüft bei zwei Ja/Nein-Gruppen — hier: Konzern ja/nein und viele Probleme ja/nein — "
    "ob die Verteilung zwischen den Gruppen vom Zufall abweicht. "
    "Das Ergebnis: p = 0,90. Zur Erinnerung: p-Wert nahe 0 heißt, der Unterschied ist kein Zufall. "
    "p-Wert nahe 1 heißt das Gegenteil — die Verteilung zwischen Konzernhäusern und unabhängigen Häusern "
    "ist so ähnlich, dass sie genauso gut zufällig entstanden sein könnte. Kein Zusammenhang."
)
uebergang("Diese einzelnen Befunde zusammenzufassen ist die Aufgabe der nächsten Folie.")

# ── Folie 11 ──────────────────────────────────────────────────────
folie_heading(13, "Gesamtübersicht: Korrelationsmatrix", "ca. 2 Min")
absatz(
    "Die Korrelationsmatrix fasst alle Zusammenhänge auf einmal zusammen — als Heatmap. "
    "Jede Zahl darin ist ein r, ein sogenannter Korrelationskoeffizient. "
    "r misst, wie stark zwei Merkmale zusammenhängen, auf einer Skala von −1 bis +1. "
    "r = 0 bedeutet kein Zusammenhang. r = 1 bedeutet: steigt das eine, steigt auch das andere. "
    "r = −1 bedeutet: steigt das eine, sinkt das andere. "
    "Alles zwischen −0,3 und +0,3 gilt in der Praxis als schwacher Zusammenhang."
)
absatz(
    "Das wichtigste Ergebnis: Alle unsere r-Werte liegen in diesem schwachen Bereich. "
    "Den rechnerisch stärksten Zusammenhang zeigt die Anzahl der bewerteten Indikatoren — "
    "r = −0,28. Das ist aber kein inhaltlicher Befund, sondern ein methodisches Artefakt: "
    "Größere Häuser haben mehr Indikatoren und performen im Durchschnitt stabiler."
)
absatz(
    "Von den inhaltlich interessanten Merkmalen haben Ärzte pro Bett und Pflegekräfte "
    "pro Bett je r = −0,14 — schwach, aber der stärkste inhaltliche Zusammenhang. "
    "Das Minuszeichen sagt dabei: mehr Personal pro Bett geht mit weniger Problemen einher. "
    "Bettenzahl: r = −0,08. Fortbildungsquote: r ≈ 0,01. Konzern: r ≈ 0,00 — "
    "praktisch null, kein messbarer Zusammenhang."
)
absatz(
    "Das ist das zentrale Ergebnis dieses Projekts: Die verfügbaren Strukturmerkmale "
    "erklären die Qualitätsunterschiede zwischen Krankenhäusern nur sehr begrenzt. "
    "Andere Faktoren — Patientenstruktur, Dokumentationsverhalten, regionale Besonderheiten — "
    "spielen wahrscheinlich eine viel größere Rolle. Das macht dieses Ergebnis nicht "
    "wertlos — es zeigt klar, wo die Grenzen strukturbasierter Erklärungsansätze liegen."
)
uebergang("Wir haben die Erkenntnisse aus dieser Analyse genutzt, um ein Machine-Learning-Modell zu trainieren.")
trennlinie()

# ═════════════════════════════════════════════════════════════════
# BLOCK 4 — MACHINE LEARNING & DASHBOARD  (Folien 12–13)
# ═════════════════════════════════════════════════════════════════
block_heading("Block 4 — Machine Learning & Dashboard  (Folien 14–15 · ca. 6 Min)")

# ── Folie 12 ──────────────────────────────────────────────────────
folie_heading(14, "Decision Tree: Modell & Ergebnis", "ca. 3 Min")
absatz(
    "Als nächsten Schritt haben wir einen Entscheidungsbaum trainiert. "
    "Wir haben uns bewusst für ein einfaches Modell entschieden, das man Schritt für Schritt "
    "nachvollziehen kann — kein Black-Box-Algorithmus, sondern ein klarer Entscheidungspfad. "
    "Die maximale Tiefe ist auf drei Ebenen begrenzt: Das Modell stellt höchstens "
    "drei Fragen hintereinander, bevor es zu einer Einschätzung kommt."
)
absatz(
    "Die erste und wichtigste Frage lautet: Hat das Haus weniger als 0,271 Ärzte pro Bett? "
    "Dieser Schwellenwert wurde vom Modell automatisch ermittelt und bestätigt, was wir "
    "in der deskriptiven Analyse gesehen haben: Ärzte pro Bett ist der stärkste Prädiktor. "
    "Er trägt 53,6 Prozent zur Vorhersagequalität bei. "
    "Pflegekräfte pro Bett kommen auf 23,8 Prozent, Bettenzahl auf 22,6 Prozent."
)
absatz(
    "Die Accuracy liegt bei 63,6 Prozent — das Modell trifft in knapp zwei von drei Fällen die richtige Einschätzung. "
    "Zum Vergleich: Reines Raten würde 50,7 Prozent erreichen. "
    "Das Modell ist also besser als Zufall, aber kein starkes Modell. "
    "Das passt zum R²-Wert von 0,033: Die Strukturmerkmale erklären nur 3,3 Prozent "
    "der Unterschiede zwischen den Häusern. Das Modell zeigt, was mit diesen Daten möglich ist — "
    "und macht damit die Grenze der verfügbaren Information sichtbar."
)
uebergang("Alle diese Ergebnisse haben wir in einem interaktiven Dashboard zusammengeführt, das wir jetzt kurz live zeigen.")

# ── Folie 13 ──────────────────────────────────────────────────────
folie_heading(15, "Streamlit-Dashboard — Live-Demo", "ca. 3 Min")
absatz(
    "Das Dashboard ist in Python gebaut — mit Streamlit als Web-Framework und Plotly "
    "für die interaktiven Grafiken. Die Datenbasis ist analysetabelle.csv, das Modell ist "
    "das trainierte scikit-learn-Entscheidungsbaum-Objekt. "
    "Für die Präsentation läuft das Dashboard über Streamlit Community Cloud: "
    "Der Code liegt auf GitHub, Streamlit liest das Repository automatisch ein und "
    "stellt die App als öffentliche URL bereit — kein lokales Setup nötig, "
    "nur ein Browser und der Link."
)
absatz(
    "Das Dashboard hat vier Seiten. Auf der ersten Seite — Gesamtüberblick — sieht man "
    "die wichtigsten Kennzahlen auf einen Blick: Wie viele Häuser gibt es, wie viele "
    "davon haben viele Probleme, wie hoch ist die durchschnittliche Auffälligkeitsquote. "
    "Dazu gibt es eine interaktive Deutschlandkarte, auf der jedes Haus als farbiger "
    "Punkt erscheint — grün für wenige Probleme, rot für viele. "
    "Filter für Bundesland, Träger und Kliniktyp erlauben gezielte Auswertungen."
)
absatz(
    "Die zweite Seite — Einflussfaktoren — ist die wichtigste für unsere Projektfrage. "
    "Hier sieht man direkt, ob sich Gruppen unterscheiden: "
    "Trägerschaftsvergleich, Boxplots für Personal pro Bett, Streudiagramme. "
    "Die dritte Seite ermöglicht es, ähnliche Häuser zu finden und "
    "einen Einzelhaus-Steckbrief anzuzeigen — ideal für konkrete Vergleiche. "
    "Die vierte Seite ist der Risiko-Rechner: Wir geben Merkmale ein, "
    "und das Modell gibt eine Einschätzung aus — inklusive der Entscheidungsbaum-Visualisierung, "
    "die zeigt, wie die Entscheidung zustande gekommen ist."
)
absatz(
    "[Live-Demo des Dashboards — ca. 3 Minuten]",
    italic=True, farbe=ORANGE, size=10
)
uebergang("Jetzt ziehen wir ein Gesamtbild: Was haben wir gelernt?")
trennlinie()

# ═════════════════════════════════════════════════════════════════
# BLOCK 5 — FAZIT & GRENZEN  (Folien 14–15)
# ═════════════════════════════════════════════════════════════════
block_heading("Block 5 — Fazit & Grenzen  (Folien 16–17 · ca. 4 Min)")

# ── Folie 14 ──────────────────────────────────────────────────────
folie_heading(16, "Gesamteinschätzung", "ca. 2 Min")
absatz(
    "Die Ampel-Tabelle fasst zusammen, was wir gefunden haben — und was nicht. "
    "Grün, also schwacher aber nachweisbarer Zusammenhang: Personal pro Bett. "
    "Sowohl Ärzte als auch Pflegekräfte zeigen einen statistisch signifikanten, "
    "wenn auch schwachen negativen Zusammenhang mit Qualitätsproblemen. "
    "Das ist inhaltlich plausibel und methodisch belastbar."
)
absatz(
    "Gelb, also sichtbar aber mit Einschränkungen: Trägerschaft, Bundesland und Bettenzahl. "
    "Diese Merkmale zeigen Unterschiede, die aber durch Störfaktoren wie Hausgröße "
    "oder kleine Stichproben beeinflusst sein können."
)
absatz(
    "Rot, also kein messbarer Zusammenhang: Uni-Status, Konzernzugehörigkeit und "
    "Fortbildungsquote. Diese drei Merkmale zeigen in keiner unserer Analysen "
    "einen bedeutsamen Zusammenhang mit Qualitätsproblemen."
)
absatz(
    "Das Gesamtfazit lautet: Wir finden keine starken, eindeutigen Zusammenhänge. "
    "Die Strukturmerkmale allein reichen nicht aus, um Qualitätsprobleme zu erklären. "
    "Das ist ein valides Ergebnis — und der ehrlichste Befund, den wir präsentieren können. "
    "Eine Analyse, die nur positive Befunde zeigt, wäre keine ehrliche Wissenschaft."
)
uebergang("Zum Abschluss reflektieren wir, wo diese Analyse ihre Grenzen hat — und was als nächstes möglich wäre.")

# ── Folie 15 ──────────────────────────────────────────────────────
folie_heading(17, "Grenzen & Ausblick", "ca. 2 Min")
absatz(
    "Jede Analyse hat Grenzen, und wir möchten die drei wichtigsten benennen."
)
absatz(
    "Erstens: Patientenmix. Ein Krankenhaus, das besonders schwere oder seltene Fälle behandelt, "
    "wird häufiger auffällig bewertet — nicht weil die Qualität schlechter ist, sondern weil "
    "die Indikatoren auf typische Normalfälle kalibriert sind. Diesen Effekt können wir "
    "mit unseren Daten nicht herausrechnen."
)
absatz(
    "Zweitens: Dokumentationsqualität. Manche rechnerischen Auffälligkeiten entstehen "
    "nicht durch schlechte Versorgung, sondern durch fehlerhafte oder unvollständige "
    "Dokumentation. Wir können diese Fälle nicht von echten Qualitätsproblemen trennen."
)
absatz(
    "Drittens — und das ist der grundlegende Punkt: Korrelation ist nicht Kausalität. "
    "Selbst die schwachen Zusammenhänge, die wir gefunden haben, sagen nichts darüber aus, "
    "ob ein bestimmtes Merkmal Qualitätsprobleme verursacht oder nur mit ihnen zusammenhängt."
)
absatz(
    "Für mögliche nächste Schritte hätten wir drei konkrete Ansätze: "
    "Erstens die Dokumentationsrate aus QS.Leistungsbereich.csv als weiteres Merkmal einbinden — "
    "das würde den zweiten Störfaktor teilweise kontrollieren. "
    "Zweitens eine Sentiment-Analyse der Freitext-Kommentare: "
    "QSQI.KommentarKrankenhaus enthält individuelle Aussagen der Häuser, "
    "die NLP-Methoden zugänglich wären. "
    "Drittens eine mehrstufige Regression, die Patientenstruktur als Kontrollvariable einschließt."
)
absatz(
    "Das Dashboard und alle Notebooks sind vollständig reproduzierbar und können "
    "direkt erweitert werden — der Code liegt strukturiert vor und ist dokumentiert."
)
absatz(
    "Damit sind wir am Ende unserer Präsentation. Wir freuen uns auf Ihre Fragen — "
    "ob zum methodischen Vorgehen, zu den Ergebnissen oder zum Dashboard.",
    fett=False, farbe=BLAU, size=11
)
trennlinie()

# ═════════════════════════════════════════════════════════════════
# ZEITPLAN-ANHANG
# ═════════════════════════════════════════════════════════════════
doc.add_page_break()
absatz("Zeitplan-Übersicht", fett=True, farbe=BLAU, size=12, abstand_vor=0, abstand_nach=8)

zeitplan = [
    ("Block 1", "Einstieg & Projektfrage",    "Folien 1–3",  "~4 Min"),
    ("Block 2", "Datenvorbereitung",           "Folien 4–8",  "~8 Min"),
    ("Block 3", "Ergebnisse der Analyse",      "Folien 9–13", "~10 Min"),
    ("Block 4", "Machine Learning & Dashboard","Folien 14–15","~6 Min"),
    ("Block 5", "Fazit & Grenzen",             "Folien 16–17","~4 Min"),
    ("Gesamt",  "",                            "17 Folien",   "~32 Min"),
]
tbl = doc.add_table(rows=1, cols=4)
tbl.style = "Table Grid"
hdr = tbl.rows[0].cells
for i, h in enumerate(["Block", "Inhalt", "Folien", "Zeit"]):
    p = hdr[i].paragraphs[0]
    p.clear()
    run = p.add_run(h)
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    hdr[i]._tc.get_or_add_tcPr()
for cell in tbl.rows[0].cells:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "1F497D")
    tcPr.append(shd)

for row_data in zeitplan:
    row = tbl.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val
        if row_data[0] == "Gesamt":
            run = row[i].paragraphs[0].runs
            if run:
                run[0].bold = True

# ── Speichern ─────────────────────────────────────────────────────
doc.save(ZIEL)
print(f"Dokument gespeichert: {ZIEL}")
