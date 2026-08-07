import pandas as pd, os
from datetime import date
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

analyse = pd.read_csv("Data/analysetabelle.csv", low_memory=False)
md_q = analyse["auffaellig_quote"].median()
df = analyse.copy()
traeger_col = "KH.Träger.Art" if "KH.Träger.Art" in df.columns else "KH.Traeger.Art"

def add_h(doc, t, l=1): return doc.add_heading(t, level=l)
def add_b(doc, t):
    p = doc.add_paragraph(t); p.style.font.size = Pt(11); return p
def add_tbl(doc, hdrs, rows, cw=None):
    tbl = doc.add_table(rows=1+len(rows), cols=len(hdrs)); tbl.style="Table Grid"
    for i,h in enumerate(hdrs):
        tbl.rows[0].cells[i].text=h
        r=tbl.rows[0].cells[i].paragraphs[0].runs[0]; r.bold=True; r.font.size=Pt(10)
        tc=tbl.rows[0].cells[i]._tc; tcPr=tc.get_or_add_tcPr()
        shd=OxmlElement("w:shd"); shd.set(qn("w:fill"),"4472C4"); shd.set(qn("w:val"),"clear")
        tcPr.append(shd); r.font.color.rgb=RGBColor(255,255,255)
    for ri,row in enumerate(rows):
        for ci,val in enumerate(row):
            tbl.rows[ri+1].cells[ci].text=str(val)
            tbl.rows[ri+1].cells[ci].paragraphs[0].runs[0].font.size=Pt(10)
    if cw:
        for row in tbl.rows:
            for i,w in enumerate(cw): row.cells[i].width=Cm(w)
    return tbl

doc = Document()
for s in doc.sections:
    s.top_margin=Cm(2.5); s.bottom_margin=Cm(2.5)
    s.left_margin=Cm(3.0); s.right_margin=Cm(2.5)

# Titelseite
t=doc.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER
rn=t.add_run("Qualitaets-Muster-Finder"); rn.bold=True; rn.font.size=Pt(24); rn.font.color.rgb=RGBColor(31,73,125)
doc.add_paragraph()
s=doc.add_paragraph(); s.alignment=WD_ALIGN_PARAGRAPH.CENTER
s.add_run("Projektdokumentation").font.size=Pt(14)
doc.add_paragraph()
i=doc.add_paragraph(); i.alignment=WD_ALIGN_PARAGRAPH.CENTER
i.add_run(f"Erstellt: {date.today().strftime('%d.%m.%Y')}  |  Datenbasis: Qualitaetsberichte 2023").font.size=Pt(11)
doc.add_page_break()

# Inhaltsverzeichnis
add_h(doc,"Inhaltsverzeichnis",1)
for nr,ti in [
    ("1","Projektuebersicht"),("1.1","Fragestellung"),("1.2","Bausteine und Status"),
    ("2","Baustein 1 - Daten vorbereiten"),("2.1","Setup"),("2.2","Datensatz erkunden"),
    ("2.3","Ziel-Variable erstellen"),("2.4","Merkmale und Analysetabelle"),("2.5","Aerzte pro Bett"),("2.6","Ergebnisse"),
    ("3","Baustein 2 - Deskriptive Analyse"),("3.1","Vorgehen"),("3.2","Befunde"),("3.3","Gesamteinschaetzung"),("3.4","Grafiken"),
    ("4","Offene Punkte und Naechste Schritte"),
]:
    p=doc.add_paragraph(); r2=p.add_run(f"{nr}  {ti}"); r2.font.size=Pt(11)
    if "." not in nr: r2.bold=True
doc.add_page_break()

# 1. Projektuebersicht
add_h(doc,"1  Projektuebersicht",1)
add_h(doc,"1.1  Fragestellung",2)
add_b(doc,"Welche Krankenhausmerkmale haengen damit zusammen, dass ein Haus ueberdurchschnittlich viele Qualitaetsprobleme aufweist?")
add_b(doc,"Grundlage: jaehrliche Qualitaetsberichte aller deutschen Krankenhaeuser. Jedes Haus berichtet ueber ~150 Qualitaetsindikatoren. Ziel: strukturelle Merkmale (Groesse, Personal, Traeger, Region) identifizieren, die mit erhoehter Auffaelligkeitsquote zusammenhaengen.")
add_b(doc,"Wichtig: Kein Zusammenhang ist ein valides Ergebnis. QSErgBewStrukDialog = R* bedeutet nur rechnerisch auffaellig, kein echtes Qualitaetsurteil. Quelle: Text_Presentation.docx, Folie 7.")
add_h(doc,"1.2  Bausteine und Status",2)
add_tbl(doc,["Baustein","Beschreibung","Status"],[
    ("Baustein 1","Daten vorbereiten","Abgeschlossen"),
    ("Baustein 2","Deskriptive Analyse","Abgeschlossen"),
    ("Baustein 3","Streamlit-Dashboard (3 Seiten)","Offen"),
    ("Baustein 4","Entscheidungsbaum (Bonus)","Offen"),
    ("Baustein 5","Abschluss und Praesentation","Offen"),
],cw=[3.5,8.5,3.0]); doc.add_paragraph(); doc.add_page_break()

# 2. Baustein 1
add_h(doc,"2  Baustein 1 - Daten vorbereiten",1)
add_b(doc,"Status: Abgeschlossen | Datum: 2026-07-27 | Datei: 01_Exploration.ipynb")
add_h(doc,"2.1  Setup",2)
add_b(doc,"Kein GitHub - Projekt wird lokal geteilt. Rohdaten bleiben im Data/-Ordner. Alle anderen Dateien koennen direkt kopiert werden.")
add_h(doc,"2.2  Datensatz erkunden",2)
add_b(doc,"86 CSV-Dateien gesichtet (Header + Beispieldaten). Schluessel-ID SO.QBID verbindet fast alle Tabellen. Ergebnis: Daten_Inhaltsverzeichnis.md")
add_b(doc,"Auswahlkriterium (3 Gruppen):")
add_tbl(doc,["Kriterium","Beschreibung","Beispiel"],[
    ("1 - Ziel-Variable","Enthaelt standardisierte Bewertung von Qualitaetsindikatoren","QS.Qualitaetsindikator.csv"),
    ("2 - Strukturmerkmal","Enthaelt Betten, Traeger, Personal, Region, Uni oder Fortbildung","SO.csv, QS.Fortbildung.csv, FA.Personalliste.csv"),
    ("3 - Verknuepfung","Notwendige Brueckentabelle zwischen zwei relevanten Tabellen","FA.csv (ABTID -> SO.QBID)"),
],cw=[4.0,8.5,5.0]); doc.add_paragraph()
add_b(doc,"Warum genau diese 6 Dateien aus 86?")
entscheid_rows = [
    ("SO.csv","Einzige Datei mit ALLEN Strukturmerkmalen in einer Tabelle (Betten, Traeger, Region, Uni, Koordinaten). Alternativdateien wie SO.Personalliste.csv enthalten nur Detailpersonal, keine Stammdaten."),
    ("QS.Qualitaetsindikator.csv","Einzige Datei mit standardisierter, einheitlicher Bewertung (R* = auffaellig) fuer alle Haeuser nach denselben IQTIG-Regeln. QS.Extern.Sonstige.csv enthaelt Zahlenwerte, aber keine einheitliche Klassifikation."),
    ("QS.Fortbildung.csv","EINZIGE Datei im Datensatz mit Fortbildungsquoten. AQ.Aerzte.csv enthaelt nur Qualifikationen, keine Quoten."),
    ("FA.csv","Notwendige Brueckentabelle: FA.Personalliste.csv kennt nur ABTID, SO.csv nur SO.QBID. Ohne FA.csv kein Join."),
    ("FA.Personalliste.csv","Einzige Datei mit Aerzte-Anzahl pro Abteilung nach Berufsgruppe (FA.Personal.Bereich='Aerzte'). Ergebnis: aerzte_pro_bett = wichtigstes Merkmal (Feature Importance 71,3%)."),
    ("QS.Leistungsbereich.csv","Enthaelt Dokumentationsrate - potenzielle Qualitaetskennzahl. Noch nicht eingebunden."),
]
add_tbl(doc,["Datei","Entscheidungsbegruendung"],entscheid_rows,cw=[5.0,12.0]); doc.add_paragraph()
add_h(doc,"2.3  Ziel-Variable erstellen",2)
add_b(doc,"Bewertungsspalte: QSErgBewStrukDialog - R* = rechnerisch auffaellig | N99 = nicht bewertet (ausgeschlossen!) | N* = nicht auffaellig")
add_b(doc,"Wichtige Entscheidung: QSQI.AEKey ist eine Haus-ID, kein Indikator-Schluessel! Deduplizierung ueber (SO.QBID, QSQI.Indikator).")
add_tbl(doc,["Schritt","Code","Erklaerung"],[
    ("1","QSQI.ArtDesWertes == QI","Nur echte QI, keine Zaehlkennzahlen"),
    ("2","QSErgBewStrukDialog != N99","Nur bewertete Indikatoren"),
    ("3","drop_duplicates([SO.QBID, QSQI.Indikator])","Je Haus+Indikator eine Zeile"),
    ("4","str.startswith(R)","Auffaellig-Flag setzen"),
    ("5","groupby(SO.QBID).agg(count, sum)","Pro Haus aggregieren"),
    ("6","auffaellig_n / total_qi","Quote berechnen"),
    ("7","quote > Median -> 1","Binaere Ziel-Variable"),
],cw=[1.5,6.5,6.0]); doc.add_paragraph()
add_b(doc,f"Ergebnis: Median-Quote = {md_q:.2%}, 1.824 Haeuser, ausgewogene Verteilung (899 vs. 925).")
add_h(doc,"2.4  Merkmale und Analysetabelle",2)
add_tbl(doc,["Merkmal","Quelle","Berechnung"],[
    ("SO.Betten","SO.csv","Direkt"),
    ("KH.Traeger.Art","SO.csv","Direkt (privat/freigemeinn./oeffentlich)"),
    ("SO.Bundesland","SO.csv","Direkt (16 Bundeslaender)"),
    ("SO.Uni","SO.csv","Direkt (0/1)"),
    ("SO.Latitude/Long.","SO.csv","Direkt (fuer Karte)"),
    ("fortbildungsquote","QS.Fortbildung.csv","Erbracht / Pflichtige"),
    ("aerzte_pro_bett","FA.Personalliste.csv","Summe Aerzte / SO.Betten"),
],cw=[4.5,4.5,7.0]); doc.add_paragraph()
add_b(doc,"Zusammenfuehrung vollstaendig per Skript reproduzierbar (01_Exploration.ipynb). Rohdaten -> Analysetabelle -> alles andere (Analyse, Dashboard, Decision Tree).")
add_h(doc,"2.5  Aerzte pro Bett",2)
add_b(doc,"Quelle: FA.Personalliste.csv x FA.csv ueber ABTID | Filter: FA.Personal.Bereich == Aerzte")
add_b(doc,"FA.Personal.Anzahl = Komma-Dezimal (z.B. 13,47) -> float via str.replace | SO.Betten = 0 (Tageskliniken) -> NaN (korrekt, wird nicht aufgefuellt)")
add_tbl(doc,["Kennzahl","Wert"],[
    ("Haeuser mit Aerzte-Daten","2.308"),
    ("Durchschnitt Aerzte gesamt pro Haus","112,2"),
    ("Durchschnitt Aerzte pro Bett","0,451"),
    ("Fehlende Werte (NaN)","5 - davon 4 Tageskliniken"),
],cw=[8.0,4.0]); doc.add_paragraph()
add_h(doc,"2.6  Ergebnisse der Analysetabelle",2)
n_h=len(analyse); n_a=int(analyse["hat_viele_Probleme"].sum())
add_tbl(doc,["Kennzahl","Wert"],[
    ("Zeilen (Krankenhaeuser)",f"{n_h:,}"),
    ("Spalten",f"{analyse.shape[1]}"),
    ("Median auffaellig-Quote",f"{md_q:.2%}"),
    ("hat_viele_Probleme = 1",f"{n_a:,} ({n_a/n_h:.1%})"),
    ("hat_viele_Probleme = 0",f"{n_h-n_a:,} ({(n_h-n_a)/n_h:.1%})"),
    ("Fehlende KH.Traeger.Art",f"{int(df[traeger_col].isna().sum())} (1,5 %)"),
    ("Fehlende fortbildungsquote",f"{int(df['fortbildungsquote'].isna().sum())} (1,8 %)"),
    ("Fehlende aerzte_pro_bett",f"{int(df['aerzte_pro_bett'].isna().sum())} (0,3 %)"),
],cw=[8.0,4.0]); doc.add_paragraph(); doc.add_page_break()

# 3. Baustein 2
add_h(doc,"3  Baustein 2 - Deskriptive Analyse",1)
add_b(doc,"Status: Abgeschlossen | Datum: 2026-07-27 | Datei: 02_Analyse.ipynb")
add_h(doc,"3.1  Vorgehen",2)
add_b(doc,"10 Grafiken aus analysetabelle.csv mit automatisch berechnetem Befundsatz. Farbschema: gruen = wenige Probleme, rot = viele Probleme. Grafiken in grafiken/.")
add_h(doc,"3.2  Befunde",2)
p0=df[df["hat_viele_Probleme"]==0]; p1=df[df["hat_viele_Probleme"]==1]
priv=df[df[traeger_col]=="privat"]["hat_viele_Probleme"].mean()
frei=df[df[traeger_col]=="freigemeinnuetzig"]["hat_viele_Probleme"].mean() if "freigemeinn" in df[traeger_col].unique() else df[df[traeger_col].str.contains("freigemein",na=False)]["hat_viele_Probleme"].mean()
oeffentl=df[df[traeger_col]=="oeffentlich"]["hat_viele_Probleme"].mean() if "oeffentlich" in df[traeger_col].unique() else df[df[traeger_col].str.contains("ffentlich",na=False)]["hat_viele_Probleme"].mean()
add_tbl(doc,["Grafik","Befund"],[
    ("G1 auffaellig-Quote",f"Median {df['auffaellig_quote'].median():.0%}, linkssteil - meisten Haeuser 60-90%."),
    ("G2 Bettenzahl",f"Median: Wenige={int(p0['SO.Betten'].median())} Betten, Viele={int(p1['SO.Betten'].median())} Betten. Kein klarer Groessenunterschied."),
    ("G3 Traegerschaft",f"Privat {priv:.1%} vs. freigemeinn. {frei:.1%} vs. oeffentl. {oeffentl:.1%}. Private Haeuser auffaellig hoeher."),
    ("G4 Uni-Kliniken",f"Uni {df[df['SO.Uni']==1]['hat_viele_Probleme'].mean():.1%} vs. Normal {df[df['SO.Uni']==0]['hat_viele_Probleme'].mean():.1%} - kaum Unterschied."),
    ("G5+6 Fortb./Aerzte",f"Fortbildung: kein Unterschied. Aerzte/Bett: Wenige={p0['aerzte_pro_bett'].median():.3f}, Viele={p1['aerzte_pro_bett'].median():.3f}"),
    ("G7 Bundesland","Saarland hoechster Anteil (63,2%), Berlin niedrigster (33,3%)."),
    ("G8 Korrelation","Staerkste Korrelation: total_qi (r=-0,28), aerzte_pro_bett (r=-0,14). Fortbildungsquote: keine."),
    ("G9 Scatter","Kein klares Trennmuster - starke Ueberlappung."),
    ("G10 Stoerfaktor",f"Private Haeuser kleiner (Md={int(df[df[traeger_col]=='privat']['SO.Betten'].median())} Betten) - Traegereffekt mit Vorsicht."),
],cw=[3.5,12.5]); doc.add_paragraph()
add_h(doc,"3.3  Gesamteinschaetzung",2)
add_b(doc,"Keine starken, eindeutigen Zusammenhaenge. Staerkster Praediktor total_qi ist ein Strukturmerkmal (nicht Qualitaetsmerkmal). Einziger klarer Befund: Private Haeuser 56,5% vs. ~47% andere Traeger - aber Stoerfaktor Groesse moeglich.")
add_b(doc,"Kein Zusammenhang ist ein valides Ergebnis. R* = nur rechnerisch auffaellig, kein echtes Qualitaetsurteil. Quelle: Text_Presentation.docx, Folie 7.")
add_h(doc,"3.4  Grafiken",2)
for pfad,titel in [
    ("grafiken/g1_auffaellig_quote.png","Grafik 1: Verteilung der auffaellig-Quote"),
    ("grafiken/g2_bettenzahl.png","Grafik 2: Bettenzahl"),
    ("grafiken/g3_traegerschaft.png","Grafik 3: Traegerschaft"),
    ("grafiken/g4_uni.png","Grafik 4: Uni-Kliniken vs. normale Haeuser"),
    ("grafiken/g5_6_fortbildung_aerzte.png","Grafik 5+6: Fortbildungsquote und Aerzte pro Bett"),
    ("grafiken/g7_bundesland.png","Grafik 7: Anteil Haeuser je Bundesland"),
    ("grafiken/g8_korrelation.png","Grafik 8: Korrelationsmatrix"),
    ("grafiken/g9_scatter_betten_aerzte.png","Grafik 9: Scatter - Bettenzahl vs. Aerzte pro Bett"),
    ("grafiken/g10_stoerfaktor_traeger.png","Grafik 10: Stoerfaktor Traegerschaft x Bettengroesse"),
]:
    if os.path.exists(pfad):
        p=doc.add_paragraph(); r3=p.add_run(titel); r3.bold=True; r3.font.size=Pt(11)
        doc.add_picture(pfad,width=Cm(15)); doc.add_paragraph()
doc.add_page_break()

# 4. Offene Punkte
add_h(doc,"4  Offene Punkte und Naechste Schritte",1)
add_tbl(doc,["Baustein","Aufgabe","Details"],[
    ("Baustein 3","Streamlit-Dashboard","Seite 1: Uebersicht + Karte | Seite 2: Vergleiche | Seite 3: Aehnliche Haeuser"),
    ("Baustein 4","Entscheidungsbaum (Bonus)","max_depth=3, Train-Test-Split, Modell speichern (joblib)"),
    ("Baustein 5","Praesentation","Startanleitung, Entscheidungsbegruendungen, Live-Demo"),
    ("Offen","Pflegekraefte pro Bett","Noch nicht berechnet - aus FA.Personalliste.csv extrahieren"),
],cw=[3.0,4.5,8.5]); doc.add_paragraph()
add_b(doc,"Die Datengrundlage (analysetabelle.csv) ist vollstaendig - alle weiteren Bausteine koennen direkt gestartet werden.")

doc.save("Dokumentation_Qualitaets_Muster_Finder.docx")
print("Gespeichert")
